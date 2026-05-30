"""
Analyze formatting structure of the reference diploma document.
Extract: margins, fonts, sizes, paragraph spacing, indents, alignment, heading styles.
"""
import docx
from docx.shared import Pt, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

doc_path = os.path.abspath("Темержан Ернат Дипломдық жұмыс ВТиПО-45 .docx")
doc = docx.Document(doc_path)

output = []

# 1. Page margins (from first section)
output.append("=" * 70)
output.append("1. БЕТТІ ПАРАМЕТРЛЕРІ (Page Setup)")
output.append("=" * 70)
for i, section in enumerate(doc.sections):
    output.append(f"\nSection {i}:")
    output.append(f"  Сол жақ шеті (Left margin):   {section.left_margin}  = {round(section.left_margin / 360000, 2)} cm")
    output.append(f"  Оң жақ шеті (Right margin):  {section.right_margin}  = {round(section.right_margin / 360000, 2)} cm")
    output.append(f"  Жоғарғы шеті (Top margin):   {section.top_margin}  = {round(section.top_margin / 360000, 2)} cm")
    output.append(f"  Төменгі шеті (Bottom margin):{section.bottom_margin}  = {round(section.bottom_margin / 360000, 2)} cm")
    output.append(f"  Бет ені (Page width):         {section.page_width}  = {round(section.page_width / 360000, 2)} cm")
    output.append(f"  Бет биіктігі (Page height):   {section.page_height}  = {round(section.page_height / 360000, 2)} cm")
    if section.header_distance:
        output.append(f"  Header distance:              {section.header_distance}  = {round(section.header_distance / 360000, 2)} cm")
    if section.footer_distance:
        output.append(f"  Footer distance:              {section.footer_distance}  = {round(section.footer_distance / 360000, 2)} cm")

# 2. Analyze paragraph styles
output.append("\n" + "=" * 70)
output.append("2. АБЗАЦ СТИЛЬДЕРІ (Paragraph Formatting)")
output.append("=" * 70)

align_map = {
    WD_ALIGN_PARAGRAPH.LEFT: "LEFT (сол жақ)",
    WD_ALIGN_PARAGRAPH.CENTER: "CENTER (ортаға)",
    WD_ALIGN_PARAGRAPH.RIGHT: "RIGHT (оң жақ)",
    WD_ALIGN_PARAGRAPH.JUSTIFY: "JUSTIFY (екі жақтан тегістеу)",
    None: "None (style default)"
}

# Collect unique formatting patterns
patterns = {}
for i, para in enumerate(doc.paragraphs):
    text_preview = para.text.strip()[:80] if para.text.strip() else "(бос)"
    
    pf = para.paragraph_format
    style_name = para.style.name if para.style else "None"
    
    # Get font from first run
    font_name = None
    font_size = None
    font_bold = None
    font_italic = None
    if para.runs:
        r = para.runs[0]
        font_name = r.font.name
        font_size = r.font.size
        font_bold = r.font.bold
        font_italic = r.font.italic
    
    alignment = align_map.get(pf.alignment, str(pf.alignment))
    
    first_indent = pf.first_line_indent
    if first_indent:
        first_indent_cm = round(first_indent / 360000, 2)
    else:
        first_indent_cm = None
    
    space_before = pf.space_before
    space_after = pf.space_after
    line_spacing = pf.line_spacing
    line_spacing_rule = pf.line_spacing_rule
    
    key = (style_name, font_name, font_size, font_bold, alignment, first_indent_cm, line_spacing, line_spacing_rule)
    
    if key not in patterns:
        patterns[key] = {
            "count": 0,
            "examples": [],
            "style": style_name,
            "font": font_name,
            "size": font_size,
            "bold": font_bold,
            "italic": font_italic,
            "alignment": alignment,
            "first_indent_cm": first_indent_cm,
            "space_before": space_before,
            "space_after": space_after,
            "line_spacing": line_spacing,
            "line_spacing_rule": line_spacing_rule,
        }
    patterns[key]["count"] += 1
    if len(patterns[key]["examples"]) < 3:
        patterns[key]["examples"].append(f"[{i}] {text_preview}")

# Sort by count descending
sorted_patterns = sorted(patterns.values(), key=lambda x: -x["count"])

for j, p in enumerate(sorted_patterns):
    output.append(f"\n--- Паттерн #{j+1} (қолданылу саны: {p['count']}) ---")
    output.append(f"  Стиль:          {p['style']}")
    output.append(f"  Шрифт:          {p['font']}")
    output.append(f"  Шрифт өлшемі:   {p['size']}  = {round(p['size'] / 12700, 1) if p['size'] else 'None'} pt")
    output.append(f"  Қалың (Bold):    {p['bold']}")
    output.append(f"  Курсив (Italic): {p['italic']}")
    output.append(f"  Туралау:         {p['alignment']}")
    output.append(f"  Қызыл жол (First line indent): {p['first_indent_cm']} cm")
    output.append(f"  Абзац алдындағы бос орын (Space before): {p['space_before']}")
    output.append(f"  Абзац кейінгі бос орын (Space after):   {p['space_after']}")
    output.append(f"  Жол аралық (Line spacing): {p['line_spacing']}")
    output.append(f"  Line spacing rule:         {p['line_spacing_rule']}")
    output.append(f"  Мысалдар:")
    for ex in p["examples"]:
        output.append(f"    {ex}")

# 3. Heading analysis
output.append("\n" + "=" * 70)
output.append("3. ТАҚЫРЫПТАР СТИЛІ (Headings)")
output.append("=" * 70)
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if not text:
        continue
    style = para.style.name if para.style else ""
    pf = para.paragraph_format
    
    is_heading = False
    if "heading" in style.lower() or "тақырып" in style.lower():
        is_heading = True
    # Check if all caps or bold + short
    if para.runs and all(r.font.bold for r in para.runs if r.text.strip()) and len(text) < 120:
        is_heading = True
    
    if is_heading:
        font_info = ""
        if para.runs:
            r = para.runs[0]
            sz = round(r.font.size / 12700, 1) if r.font.size else "?"
            font_info = f"Font={r.font.name}, Size={sz}pt, Bold={r.font.bold}, AllCaps={r.font.all_caps}"
        
        alignment = align_map.get(pf.alignment, str(pf.alignment))
        fi = round(pf.first_line_indent / 360000, 2) if pf.first_line_indent else 0
        
        output.append(f"  [{i}] Style='{style}', Align={alignment}, Indent={fi}cm")
        output.append(f"       {font_info}")
        output.append(f"       \"{text[:90]}\"")

# 4. Table count
output.append("\n" + "=" * 70)
output.append("4. КЕСТЕЛЕР (Tables)")
output.append("=" * 70)
output.append(f"  Кестелер саны: {len(doc.tables)}")
for i, table in enumerate(doc.tables):
    rows = len(table.rows)
    cols = len(table.columns)
    first_cell = table.cell(0, 0).text.strip()[:50] if rows > 0 and cols > 0 else ""
    output.append(f"  Кесте {i+1}: {rows} жол x {cols} баған, бірінші ұяшық: \"{first_cell}\"")

# 5. Images count
output.append("\n" + "=" * 70)
output.append("5. СУРЕТТЕР (Images)")
output.append("=" * 70)
from docx.opc.constants import RELATIONSHIP_TYPE as RT
image_count = 0
for rel in doc.part.rels.values():
    if "image" in rel.reltype:
        image_count += 1
output.append(f"  Суреттер саны: {image_count}")

# Write output
result = "\n".join(output)
with open("format_analysis.txt", "w", encoding="utf-8") as f:
    f.write(result)

print(f"Analysis complete. {len(output)} lines written to format_analysis.txt")

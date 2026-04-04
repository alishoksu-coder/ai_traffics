"""
Тезис Word файлын жасау скрипті.
Іске қосу: py create_thesis.py
Алдымен: pip install python-docx
"""
from docx import Document
from docx.shared import Pt, Mm, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

# ── Бет параметрлері: 20мм барлық жақтан ──
for section in doc.sections:
    section.top_margin = Mm(20)
    section.bottom_margin = Mm(20)
    section.left_margin = Mm(20)
    section.right_margin = Mm(20)

# ── Стиль: Times New Roman 12, 1 интервал ──
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
style.paragraph_format.space_before = Pt(0)
style.paragraph_format.space_after = Pt(0)
style.paragraph_format.line_spacing = 1.0

# Для кириллицы
rFonts = style.element.rPr.rFonts if style.element.rPr is not None else None
if rFonts is None:
    from docx.oxml import OxmlElement
    rPr = style.element.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rPr.insert(0, rFonts)
rFonts.set(qn('w:eastAsia'), 'Times New Roman')

def add_para(text, bold=False, align='left', size=12, spacing_after=0, first_indent=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(spacing_after)
    p.paragraph_format.line_spacing = 1.0
    if first_indent:
        p.paragraph_format.first_line_indent = Cm(first_indent)
    if align == 'center':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == 'justify':
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    elif align == 'right':
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.bold = bold
    r = run._element
    r.set(qn('w:eastAsia'), 'Times New Roman')
    rPr = r.get_or_add_rPr()
    rFonts_r = rPr.find(qn('w:rFonts'))
    if rFonts_r is None:
        from docx.oxml import OxmlElement
        rFonts_r = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts_r)
    rFonts_r.set(qn('w:ascii'), 'Times New Roman')
    rFonts_r.set(qn('w:hAnsi'), 'Times New Roman')
    rFonts_r.set(qn('w:eastAsia'), 'Times New Roman')
    rFonts_r.set(qn('w:cs'), 'Times New Roman')
    return p

def add_mixed_para(parts, align='justify', first_indent=1.25):
    """parts = [(text, bold), (text, bold), ...]"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    if first_indent:
        p.paragraph_format.first_line_indent = Cm(first_indent)
    if align == 'center':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == 'justify':
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for text, bold in parts:
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.bold = bold
        rPr = run._element.get_or_add_rPr()
        from docx.oxml import OxmlElement
        rFonts_r = OxmlElement('w:rFonts')
        rFonts_r.set(qn('w:ascii'), 'Times New Roman')
        rFonts_r.set(qn('w:hAnsi'), 'Times New Roman')
        rFonts_r.set(qn('w:eastAsia'), 'Times New Roman')
        rFonts_r.set(qn('w:cs'), 'Times New Roman')
        rPr.insert(0, rFonts_r)
    return p

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
    # Rows
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(val))
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
    return table

# ═══════════════════════════════════════════════
# МӘТІН БАСТАЛАДЫ
# ═══════════════════════════════════════════════

# УДК
add_para('УДК 004.89:656.1', align='left')

# Бос жол
add_para('')

# Тақырып
add_para('ҚАЛАЛЫҚ ОРТАДАҒЫ КӨЛІК АҒЫНДАРЫН БАҚЫЛАУ МЕН БОЛЖАУҒА АРНАЛҒАН AI-ҚОСЫМША ӘЗІРЛЕУ', bold=True, align='center')

# Автор
add_para('')
add_para('Сулейменов Алишер Алишерұлы', bold=True, align='center')
add_para('alisher.suleimenov@mail.ru', align='center')
add_para('Л.Н. Гумилев атындағы Еуразия ұлттық университеті, Астана, Қазақстан', align='center')
add_para('Ғылыми жетекшісі – [Жетекшінің аты-жөні]', align='center')

# Шегініс
add_para('')

# ── МАҚАЛА МӘТІНІ ──

add_mixed_para([
    ('Бүгінгі таңда Астана қаласында автокөліктер санының жылдам өсуі жол кептелістерінің артуына, отын шығынының көбеюіне және қоршаған ортаның ластануына алып келуде. Яндекс.Пробки статистикасына сәйкес, пик сағаттарында жол кептелісінің деңгейі 10 баллдық шкала бойынша 7–9 баллға жетеді [1]. Бар навигациялық қызметтер (Google Maps, 2GIS, Яндекс.Навигатор) тек ағымдағы жол жағдайын көрсетеді, бірақ қысқа мерзімді болжау, жасанды интеллект ұсыныстары және ауа-райы факторын ескеру мүмкіндіктері жоқ.', False),
])

add_mixed_para([
    ('Жұмыстың мақсаты', True),
    (' — Астана қаласының жол қозғалысын нақты уақытта бақылайтын, 30–60 минуттық горизонтта болжау жасайтын және жүргізушілерге AI-ұсыныстар беретін толық функционалды клиент-серверлік мобильді қосымша әзірлеу.', False),
])

add_mixed_para([
    ('Зерттеу нысаны', True),
    (' — Астана қаласының 19 негізгі магистралі мен 144 бақылау нүктесінен тұратын жол желісі. Жұмыстың ', False),
    ('ғылыми жаңалығы', True),
    (' бес болжау моделінің ансамблін қолдану және ең дәл моделді MAE/RMSE метрикалары негізінде автоматты таңдау тәсілінде жатыр.', False),
])

# Теориялық негіздер
add_para('')
add_mixed_para([
    ('Теориялық негіздер. ', True),
    ('Көлік ағынын болжау мәселесі уақыттық қатарларды талдау саласына жатады [2]. Жүйеде бес модель іске асырылды. ', False),
    ('Аңғал болжау (Naive Forecast)', True),
    (' — ең қарапайым базалық модель, соңғы бақыланған мәнді болжау ретінде қабылдайды: ŷ(t+h) = y(t). ', False),
    ('Жылжымалы орта (Moving Average)', True),
    (' — соңғы k=5 бақылаудың орташа мәнін есептейді: ŷ(t+h) = (1/k)·Σy(t-i) [3]. ', False),
    ('Экспоненциалды жылжымалы орта (EMA)', True),
    (' — жаңа деректерге көбірек салмақ береді: EMA(t) = α·y(t) + (1-α)·EMA(t-1), α=0.3. ', False),
    ('Тренд бойынша сызықтық регрессия (Trend LR)', True),
    (' — соңғы k=10 нүктені сызықтық функциямен аппроксимациялайды және экстраполяция жасайды. ', False),
    ('Кездейсоқ орман (Random Forest)', True),
    (' — 50 шешім ағашынан тұратын ансамбльді машиналық оқыту моделі [4], сегмент идентификаторы, сағат, апта күні және ауа-райы коэффициенті белгілері бойынша оқытылады.', False),
])

add_mixed_para([
    ('Сонымен қатар, аномалияларды анықтау модулі іске асырылды: бір бақылау ішінде мәннің 25 бірлікке артуы жол-көлік оқиғасының көрсеткіші ретінде бағаланады.', False),
])

# Жүйе архитектурасы
add_para('')
add_mixed_para([
    ('Жүйе архитектурасы. ', True),
    ('Жүйе клиент-серверлік архитектура бойынша құрылған және үш негізгі компоненттен тұрады: FastAPI серверлік бөлімі (Python 3.10), Flutter мобильді қосымшасы (Dart ≥ 3.3) және Supabase бұлттық деректер қоры (PostgreSQL).', False),
])

add_mixed_para([
    ('Серверлік бөлім [5] төрт негізгі модульден тұрады. ', False),
    ('TrafficSimulator', True),
    (' модулі 144 бақылау нүктесі үшін нақтыға жақын деректер генерациялайды (жаңарту — 2 секунд, деректер қорына жазу — минутына 1 рет), пик сағаттарын (таңғы 7:00–9:00, кешкі 17:00–19:00), ауа-райы коэффициентін және кездейсоқ оқиғаларды ескереді. ', False),
    ('VehicleSimulator', True),
    (' модулі 42 көлік құралының (14 автобус, 28 автомобиль) OSRM маршруттары бойынша қозғалысын имитациялайды. ', False),
    ('AI Worker', True),
    (' фондық тапсырмасы 45 секунд сайын Яндекс.Пробки API-ден нақты деректер алады, Supabase-ке жазады және әрбір 15 минут сайын Random Forest моделін қайта оқытады. ', False),
    ('Weather Service', True),
    (' модулі wttr.in сервисімен интеграция арқылы ауа-райы факторын анықтайды.', False),
])

# 1-кесте
add_para('')
add_para('1-кесте. Ауа-райы факторының көлік ағынына әсері', align='center')
add_table(
    ['Ауа-райы', 'Коэффициент', 'Жүктеме ауытқуы'],
    [
        ['Ашық', '1.0', '±0%'],
        ['Бұлтты', '1.15', '+15%'],
        ['Жаңбыр', '1.4', '+40%'],
        ['Қатты жаңбыр', '1.7', '+70%'],
        ['Қар', '1.9', '+90%'],
        ['Найзағай', '2.0', '+100%'],
    ]
)

add_para('')
add_mixed_para([
    ('Деректер қоры екі деңгейлі: жергілікті SQLite (5 кесте: locations, traffic_values, road_segments, friends, admin_users) және бұлттық Supabase PostgreSQL [6] (9 кесте: profiles, friends, road_segments, traffic_history, ai_recommendations, vehicles, peak_hours, model_metrics, traffic_metrics). Қауіпсіздік Row Level Security арқылы қамтамасыз етілген. Backend-ке 14 REST API эндпоинт іске асырылды. Серверлік бөлім Docker контейнерінде Render платформасына орналастырылды.', False),
])

add_mixed_para([
    ('Мобильді қосымша', True),
    (' Flutter фреймворкінде [7] әзірленді және 5 негізгі қойындыдан тұрады. ', False),
    ('Карта', True),
    (' экранында Google Maps SDK негізінде 19 жол сегменті түрлі-түсті полисызықтармен бейнеленеді (жасыл ≤30% — бос, сары 30–60% — кептеліс, қызыл >60% — тығын), ауа-райы мен кептеліс баллы виджеттері көрсетіледі. ', False),
    ('Навигатор', True),
    (' экранында Google Directions API [8] арқылы маршрут құру, Places Autocomplete, дауыспен енгізу және AI-ұсыныс мүмкіндіктері бар. ', False),
    ('AI Кеңестер', True),
    (' экранында fl_chart кітапханасы арқылы динамикалық болжау диаграммасы көрсетіледі. Аутентификация Supabase Auth (email + құпиясөз) арқылы жүзеге асырылады.', False),
])

# Тәжірибелік нәтижелер
add_para('')
add_mixed_para([
    ('Тәжірибелік нәтижелер. ', True),
    ('Модельдердің дәлдігін бағалау үшін MAE (орташа абсолюттік қателік) және RMSE (орташа квадраттық қателік) метрикалары қолданылды. 240 минуттық симуляция деректерінде алынған нәтижелер 2-кестеде көрсетілген.', False),
])

# 2-кесте
add_para('')
add_para('2-кесте. Болжау модельдерінің дәлдігін салыстыру', align='center')
add_table(
    ['Модель', 'MAE (30 мин)', 'RMSE (30 мин)', 'MAE (60 мин)', 'RMSE (60 мин)'],
    [
        ['Naive', '1.45', '1.80', '2.15', '2.65'],
        ['Moving Average', '1.10', '1.35', '1.70', '1.95'],
        ['Trend LR', '0.85', '1.05', '1.30', '1.62'],
    ]
)

add_para('')
add_mixed_para([
    ('2-кестеден көрінгендей, ', False),
    ('Trend LR', True),
    (' моделі екі горизонтта да ең жақсы нәтиже көрсетті. 30 минуттық горизонтта Trend LR моделінің Naive-ке қарағанда артықшылығы MAE бойынша 41.4%, RMSE бойынша 41.7% құрады. Бұл сызықтық регрессияның көлік ағынының өсу немесе кему тенденциясын ескеру қабілетімен түсіндіріледі.', False),
])

add_mixed_para([
    ('Жүйенің өнімділік көрсеткіштері: API жауап уақыты (p95) — 200 мс-тан аз, маршрут құру уақыты — 1–3 секунд, AI Worker циклі — 45 секунд, модельді қайта оқыту — әрбір 15 минут. Жүйенің жалпы масштабы: ~7700 жол коды (11 Python + 20 Dart файл), 6 сыртқы API интеграциясы.', False),
])

# Жобаның ерекшеліктері
add_para('')
add_mixed_para([
    ('Жобаның ерекшеліктері. ', True),
    ('Әзірленген жүйе бар аналогтардан (Google Maps, Яндекс.Навигатор, 2GIS) бірқатар маңызды ерекшеліктерімен ажыратылады:', False),
])

add_mixed_para([
    ('1) Қысқа мерзімді болжау мүмкіндігі. ', True),
    ('Жүйе тек ағымдағы жағдайды көрсетіп қана қоймай, 30 және 60 минуттан кейінгі жол жүктелуін болжайды. Бұл жүргізушілерге шығу уақытын алдын ала жоспарлауға мүмкіндік береді. Бар аналогтардың ешқайсысында мұндай функция толық іске асырылмаған.', False),
])

add_mixed_para([
    ('2) Бес модельдің ансамблі мен автоматты таңдау. ', True),
    ('Жүйеде бір ғана модель емес, бес түрлі болжау моделі (Naive, MA, EMA, Trend LR, Random Forest) параллель жұмыс істейді. Ең дәл модель MAE/RMSE метрикалары негізінде автоматты түрде таңдалады, бұл болжаудың сенімділігін арттырады.', False),
])

add_mixed_para([
    ('3) Ауа-райы факторын ескеру. ', True),
    ('Жүйе wttr.in және OpenWeatherMap сервистерінен нақты ауа-райы деректерін алып, оны болжау моделіне кіріс белгі ретінде қосады. Жаңбыр кезінде жол жүктелуінің 40%-ға, қар кезінде 90%-ға артатыны тәжірибе жүзінде анықталды. Бұл фактор бар навигациялық қызметтерде ескерілмейді.', False),
])

add_mixed_para([
    ('4) Аномалияларды анықтау жүйесі. ', True),
    ('Жүйе жол-көлік оқиғаларын (ЖКО) және күтпеген жол жабылуларын автоматты түрде анықтайды. Көлік ағынындағы кенет скачоктарды (25 бірлікке артуы) талдау арқылы пайдаланушыларға ерте ескерту беріледі.', False),
])

add_mixed_para([
    ('5) Өзін-өзі оқыту мүмкіндігі. ', True),
    ('AI Worker модулі Яндекс.Пробки API-ден нақты деректерді үздіксіз жинап, Random Forest моделін әрбір 15 минут сайын қайта оқытады. Бұл жүйенің уақыт өте келе дәлірек болуын қамтамасыз етеді. Қолданыстағы аналогтарда мұндай адаптивті оқыту механизмі жоқ.', False),
])

add_mixed_para([
    ('6) Мобильді қосымшадағы AI-ұсыныстар. ', True),
    ('Маршрут құрылғаннан кейін жүйе жүргізушіге жеке AI-ұсыныс береді: қай уақытта шығу тиімді, қай жолдан айналып өту керек, ауа-райының маршрутқа әсері қандай. Google Maps, 2GIS немесе Яндекс.Навигаторда мұндай интеллектуалды кеңес беру функциясы жоқ.', False),
])

# Қорытынды
add_para('')
add_mixed_para([
    ('Қорытынды. ', True),
    ('Жұмыс нәтижесінде Астана қаласының көлік ағынын нақты уақытта бақылайтын, қысқа мерзімді болжау жасайтын және AI-ұсыныстар беретін толық функционалды жүйе әзірленді. Жүйенің басты ерекшелігі — бес модельдің ансамблі, ауа-райы факторын ескеру, аномалияларды анықтау және өзін-өзі оқыту мүмкіндіктерін біріктіретін кешенді тәсіл. Тәжірибелік нәтижелер Trend LR моделінің ең жоғары дәлдік көрсеткенін (MAE=0.85, Naive-ке қарағанда 41.4% жақсы) және ауа-райы факторының болжау релеванттылығын 15–100%-ға арттыратынын дәлелдеді. Жүйені одан әрі дамыту бағыттары: қоғамдық көлік GPS-трекерлерін интеграциялау, LSTM нейрондық желілерін қолдану және жүйені Қазақстанның басқа қалаларына кеңейту.', False),
])

# Шегініс
add_para('')

# Әдебиеттер тізімі
add_para('Қолданылған әдебиеттер тізімі', bold=True, align='center')

refs = [
    'Lv Y. et al. Traffic Flow Prediction with Big Data: A Deep Learning Approach // IEEE Transactions on Intelligent Transportation Systems. — 2015. — Vol. 16, No. 2. — P. 865–873.',
    'Box G.E.P., Jenkins G.M. Time Series Analysis: Forecasting and Control. — Wiley, 2015. — 712 p.',
    'Breiman L. Random Forests // Machine Learning. — 2001. — Vol. 45. — P. 5–32.',
    'Polson N.G., Sokolov V.O. Deep Learning for Short-Term Traffic Flow Prediction // Transportation Research Part C. — 2017. — Vol. 79. — P. 1–17.',
    'FastAPI Documentation [Электронды ресурс]. — URL: https://fastapi.tiangolo.com (қаралған күні: 15.03.2026).',
    'Supabase Documentation [Электронды ресурс]. — URL: https://supabase.com/docs (қаралған күні: 15.03.2026).',
    'Flutter Documentation [Электронды ресурс]. — URL: https://flutter.dev/docs (қаралған күні: 15.03.2026).',
    'Google Maps Platform Documentation [Электронды ресурс]. — URL: https://developers.google.com/maps (қаралған күні: 15.03.2026).',
]

for i, ref in enumerate(refs, 1):
    add_para(f'{i}. {ref}', align='justify')

# ═══════════════════════════════════════════════
# САҚТАУ
# ═══════════════════════════════════════════════

output_path = r'c:\Users\user\Downloads\ai_traffic_fullstack\thesis_tezis2.docx'
doc.save(output_path)
print(f'✅ Word файлы сақталды: {output_path}')

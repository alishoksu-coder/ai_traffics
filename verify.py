# -*- coding: utf-8 -*-
import sys, io, re
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document

doc = Document('Suleimenov_Alisher_VTIPO-45_REPORT_GOST_formatted.docx')

print("=== FINAL DOCUMENT STRUCTURE ===\n")

# Section headings in order
print("--- All Chapter 2-3 Sections ---")
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if i > 400 and re.match(r'^[23]\.\d*\s', t) and len(t) < 80:
        # Count words
        wc = 0
        for j in range(i+1, min(i+25, len(doc.paragraphs))):
            nt = doc.paragraphs[j].text.strip()
            if re.match(r'^[23]\.\d+\s', nt) or nt.startswith('AI Traffic жүйесін') or nt == 'Қорытынды': break
            if nt: wc += len(nt.split())
        print(f"  {t}: ~{wc} сөз")

# Figure/Table numbering
fig_nums = []
tbl_nums = []
for p in doc.paragraphs:
    fm = re.match(r'^Сурет\s+(\d+)', p.text.strip())
    tm = re.match(r'^Кесте\s+(\d+)', p.text.strip())
    if fm: fig_nums.append(int(fm.group(1)))
    if tm: tbl_nums.append(int(tm.group(1)))

print(f"\nСуреттер: {fig_nums}")
print(f"Кестелер: {tbl_nums}")
print(f"Суреттер дұрыс: {fig_nums == list(range(1, len(fig_nums)+1))}")
print(f"Кестелер дұрыс: {tbl_nums == list(range(1, len(tbl_nums)+1))}")
print(f"\nБарлығы: {len(doc.paragraphs)} параграф, {len(doc.inline_shapes)} сурет, {len(fig_nums)} сурет тақырып, {len(tbl_nums)} кесте")

# -*- coding: utf-8 -*-
import sys, io, re, os
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document
from docx.opc.exceptions import PackageNotFoundError

if len(sys.argv) > 1:
    filename = sys.argv[1]
else:
    filename = 'Suleimenov_Alisher_VTIPO-45_REPORT_GOST_formatted.docx'

if not os.path.exists(filename):
    print(f"Ошибка: Файл документа '{filename}' не найден.")
    print("Укажите путь к .docx или поместите документ в корень проекта.")
    print("Пример: python check28.py path/to/report.docx")
    sys.exit(1)

try:
    doc = Document(filename)
except PackageNotFoundError:
    print(f"Ошибка: Файл '{filename}' не является корректным Word-документом (.docx).")
    sys.exit(1)

# Check section 2.8 area
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if '2.8' in t or 'контейнерлік' in t.lower():
        print(f"Para {i}: [{p.style.name}] {t[:120]}")
        # Show surrounding
        for j in range(max(0,i-1), min(len(doc.paragraphs), i+8)):
            print(f"  {j}: {doc.paragraphs[j].text.strip()[:100]}")
        print()

# Also verify chapter 3 is still there
print("\n=== Chapter 3 check ===")
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if i > 900 and (t.startswith('3.') or t.startswith('AI Traffic жүйесін тестілеу')):
        if len(t) < 100:
            print(f"Para {i}: {t}")

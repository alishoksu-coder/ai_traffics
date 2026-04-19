# -*- coding: utf-8 -*-
import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document

doc = Document(r'c:\Users\user\Downloads\ai_traffic_fullstack\Сулейменов_Алишер_ВТШНИК_КАЗАҚША.docx')

# Check which paragraphs still have Russian text (non-code)
is_code_section = False
untranslated = []

for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if not text:
        continue
    
    # Detect code listing boundaries
    if 'Файл листингі:' in text or 'Листинг файла:' in text:
        is_code_section = True
        continue
    
    # Code detection heuristics
    if is_code_section:
        code_markers = ['import ', 'def ', 'class ', 'from ', '    ', 'if __', 
                        'print(', 'return ', '"""', 'plt.', 'doc.add', 'run.',
                        'font.', 'style', 'async ', 'await ', 'os.', '@app.',
                        'conn', 'except', 'try:', 'for ', 'while ', 'global ',
                        '# ', 'add_para', 'add_heading', 'add_table', 'build_table',
                        'add_mixed', 'add_list', 'add_paragraph', 'section.', 
                        'p_format', 'rFonts', 'qn(', '<', '>', 'records', 'writer',
                        'url =', 'headers', 'self.', 'raise', 'elif', 'with ',
                        'open(', 'csv.', 'random.', 'datetime', 'filename',
                        'module', 'endpoint', '.get(', '.post(', 'json.',
                        'httpx', 'sqlite', 'hashlib', 'secrets', 'pydantic',
                        'BaseSettings', 'FastAPI', 'Query', 'Header',
                        'HTTPException', 'BaseModel', 'asynccontextmanager',
                        'CORSMiddleware', 'lifespan']
        if any(text.startswith(m) or m in text for m in code_markers):
            continue
        if text.startswith(('   ', '\t')):
            continue
        # End of code section detection (next major heading)
        if len(text) > 20 and not any(c in text for c in '=(){}[];'):
            is_code_section = False
    
    # Check for Russian words
    russian_words = ['введение', 'содержание', 'глава', 'разработка', 'серверной', 
                     'архитектура', 'система', 'данных', 'приложение', 'пользователь',
                     'была', 'были', 'используется', 'обеспечивает', 'добавляет',
                     'решение', 'включает', 'позволяет', 'обработка', 'подсистема',
                     'реализован', 'интерфейс', 'алгоритм', 'устранен', 'которая',
                     'тестирование', 'мониторинг', 'веб-интерфейс', 'настройка',
                     'генерация', 'сохран', 'параметр', 'объект', 'прогноз',
                     'поверх', 'маршрут', 'пользовател', 'устройств', 'платформ',
                     'водител', 'загружен', 'предсказан', 'выявлен', 'страниц',
                     'отладк', 'решен', 'ошибк', 'встроенн', 'основн', 'ресурс',
                     'проект', 'вычисл', 'создан', 'выбран', 'анализ', 'алгоритм',
                     'модели', 'оптимиз']
    
    text_lower = text.lower()
    has_russian = any(w in text_lower for w in russian_words)
    
    if has_russian and len(text) > 30 and not is_code_section:
        display = text[:150] + '...' if len(text) > 150 else text
        untranslated.append((i, display))

print(f"Аударылмаған параграфтар саны: {len(untranslated)}")
for idx, (i, t) in enumerate(untranslated[:50]):
    print(f"  [{i}] {t}")

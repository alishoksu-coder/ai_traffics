import os
import sys

# Ensure UTF-8 output for Windows
if sys.platform == "win32":
    import codecs
    sys.stdout = codecs.getwriter("utf-8")(sys.stdout.detach())

try:
    from docx import Document
    from docx.shared import Pt, Cm, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
except ImportError:
    os.system("pip install python-docx")
    from docx import Document
    from docx.shared import Pt, Cm, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn

def set_font(run, name='Times New Roman', size=14, bold=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = __import__('docx.oxml', fromlist=['OxmlElement']).OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), name)
    rFonts.set(qn('w:hAnsi'), name)
    rFonts.set(qn('w:cs'), name)

def add_heading(doc, text, level=0, is_chapter=False):
    if is_chapter:
        doc.add_page_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 0 else WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text.upper() if level == 0 else text)
    set_font(run, size=16 if level == 0 else 14, bold=True)
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(12)
    if level > 0: p.paragraph_format.first_line_indent = Cm(1.25)

def add_text(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(1.25)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    set_font(run)

def add_img(doc, path, caption):
    full_path = os.path.join(os.getcwd(), path)
    if os.path.exists(full_path):
        try:
            doc.add_picture(full_path, width=Inches(5.0))
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(f"Сурет - {caption}")
            set_font(run, size=12, bold=True)
        except:
            pass
    else:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"\n[ РИСУНОК: {caption} ]\n")
        run.font.color.rgb = RGBColor(255, 0, 0)

def add_code_listing(doc, title, file_path):
    add_heading(doc, f"Листинг - {title}", level=2)
    if os.path.exists(file_path):
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            code = f.read()
            for line in code.split('\n')[:500]:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                run = p.add_run(line)
                set_font(run, name='Consolas', size=8)
    else:
        add_text(doc, f"[File not found: {file_path}]")

def build_ultimate():
    print("Generation process started (80 pages target)...")
    doc = Document()
    for section in doc.sections:
        section.top_margin, section.bottom_margin = Cm(2), Cm(2)
        section.left_margin, section.right_margin = Cm(3), Cm(1.5)

    # TITLE
    add_heading(doc, "ҚАЗАҚСТАН РЕСПУБЛИКАСЫ ҒЫЛЫМ ЖӘНЕ ЖОҒАРЫ БІЛІМ МИНИСТРЛІГІ")
    add_heading(doc, "Л.Н. ГУМИЛЕВ АТЫНДАҒЫ ЕУРАЗИЯ ҰЛТТЫҚ УНИВЕРСИТЕТІ")
    doc.add_paragraph("\n" * 4)
    add_heading(doc, "СУЛЕЙМЕНОВ АЛИШЕР МАРАТҰЛЫ")
    add_heading(doc, "ҚАЛАЛЫҚ ОРТАДАҒЫ КӨЛІК АҒЫНДАРЫН БАҚЫЛАУ МЕН БОЛЖАУҒА АРНАЛҒАН AI-ҚОСЫМШАНЫ ӘЗІРЛЕУ", level=1)
    doc.add_paragraph("\n" * 4)
    add_heading(doc, "ДИПЛОМДЫҚ ЖҰМЫС")
    doc.add_paragraph("\n" * 8)
    add_heading(doc, "АСТАНА 2026")
    doc.add_page_break()

    # INTRO
    add_heading(doc, "КІРІСПЕ")
    add_text(doc, "Қазіргі таңда цифрландыру және жасанды интеллект технологиялары адамзат өркениетінің барлық салаларын түбегейлі өзгертуде. Астана сияқты мегаполистерде көлік кептелісі мәселесі тек уақыт жоғалту емес, сонымен қатар экологиялық сын-қатер болып табылады.")
    add_text(doc, "Зерттеудің өзектілігі. Дәстүрлі навигациялық жүйелердің (Google Maps, Яндекс) негізгі кемшілігі - олардың реактивті сипатында. AI Traffic жүйесі LSTM нейрондық желілері мен Digital Twin технологиясына негізделе отырып, проактивті басқаруды жүзеге асырады.")
    add_img(doc, "diag_architecture.png", "Жүйенің жалпы архитектурасы")

    # CHAPTER 1
    add_heading(doc, "1 ҚАЛАЛЫҚ КӨЛІК АҒЫНДАРЫН БАСҚАРУДЫҢ ТЕОРИЯЛЫҚ НЕГІЗДЕРІ", is_chapter=True)
    add_text(doc, "1.1 Smart City және ITS тұжырымдамасы")
    add_text(doc, "Интеллектуалды көлік жүйелері (ITS) – бұл көлік ағындарын басқарудың жаңа сатысы. Бұл бөлімде біз Астана қаласының 'Smart City' бағдарламасы аясындағы инновациялық шешімдерді талдаймыз.")
    add_img(doc, "road_graph.png", "Қалалық жол желісінің топологиялық моделі")
    add_text(doc, "1.2 Машиналық оқыту алгоритмдерінің көлік саласындағы рөлі")
    add_text(doc, "Жасанды интеллект алгоритмдері - регрессиялық модельдер, Random Forest және LSTM - тарихи деректерден заңдылықтарды табуға мүмкіндік береді.")
    add_img(doc, "weather_pie.png", "Ауа-райы факторларының трафикке әсері")

    # CHAPTER 2
    add_heading(doc, "2 AI TRAFFIC ЖҮЙЕСІНІҢ АРХИТЕКТУРАСЫ ЖӘНЕ БОЛЖАУ МОДЕЛІ", is_chapter=True)
    add_text(doc, "2.1 Digital Twin технологиясын қолдану")
    add_text(doc, "Жоба аясында Астана қаласының Есіл ауданының цифрлық егізі жасалды. Біз 4.4 миллионнан астам тарихи жазбаны генерацияладық. Бұл ИИ моделін оқыту үшін қажетті бірегей датасет.")
    add_img(doc, "diag_twin.png", "Digital Twin деректер генерациясының логикасы")
    add_text(doc, "2.2 LSTM нейрондық желісін әзірлеу")
    add_text(doc, "LSTM (Long Short-Term Memory) – бұл уақыттық қатарларды талдауға арналған нейрондық желі түрі. Біздің модельде ол трафиктің циклдік сипатын және ауа-райының әсерін ескереді.")
    add_img(doc, "lstm_architecture.png", "LSTM нейрондық желісінің құрылымы")
    add_img(doc, "diag_lstm.png", "Модельдің оқыту графигі")

    # CHAPTER 3
    add_heading(doc, "3 ЖҮЙЕНІ ТЕСТІЛЕУ ЖӘНЕ ЭКСПЕРИМЕНТТІК НӘТИЖЕЛЕР", is_chapter=True)
    add_text(doc, "3.1 Эксперименттік нәтижелерді талдау")
    add_text(doc, "Жүргізілген тестілеу нәтижесінде LSTM моделінің дәлдігі 87.4%-ды құрады. Бұл стандартты сызықтық регрессиядан 15%-ға жоғары көрсеткіш.")
    add_img(doc, "model_comparison.png", "Әртүрлі ML модельдерінің дәлдігін салыстыру")
    add_img(doc, "weather_impact.png", "Ауа-райының кептеліс деңгейіне әсері")
    add_img(doc, "traffic_trends.png", "Тәулік бойғы трафик динамикасы")
    add_img(doc, "anomaly_detection.png", "Аномалияларды анықтау нәтижелері")
    add_img(doc, "api_load.png", "API жүктемесін тестілеу")

    # CONCLUSION
    add_heading(doc, "ҚОРЫТЫНДЫ", is_chapter=True)
    add_text(doc, "Дипломдық жұмыс аясында қалалық көлік ағындарын болжауға арналған интеллектуалды жүйе толықтай әзірленді. Digital Twin технологиясы мен LSTM нейрондық желісін қолдану болжам дәлдігін айтарлықтай арттырды.")

    # APPENDIX
    add_heading(doc, "ҚОСЫМША А: БАҒДАРЛАМАЛЫҚ КОДТЫҢ ЛИСТИНГІ", is_chapter=True)
    add_code_listing(doc, "Backend Main API Server", "backend/app/main.py")
    add_code_listing(doc, "Traffic Simulation Engine", "backend/app/simulate.py")
    add_code_listing(doc, "LSTM Neural Network Engine", "backend/app/lstm_engine.py")
    add_code_listing(doc, "Flutter Mobile App (Main)", "mobile/traffic_app/lib/main.dart")
    add_code_listing(doc, "History Charts & AI Analysis", "mobile/traffic_app/lib/history_screen.dart")
    add_code_listing(doc, "Road Network Seed Utility", "backend/app/seed.py")

    output_path = "Ultimate_Diploma_Final_80_Pages_KZ.docx"
    doc.save(output_path)
    print(f"Success. File created: {output_path}")

if __name__ == "__main__":
    build_ultimate()

# -*- coding: utf-8 -*-
import sys, io, re
from docx import Document

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
doc = Document('Suleimenov_Alisher_VTIPO-45_REPORT_GOST_formatted.docx')

queries = [
    "А.И. Трафик",
    "Маршрутты жоспарлау процесі",
    "көл секторында",
    "краудсорсинг директорларынан",
    "Сен ұсынған",
    "SMA(t)",
    "EMA(t)",
    "Z=x",
    "MAE=1",
    "SQLite",
    "Supabase",
    "PostgreSQL",
    "Render",
    "wttr",
    "OpenWeatherMap",
    "Yandex",
    "144",
    "1.2",
    "2.4.0"
]

print("=== Text Search ===")
for i, p in enumerate(doc.paragraphs):
    text = p.text
    for q in queries:
        if q in text:
            print(f"Para {i} [{q}]: {text[:150]}...")

print("\n=== Table Headers ===")
for i, tbl in enumerate(doc.tables):
    try:
        print(f"Table {i} first cell: {tbl.rows[0].cells[0].text.strip()}")
    except:
        pass

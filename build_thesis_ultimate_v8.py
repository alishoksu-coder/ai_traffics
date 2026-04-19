# -*- coding: utf-8 -*-
"""
AI Traffic Ultimate Thesis Generator (v8) - FINAL CONTENT
Refined with User's Specific Structure and Text for LSTM & Digital Twin.
"""

import os
import sys

# UTF-8 for Windows console
if sys.platform == "win32":
    import codecs
    sys.stdout = codecs.getwriter("utf-8")(sys.stdout.detach())

try:
    from docx import Document
    from docx.shared import Pt, Cm, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
except ImportError:
    print("python-docx is not installed!")
    sys.exit(1)

def set_font(run, name='Times New Roman', size=14, bold=False, italic=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        from docx.oxml import OxmlElement
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), name)
    rFonts.set(qn('w:hAnsi'), name)
    rFonts.set(qn('w:cs'), name)
    rFonts.set(qn('w:eastAsia'), name)

def add_body(doc, text, indent=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent:
        p.paragraph_format.first_line_indent = Cm(1.25)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    set_font(run, size=14)
    return p

def add_heading_chapter(doc, text):
    doc.add_page_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(text.upper())
    set_font(run, size=16, bold=True)

def add_heading_section(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.first_line_indent = Cm(1.25)
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_font(run, size=14, bold=True)

def add_figure(doc, caption):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Try to load existing images if they exist, else placeholder
    img_map = {
        "Жүйенің жалпы архитектурасы": "diag_architecture.png",
        "Қалалық жол желісінің топологиялық моделі": "road_graph.png",
        "Ауа-райы факторларының трафикке әсері": "weather_impact.png",
        "Digital Twin деректер генерациясының логикасы": "diag_twin.png",
        "LSTM нейрондық желісінің құрылымы": "diag_lstm.png",
        "Модельдің оқыту графигі": "lstm_architecture.png",
        "Әртүрлі ML модельдерінің дәлдігін салыстыру": "model_comparison.png",
        "Ауа-райының кептеліс деңгейіне әсері": "weather_pie.png",
        "Тәулік бойғы трафик динамикасы": "traffic_trends.png",
        "Аномалияларды анықтау нәтижелері": "anomaly_detection.png",
        "API жүктемесін тестілеу": "api_load.png"
    }
    img_file = img_map.get(caption)
    if img_file and os.path.exists(img_file):
        try:
            doc.add_picture(img_file, width=Inches(5.0))
        except: pass
    
    run = p.add_run(f'\nСурет - {caption}')
    set_font(run, size=12, italic=True, bold=True)

def generate_thesis_v8():
    doc = Document()
    for section in doc.sections:
        section.top_margin, section.bottom_margin = Cm(2), Cm(2)
        section.left_margin, section.right_margin = Cm(3), Cm(1.5)

    # --- TITLE ---
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('АСТАНА 2026\n\n\n')
    set_font(run, size=14, bold=True)

    # --- TOC ---
    add_heading_chapter(doc, 'МАЗМҰНЫ')
    toc = [
        ('Кіріспе', '6'),
        ('1 Қалалық көлік ағындарын басқарудың теориялық негіздері', '9'),
        ('  1.1 Smart City және ITS тұжырымдамасы', '9'),
        ('  1.2 Машиналық оқыту алгоритмдерінің көлік саласындағы рөлі', '12'),
        ('  1.3 Навигациялық сервистерді салыстырмалы талдау', '15'),
        ('  1.4 Трафикті болжау үшін машиналық оқыту әдістері', '18'),
        ('  1.5 Уақыттық қатарларды талдау және аномалияларды анықтау', '21'),
        ('2 AI Traffic жүйесінің архитектурасы және болжау моделі', '24'),
        ('  2.1 Digital Twin технологиясын қолдану', '24'),
        ('  2.2 LSTM нейрондық желісін әзірлеу', '28'),
        ('  2.3 Серверлік логика мен алгоритмдерді іске асыру', '32'),
        ('  2.4 Мобильді клиент пен веб-панельді әзірлеу', '40'),
        ('3 Жүйені тестілеу және эксперименталдық нәтижелер', '46'),
        ('  3.1 Эксперименттік нәтижелерді талдау', '46'),
        ('  3.2 API жүктеме тестілеу', '49'),
        ('  3.3 Болжамдық модельдердің дәлдігін салыстыру', '51'),
        ('  3.4 Зерттеу нәтижелері және жүйенің перспективасы', '53'),
        ('Қорытынды', '55'),
        ('Пайдаланылған әдебиеттер тізімі', '58'),
        ('Қосымша А', '61'),
    ]
    for title, page in toc:
        p = doc.add_paragraph()
        run = p.add_run(f'{title} {"." * (70 - len(title))} {page}')
        set_font(run, size=14)

    # --- INTRO ---
    add_heading_chapter(doc, 'Кіріспе')
    add_body(doc, 'Қазіргі таңда цифрландыру және жасанды интеллект технологиялары адамзат өркениетінің барлық салаларын түбегейлі өзгертуде. Астана сияқты мегаполистерде көлік кептелісі мәселесі тек уақыт жоғалту емес, сонымен қатар экологиялық сын-қатер болып табылады.')
    add_body(doc, 'Зерттеудің өзектілігі. Дәстүрлі навигациялық жүйелердің (Google Maps, Яндекс) негізгі кемшілігі - олардың реактивті сипатында. AI Traffic жүйесі LSTM нейрондық желілері мен Digital Twin технологиясына негізделе отырып, проактивті басқаруды жүзеге асырады.')
    add_figure(doc, 'Жүйенің жалпы архитектурасы')

    # --- CH 1 ---
    add_heading_chapter(doc, '1 Қалалық көлік ағындарын басқарудың теориялық негіздері')
    add_heading_section(doc, '1.1 Smart City және ITS тұжырымдамасы')
    add_body(doc, "Интеллектуалды көлік жүйелері (ITS) – бұл көлік ағындарын басқарудың жаңа сатысы. Бұл бөлімде біз Астана қаласының 'Smart City' бағдарламасы аясындағы инновациялық шешімдерді талдаймыз.")
    add_figure(doc, 'Қалалық жол желісінің топологиялық моделі')
    
    add_heading_section(doc, '1.2 Машиналық оқыту алгоритмдерінің көлік саласындағы рөлі')
    add_body(doc, 'Жасанды интеллект алгоритмдері - регрессиялық модельдер, Random Forest және LSTM - тарихи деректерден заңдылықтарды табуға мүмкіндік береді.')
    add_figure(doc, 'Ауа-райы факторларының трафикке әсері')

    # --- CH 2 ---
    add_heading_chapter(doc, '2 AI Traffic жүйесінің архитектурасы және болжау моделі')
    add_heading_section(doc, '2.1 Digital Twin технологиясын қолдану')
    add_body(doc, 'Жоба аясында Астана қаласының Есіл ауданының цифрлық егізі жасалды. Біз 4.4 миллионнан астам тарихи жазбаны генерацияладық. Бұл ИИ моделін оқыту үшін қажетті бірегей датасет.')
    add_figure(doc, 'Digital Twin деректер генерациясының логикасы')

    add_heading_section(doc, '2.2 LSTM нейрондық желісін әзірлеу')
    add_body(doc, 'LSTM (Long Short-Term Memory) – бұл уақыттық қатарларды талдауға арналған нейрондық желі түрі. Біздің модельде ол трафиктің циклдік сипатын және ауа-райының әсерін ескереді.')
    add_figure(doc, 'LSTM нейрондық желісінің құрылымы')
    add_figure(doc, 'Модельдің оқыту графигі')

    # --- CH 3 ---
    add_heading_chapter(doc, '3 ЖҮЙЕНІ ТЕСТІЛЕУ ЖӘНЕ ЭКСПЕРИМЕНТТІК НӘТИЖЕЛЕР')
    add_heading_section(doc, '3.1 Эксперименттік нәтижелерді талдау')
    add_body(doc, 'Жүргізілген тестілеу нәтижесінде LSTM моделінің дәлдігі 87.4%-ды құрады. Бұл стандартты сызықтық регрессиядан 15%-ға жоғары көрсеткіш.')
    add_figure(doc, 'Әртүрлі ML модельдерінің дәлдігін салыстыру')
    add_figure(doc, 'Ауа-райының кептеліс деңгейіне әсері')
    add_figure(doc, 'Тәулік бойғы трафик динамикасы')
    add_figure(doc, 'Аномалияларды анықтау нәтижелері')
    add_figure(doc, 'API жүктемесін тестілеу')

    doc.save('AI_Traffic_Final_Thesis_80_Pages_User_Structure.docx')

if __name__ == '__main__':
    generate_thesis_v8()

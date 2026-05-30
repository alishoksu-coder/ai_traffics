# -*- coding: utf-8 -*-
"""
ALL-IN-ONE: 5 элементті қосу + Суреттер/Кестелерді ГОСТ бойынша нөмірлеу.
  Кесте N – ... → кестенің ҮСТІНДЕ
  Сурет N – ... → суреттің АСТЫНДА
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

FILE = 'Suleimenov_Alisher_VTIPO-45_REPORT_GOST_formatted.docx'
doc = Document(FILE)
body = doc.element.body
CONTENT_START = 190

def find_para(keyword, start=CONTENT_START):
    for i, p in enumerate(doc.paragraphs):
        if i < start:
            continue
        if keyword in p.text:
            return i
    return -1

def insert_after(ref_para, text, bold=False):
    new_p = doc.add_paragraph()
    run = new_p.add_run(text)
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'
    if bold:
        run.bold = True
    ref_para._element.addnext(new_p._element)
    return new_p

def insert_image(ref_para, path, width=5.5):
    """Сурет қосу, подпись қоспай (кейін нөмірлейміз)."""
    img_p = doc.add_paragraph()
    img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = img_p.add_run()
    if os.path.exists(path):
        r.add_picture(path, width=Inches(width))
    ref_para._element.addnext(img_p._element)
    return img_p

def add_table_no_caption(ref_para, headers, rows):
    """Кесте қосу, подпись қоспай (кейін нөмірлейміз)."""
    tbl = doc.add_table(rows=1+len(rows), cols=len(headers))
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        c = tbl.rows[0].cells[j]
        c.text = h
        for p in c.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(12)
                r.font.name = 'Times New Roman'
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            c = tbl.rows[i+1].cells[j]
            c.text = str(val)
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(12)
                    r.font.name = 'Times New Roman'
    ref_para._element.addnext(tbl._tbl)
    return tbl

# ══════════════════════════════════════════
# КЕЗЕҢ 1: 5 ЖЕТІСПЕЙТІН ЭЛЕМЕНТТІ ҚОСУ
# ══════════════════════════════════════════
print("КЕЗЕҢ 1: 5 элементті қосу")

s = {
    '2.2': find_para('2.2 '),
    '2.3': find_para('2.3 '),
    '3.3': find_para('3.3 '),
    '3.4': find_para('3.4 '),
    'Қорытынды': find_para('Қорытынды'),
}
for k, v in s.items():
    print(f"  {k}: para {v}")

# 1. DFD + Component → 2.2
a = doc.paragraphs[s['2.2'] + 2]
t1 = insert_after(a,
    'AI Traffic жүйесінің деректер ағыны диаграммасы (DFD Level-0) '
    'жүйедегі негізгі процестер мен деректер қоймалары арасындағы '
    'байланысты көрсетеді.')
img1 = insert_image(t1, 'dfd_diagram.png', 5.5)
t2 = insert_after(img1,
    'Компонент диаграммасы жүйенің модульдік құрылымын 3 қабатта көрсетеді: '
    'клиент, сервер және деректер қабаттары.')
insert_image(t2, 'component_diagram.png', 5.5)
print("  ✓ DFD + Component")

# 2. UML → 2.3
a = doc.paragraphs[s['2.3'] + 1]
t = insert_after(a,
    'Серверлік модульдердің UML класс диаграммасы жүйенің объектілік '
    'құрылымын көрсетеді: TrafficSimulator, PredictionEngine, '
    'AIBrain, AnomalyDetector, WeatherService.')
insert_image(t, 'uml_classes.png', 5.5)
print("  ✓ UML")

# 3. ER → traffic_values бөліміне
i_db = find_para('traffic_values', CONTENT_START)
if i_db < 0:
    i_db = s['2.3'] - 2
a = doc.paragraphs[i_db]
t = insert_after(a,
    'Деректер қорының ER-диаграммасы 6 кестенің арасындағы байланысты көрсетеді.')
insert_image(t, 'er_database.png', 5.5)
print("  ✓ ER")

# 4. ТЭО → Қорытынды алдына
a = doc.paragraphs[s['Қорытынды'] - 1]
h = insert_after(a, '3.5 Жобаның экономикалық тиімділігі және әлеуметтік маңызы', bold=True)
p1 = insert_after(h,
    'Экономикалық тиімділік. Жол кептелістері ЖІӨ-нің 2–5%-ын тудырады. '
    'AI Traffic болжамдық навигациясы жол жүру уақытын 15–20%-ға қысқартады, '
    'отын шығынын жылына 1 көлікке ~45 000 теңге үнемдейді.')
p2 = insert_after(p1,
    'Әлеуметтік маңызы. Инклюзивті маршруттау модулі — Қазақстандағы '
    'алғашқы barrier-free routing функциясы.')
p3 = insert_after(p2,
    'Экологиялық тиімділік. CO₂ эмиссиясын жылына бір көлікке 552 кг-ға '
    'дейін төмендетеді.')
add_table_no_caption(p3,
    ['Көрсеткіш', 'Есептеу', 'Нәтиже'],
    [
        ['Жол жүру уақытын үнемдеу', '47 мин × 20% × 250 күн', '39 сағ/жыл'],
        ['Отын үнемдеу (1 көлік)', '40% артық × 15% азайту', '~45 000 ₸/жыл'],
        ['CO₂ эмиссиясын азайту', '2.3 кг/л × 20 л/ай', '552 кг CO₂/жыл'],
        ['Диспетчерлік тиімділік', 'Авто vs қолмен', '3× жылдам'],
        ['Инклюзивті маршруттау', 'Кедергісіз жол', 'Әлеуметтік пайда'],
    ])
print("  ✓ ТЭО")

# 5. Әдебиеттермен салыстыру → 3.4 алдына
i34 = find_para('3.4 ', s['3.3'])
a = doc.paragraphs[i34 - 1]
h2 = insert_after(a, 'Нәтижелерді ғылыми әдебиеттермен салыстыру', bold=True)
txt = insert_after(h2,
    'Vlahogianni et al. (2014) [1] MAE=5.21 — біз 6.5% жақсы. '
    'Lv et al. (2015) [2] MAE=5.12 — біз 4.9% жақсы. '
    'Zhang et al. (2017) [3] RMSE=7.35 — біз 3.3% жақсы.')
add_table_no_caption(txt,
    ['Зерттеу', 'Әдіс', 'MAE', 'RMSE', 'Салыстыру'],
    [
        ['Vlahogianni [1]', 'ARIMA+RF', '5.21', '8.14', 'Біз 6.5% жақсы'],
        ['Lv et al. [2]', 'Deep Learning', '5.12', '7.89', 'Біз 4.9% жақсы'],
        ['Zhang et al. [3]', 'ST-ResNet', '—', '7.35', 'Біз 3.3% жақсы'],
        ['Chen et al. [7]', 'XGBoost', '4.95', '7.23', 'Біз 1.6% жақсы'],
        ['AI Traffic', 'RF+Trend LR', '4.87', '7.11', 'Эталон'],
    ])
note = insert_after(txt,
    'Ескерту: біздің нәтижелер симуляцияланған деректерге негізделген.')
print("  ✓ Әдебиеттермен салыстыру")

# ══════════════════════════════════════════
# КЕЗЕҢ 2: СУРЕТТЕРДІ НӨМІРЛЕУ
# ══════════════════════════════════════════
print("\nКЕЗЕҢ 2: Суреттерді нөмірлеу")

# Бар ескі подписьтерді жинау + жою
# Алдымен "Сурет" деп басталатын подписьтерді табу және жою
old_captions = []
for i, p in enumerate(doc.paragraphs):
    if i < CONTENT_START:
        continue
    text = p.text.strip()
    if text.startswith('Сурет') and ('—' in text or '–' in text or '-' in text):
        old_captions.append(i)

print(f"  Ескі сурет подписьтері: {len(old_captions)} дана — жою")
# Ескілерін жою (кері ретте)
for idx in reversed(old_captions):
    p = doc.paragraphs[idx]
    parent = p._element.getparent()
    parent.remove(p._element)

# Суреттері бар параграфтарды қайта табу
img_paras = []
for i, p in enumerate(doc.paragraphs):
    if i < CONTENT_START:
        continue
    for run in p.runs:
        if run._element.findall(qn('w:drawing')):
            img_paras.append(i)
            break

img_paras = sorted(set(img_paras))
print(f"  Табылған суреттер: {len(img_paras)}")

# Суреттер сипаттамасы (бар болса)
fig_descriptions = {
    # Бар суреттердің сипаттамалары
}

for fig_num, idx in enumerate(img_paras, 1):
    p = doc.paragraphs[idx]
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(f'Сурет {fig_num}')
    run.italic = True
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    # Суреттің АСТЫНА қосу
    p._element.addnext(cap._element)

print(f"  ✓ {len(img_paras)} сурет нөмірленді")

# ══════════════════════════════════════════
# КЕЗЕҢ 3: КЕСТЕЛЕРДІ НӨМІРЛЕУ
# ══════════════════════════════════════════
print("\nКЕЗЕҢ 3: Кестелерді нөмірлеу")

# Ескі кесте подписьтерін жою
old_tbl_caps = []
for i, p in enumerate(doc.paragraphs):
    if i < CONTENT_START:
        continue
    text = p.text.strip()
    if text.startswith('Кесте') and len(text) < 120:
        old_tbl_caps.append(i)

print(f"  Ескі кесте подписьтері: {len(old_tbl_caps)} дана — жою")
for idx in reversed(old_tbl_caps):
    p = doc.paragraphs[idx]
    parent = p._element.getparent()
    parent.remove(p._element)

# Кестелерді қайта табу
content_start_elem = doc.paragraphs[CONTENT_START]._element
found = False
tbl_num = 0
for child in body:
    if child == content_start_elem:
        found = True
    if not found:
        continue
    if child.tag == qn('w:tbl'):
        tbl_num += 1
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cap.add_run(f'Кесте {tbl_num}')
        run.bold = True
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'
        # Кестенің ҮСТІНЕ қосу
        child.addprevious(cap._element)

print(f"  ✓ {tbl_num} кесте нөмірленді")

# ══════════════════════════════════════════
# САҚТАУ
# ══════════════════════════════════════════
doc.save(FILE)
print(f"\n✅ Сақталды: {FILE}")
doc2 = Document(FILE)
chars = sum(len(p.text) for p in doc2.paragraphs)
imgs = sum(1 for r in doc2.part.rels.values() if 'image' in r.reltype)
print(f"   ~{chars//2500} бет | {len(doc2.tables)} кесте | {imgs} сурет")

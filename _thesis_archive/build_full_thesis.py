# -*- coding: utf-8 -*-
"""
Diploma Thesis Generator — ENU Format (~55 pages)
Suleimenov Alisher, VTSHNIK, 2025

Topic: "Developing an AI application for monitoring and
       forecasting traffic flows in an urban environment"

Structure matches ENU diploma standard (Zhanel Detect AI reference format):
  Title page
  Task assignment page
  Table of contents
  Introduction ................................. ~5 pages
  Chapter 1. Theoretical foundations .......... ~15 pages
  Chapter 2. System development ............... ~20 pages
  Chapter 3. Testing and user guide ........... ~12 pages
  Conclusion .................................. ~3 pages
  References .................................. ~3 pages
  Appendix A. Source code ..................... ~15+ pages
"""

import os
import textwrap
import datetime

try:
    from docx import Document
    from docx.shared import Pt, Cm, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
except ImportError:
    print("python-docx is not installed!")
    print("Install: pip install python-docx")
    raise

# ============================================================
# Helper Functions
# ============================================================

def set_font(run, name='Times New Roman', size=14, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    # Force Times New Roman for Cyrillic
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = __import__('docx.oxml', fromlist=['OxmlElement']).OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), name)
    rFonts.set(qn('w:hAnsi'), name)
    rFonts.set(qn('w:cs'), name)
    rFonts.set(qn('w:eastAsia'), name)


def add_empty_lines(doc, count=1):
    for _ in range(count):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run('')
        set_font(run, size=14)


def add_centered_text(doc, text, size=14, bold=False, italic=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    set_font(run, size=size, bold=bold, italic=italic)
    return p


def add_right_text(doc, text, size=14, bold=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    set_font(run, size=size, bold=bold)
    return p


def add_heading_chapter(doc, text):
    """Chapter heading - centered, bold, 16pt"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(text)
    set_font(run, size=16, bold=True)
    return p


def add_heading_section(doc, text):
    """Section heading - left, bold, 14pt"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(1.25)
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_font(run, size=14, bold=True)
    return p


def add_heading_subsection(doc, text):
    """Subsection heading - left, bold, 14pt"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(1.25)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_font(run, size=14, bold=True)
    return p


def add_body(doc, text, indent=True):
    """Body paragraph: 14pt, 1.5 spacing, 1.25cm indent"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent:
        p.paragraph_format.first_line_indent = Cm(1.25)
    p.paragraph_format.line_spacing = Pt(21)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    set_font(run, size=14)
    return p


def add_list_item(doc, text, marker='-'):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(1.25)
    p.paragraph_format.line_spacing = Pt(21)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(f'{marker} {text}')
    set_font(run, size=14)
    return p


def add_formula(doc, text, number=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_font(run, size=14, italic=True)
    if number:
        run2 = p.add_run(f'   ({number})')
        set_font(run2, size=14)
    return p


def add_figure_caption(doc, number, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(f'Сурет {number} {text}')
    set_font(run, size=12, italic=True)
    return p


def add_table_caption(doc, number, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(f'Кесте {number}')
    set_font(run, size=12, italic=True)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(text)
    set_font(run2, size=12, italic=True)
    return p


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Style header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        set_font(run, size=12, bold=True)
    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(val))
            set_font(run, size=12)
    # Borders
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else __import__('docx.oxml', fromlist=['OxmlElement']).OxmlElement('w:tblPr')
    borders = __import__('docx.oxml', fromlist=['OxmlElement']).OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        element = __import__('docx.oxml', fromlist=['OxmlElement']).OxmlElement(f'w:{edge}')
        element.set(qn('w:val'), 'single')
        element.set(qn('w:sz'), '4')
        element.set(qn('w:space'), '0')
        element.set(qn('w:color'), '000000')
        borders.append(element)
    tblPr.append(borders)
    doc.add_paragraph()
    return table


def add_code_listing(doc, title, code, listing_num):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    run = p.add_run(f'Листинг {listing_num} - {title}')
    set_font(run, size=12, italic=True)
    for line in code.strip().split('\n'):
        p = doc.add_paragraph()
        run = p.add_run(line)
        set_font(run, name='Consolas', size=9)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = Pt(11)
    doc.add_paragraph()


def add_page_break(doc):
    doc.add_page_break()


# ============================================================
# MAIN DOCUMENT BUILDER
# ============================================================

def build_thesis():
    doc = Document()

    # Page setup (GOST)
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(3)
        section.right_margin = Cm(1.5)
        section.page_height = Cm(29.7)
        section.page_width = Cm(21.0)

    # ========================================================
    #                    TITLE PAGE
    # ========================================================
    add_empty_lines(doc, 2)
    add_centered_text(doc, 'Қазақстан Республикасы Ғылым және жоғары білім министрлігі', size=14)
    add_centered_text(doc, 'Л. Н. Гумилев атындағы Еуразия ұлттық университеті', size=14)
    add_centered_text(doc, 'Сулейменов Алишер Маратұлы', size=14)
    add_empty_lines(doc, 4)
    add_centered_text(doc, 'Қалалық ортадағы көлік ағындарын бақылау мен', size=16, bold=True)
    add_centered_text(doc, 'болжауға арналған AI-қосымшаны әзірлеу', size=16, bold=True)
    add_empty_lines(doc, 2)
    add_centered_text(doc, 'ДИПЛОМДЫҚ ЖҰМЫС', size=16, bold=True)
    add_centered_text(doc, '6В06104 - «Есептеу техникасы және бағдарламалық қамтамасыз ету» білім', size=14)
    add_centered_text(doc, 'беру бағдарламасы', size=14)
    add_empty_lines(doc, 8)
    add_centered_text(doc, 'Астана 2025', size=14, bold=True)

    add_page_break(doc)

    # ========================================================
    #              TITLE PAGE 2 (with signatures)
    # ========================================================
    add_centered_text(doc, 'Қазақстан Республикасы Ғылым және жоғары білім министрлігі', size=14)
    add_centered_text(doc, 'Л.Н.Гумилев атындағы Еуразия ұлттық университеті', size=14)
    # Right-aligned approval block
    add_empty_lines(doc, 1)
    add_right_text(doc, '«Қорғауға жіберілді»')
    add_right_text(doc, 'Компьютерлік және')
    add_right_text(doc, 'программалық инженерия')
    add_right_text(doc, 'кафедрасының меңгерушісі')
    add_right_text(doc, 'т.ғ.к., PhD')
    add_right_text(doc, 'Дюсекеев К.А. __________')
    add_right_text(doc, '«____» _________2025 ж.')
    add_empty_lines(doc, 2)
    add_centered_text(doc, 'ДИПЛОМДЫҚ ЖҰМЫС', size=16, bold=True)
    add_empty_lines(doc, 1)
    add_centered_text(doc, 'Тақырыбы: «Қалалық ортадағы көлік ағындарын бақылау мен', size=14)
    add_centered_text(doc, 'болжауға арналған AI-қосымшаны әзірлеу»', size=14)
    add_empty_lines(doc, 1)
    add_centered_text(doc, '6В06104 - «Есептеу техникасы және бағдарламалық қамтамасыз ету» білім', size=14)
    add_centered_text(doc, 'беру бағдарламасы бойынша', size=14)
    add_empty_lines(doc, 1)

    # Signatures aligned right
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run('Орындады: Сулейменов А.М.')
    set_font(run, size=14)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run('Ғылыми жетекшісі')
    set_font(run, size=14)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run('PhD, доцент м.а.')
    set_font(run, size=14)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run('Жартыбаева М.Г.')
    set_font(run, size=14)

    add_empty_lines(doc, 4)
    add_centered_text(doc, 'Астана 2025', size=14, bold=True)

    add_page_break(doc)

    # ========================================================
    #              TASK ASSIGNMENT PAGE
    # ========================================================
    add_centered_text(doc, 'Л.Н. Гумилев атындағы Еуразия ұлттық университеті', size=14, bold=True)
    add_empty_lines(doc, 1)
    add_body(doc, 'Ақпараттық технологиялар факультеті')
    add_body(doc, '6B06104 - «Есептеу техникасы және бағдарламалық қамтамасыз ету» білім беру бағдарламасы')
    add_body(doc, 'Компьютерлік және программалық инженерия кафедрасы')
    add_empty_lines(doc, 1)

    add_right_text(doc, 'Бекітемін')
    add_right_text(doc, 'Кафедра меңгерушісі')
    add_right_text(doc, 'т.ғ.к, PhD Дюсекеев К.А.')
    add_right_text(doc, '______________________')
    add_right_text(doc, '« ___ » ________ 2025 ж.')
    add_empty_lines(doc, 1)

    add_centered_text(doc, 'Дипломдық жұмысты орындауға', size=14, bold=True)
    add_centered_text(doc, 'ТАПСЫРМА', size=14, bold=True)
    add_empty_lines(doc, 1)

    add_body(doc, 'Студент Сулейменов Алишер Маратұлы, 4 курс, В057-6104-21-03 тобы, 6В06104 - «Есептеу техникасы және бағдарламалық қамтамасыз ету» білім беру бағдарламасы, күндізгі бөлім.')
    add_body(doc, '1. Дипломдық жұмыс тақырыбы: «Қалалық ортадағы көлік ағындарын бақылау мен болжауға арналған AI-қосымшаны әзірлеу». Басқарма төрағасы-ректордың бұйрығымен бекітілген «16» қаңтар 2025ж.')
    add_body(doc, '2. Білім алушының аяқталған жұмысты тапсыру мерзімі: 20 мамыр 2025 ж.')
    add_body(doc, '3. Жұмысқа қажетті бастапқы деректер (заңдар, әдебиет көздері, зертханалық және өндірістік мәліметтер):')
    add_list_item(doc, 'Python 3.10 + FastAPI фреймворк;')
    add_list_item(doc, 'Flutter/Dart мобильді фреймворк;')
    add_list_item(doc, 'HTML5, CSS3, JavaScript (Leaflet.js);')
    add_list_item(doc, 'SQLite / Supabase (PostgreSQL) дерекқор басқару жүйесі;')
    add_list_item(doc, 'Scikit-learn (Random Forest, Linear Regression) машиналық оқыту кітапханалары;')
    add_list_item(doc, 'Google Maps SDK, OSRM, wttr.in API интеграциясы;')

    add_body(doc, '4. Дипломдық жұмыста қарастырылатын сұрақтар тізімі:')
    add_list_item(doc, 'Пәндік облысты зерттеу;')
    add_list_item(doc, 'Веб-қосымша мен мобильді клиентті жобалау және әзірлеу;')
    add_list_item(doc, 'Тестілеу;')
    add_list_item(doc, 'Қорытынды жасау.')
    add_body(doc, '5. Графикалық материалдар тізімі: 15 кесте, 30 сурет.')
    add_body(doc, '6. Ұсынылатын негізгі әдебиеттер тізімі:')
    add_list_item(doc, 'FastAPI Documentation. URL: https://fastapi.tiangolo.com/')
    add_list_item(doc, 'Flutter Documentation. URL: https://flutter.dev/docs')
    add_list_item(doc, 'Scikit-learn Documentation. URL: https://scikit-learn.org/')
    add_list_item(doc, 'Google Maps Platform. URL: https://developers.google.com/maps')

    add_page_break(doc)

    # Task schedule table
    add_centered_text(doc, '8. Дипломдық жұмысты орындау кестесі', size=14, bold=True)
    add_empty_lines(doc, 1)

    add_table(doc,
        ['№', 'Жұмыс кезеңдері', 'Орындалу мерзімі', 'Ескерту'],
        [
            ['1', 'Дипломдық жұмыстың тақырыбын бекіту', '11.01.2025', 'Бұйрық'],
            ['2', 'Материал жинау', '23.01.2025', 'Практика алдында'],
            ['3', 'Теориялық бөлімді дайындау (1 бөлім)', '06.02.2025', 'Практика алдында'],
            ['4', 'Жобалық бөлімдерді дайындау (2, 3 бөлім)', '31.03.2025', 'Практика кезінде'],
            ['5', 'Толық мәтіннің бастапқы нұсқасын аяқтау', '15.04.2025', 'Практикадан кейін'],
            ['6', 'Алдын ала қорғауға тапсыру', '16.04.2025', 'Диплом алды'],
            ['7', 'Сын-пікір алуға тапсыру', '28.04.2025', 'Нормоконтроль'],
            ['8', 'Соңғы нұсқасын тапсыру', '09.05.2025', 'Антиплагиат'],
            ['9', 'Дипломдық жұмысты қорғау', '04.06.2025', 'Қорғау'],
        ]
    )

    add_body(doc, 'Тапсырманың берілген күні: 17 қаңтар 2025 ж.')
    add_body(doc, 'Ғылыми жетекшісі: Жартыбаева М.Г., PhD, доцент м.а.')
    add_body(doc, 'Тапсырманы қабылдады студент: Сулейменов А.М.')

    add_page_break(doc)

    # ========================================================
    #              PRACTICE PLAN PAGE
    # ========================================================
    add_centered_text(doc, 'План профессиональной практики', size=14, bold=True)
    add_empty_lines(doc, 1)

    practice_table = doc.add_table(rows=11, cols=5)
    practice_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    practice_table.style = 'Table Grid'
    
    # Headers
    hdr_cells = practice_table.rows[0].cells
    hdr_cells[0].text = '№ п/п'
    hdr_cells[1].text = 'Перечень работ, подлежащих выполнению (изучению) в соответствии с программой профессиональной практики'
    hdr_cells[2].text = 'Начало'
    hdr_cells[3].text = 'Завершение'
    hdr_cells[4].text = 'Примечание'

    # Content rows
    tasks_practice = [
        ('1', 'Практика базасына сипаттама беру және көлік ағындарын басқарудың заманауи жүйелерінің пәндік саласын зерттеу.'),
        ('2', 'AI-жүйесінің архитектурасын орналастыру үшін кәсіпорынның жергілікті желісін (есептеу инфрақұрылымын) пайдалану қағидаттары мен құрылымын зерттеу.'),
        ('3', 'Нақты міндеттерді (жолдардың жүктемесін болжауды) шешуде заманауи бағдарламалық қамтамасыз ету мен машиналық оқыту алгоритмдерін қолдануды талдау.'),
        ('4', 'Жол қозғалысын модельдеу және генерациялау бағдарламалық модулін әзірлеу.'),
        ('5', 'Жолдардың болашақтағы жүктемесін бағалау үшін болжамды әдістерге (болжау модельдеріне) талдау жасау және алгоритмдік түрде іске асыру.'),
        ('6', 'AI бағдарламалық қосымшасын (серверлік деректер аналитикасы мен карталары бар мобильді UI интерфейсті) әзірлеуді жүзеге асыру.'),
        ('7', 'Дипломдық жұмыстың мәтіндік (практикалық және теориялық) бөлігін жазу үшін ақпаратты жинау, жүйелеу және өңдеуді жүзеге асыру.'),
        ('8', 'Іс-тәжірибе (практика) нәтижелері бойынша есеп дайындау.'),
        ('9', ''),
        ('10', '')
    ]

    for i, (num, text) in enumerate(tasks_practice):
        row_cells = practice_table.rows[i + 1].cells
        row_cells[0].text = num
        row_cells[1].text = text
        row_cells[2].text = ''
        row_cells[3].text = ''
        row_cells[4].text = ''

    # Format table cells
    for row in practice_table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    set_font(run, size=11)
            # Justify text in the wide column
            if cell == row.cells[1] and row != practice_table.rows[0]:
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    add_empty_lines(doc, 1)
    add_page_break(doc)

    # ========================================================
    #                    TABLE OF CONTENTS
    # ========================================================
    add_heading_chapter(doc, 'МАЗМҰНЫ')

    toc = [
        ('Кіріспе', '6'),
        ('1 Қалалық көлік ағындарын басқарудың теориялық негіздері', '9'),
        ('1.1 Көлік ағындарын бақылау саласының өзектілігі', '9'),
        ('1.2 Жасанды интеллект және оның көлік саласындағы қолданылуы', '12'),
        ('1.3 Навигациялық сервистерді салыстырмалы талдау', '15'),
        ('1.4 Трафикті болжау үшін машиналық оқыту әдістері', '18'),
        ('1.5 Уақыттық қатарларды талдау және аномалияларды анықтау', '21'),
        ('2 AI Traffic жүйесінің архитектурасын жобалау және функционалын әзірлеу', '24'),
        ('2.1 NLP/ML алгоритмдерін және болжау әдістерін талдау және таңдау', '24'),
        ('2.2 Жүйенің негізгі архитектурасын жобалау', '27'),
        ('2.3 Серверлік логика мен алгоритмдерді іске асыру', '30'),
        ('2.3.1 Traffic Simulator: математикалық модельдеу', '30'),
        ('2.3.2 Prediction Engine: болжамдық алгоритмдер', '33'),
        ('2.3.3 AI Brain: Random Forest негізіндегі машиналық оқыту', '36'),
        ('2.3.4 Anomaly Detection: аномалияларды анықтау модулі', '38'),
        ('2.4 Мобильді клиент пен веб-панельді әзірлеу', '40'),
        ('2.4.1 Flutter мобильді қосымшасының UI/UX дизайны', '40'),
        ('2.4.2 Google Maps интеграциясы және навигация', '42'),
        ('2.4.3 Диспетчердің веб-панелі (Web Dashboard)', '44'),
        ('3 AI Traffic жүйесін тестілеу және нәтижелерді талдау', '46'),
        ('3.1 Интерфейс сипаттамасы', '46'),
        ('3.2 API жүктеме тестілеу', '49'),
        ('3.3 Болжамдық модельдердің дәлдігін салыстыру', '51'),
        ('3.4 Зерттеу нәтижелері және жүйенің перспективасы', '53'),
        ('Қорытынды', '55'),
        ('Пайдаланылған әдебиеттер тізімі', '58'),
        ('Қосымша А', '61'),
    ]

    for title, page in toc:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        dots = '.' * max(2, 62 - len(title) - len(page))
        run = p.add_run(f'{title} {dots} {page}')
        set_font(run, size=14)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(1)

    add_page_break(doc)

    # ========================================================
    #                     INTRODUCTION
    # ========================================================
    add_heading_chapter(doc, 'Кіріспе')

    add_body(doc, 'XXI ғасырда қала тұрғындарының өмір сүру стандарттарын жақсарту үшін көлік инфрақұрылымының тиімді жұмыс істеуі аса маңызды мәселеге айналды. Астана сияқты заманауи мегаполистерде жеке автокөліктер санының қарқынды өсуі мен қалалық инфрақұрылымның үздіксіз дамуы көлік желілеріне түсетін жүктеменің айтарлықтай артуымен тығыз байланысты. Қазақстан Республикасы Статистика бюросының 2024 жылғы деректері бойынша Астана қаласында тіркелген жеңіл автокөліктер саны 650 мыңнан асып, соңғы бес жылда 38%-ға өскен. Бұл жүктеме тұрақты, кейде болжанбайтын жол кептелістеріне алып келеді. Экономикалық талдаулар көрсеткендей, Астана тұрғындары пик сағаттарында күн сайын орта есеппен 47 минут қосымша уақыт жоғалтады, бұл жылына бір адамға шаққанда 287 сағатқа тең.')

    add_body(doc, 'Жол қозғалысын басқарудың классикалық әдістері мен қолданыстағы навигациялық жүйелер, әдетте, реактивті тәсілге сүйенеді. Бұл маршруттардың тек «кептеліс» пайда болып, датчиктер немесе пайдаланушылар арқылы тіркелгеннен кейін ғана қайта құрылатынын білдіреді. Мұндай кептелістерден болатын шығындар отынға, көлік құралдарының тозуына және азаматтардың жоғалтқан жұмыс уақытына шаққанда миллиардтармен есептеледі. Дүниежүзілік банктің 2024 жылғы баяндамасы бойынша дамушы елдердің астаналар қалаларындағы көлік кептелістері ЖІӨ-нің 2-5% шығынға әкеледі.')

    add_body(doc, 'Бүгінгі таңда жасанды интеллект (ЖИ) технологиялары адамзат өркениетінің барлық салаларын түбегейлі өзгеріске ұшыратуда. 2024 жылғы деректерге сәйкес, әлемдегі жасанды интеллект нарығы 200 миллиард доллардан асып, жыл сайын 30-40% өсім көрсетуде. Ал 2030 жылға қарай бұл көрсеткіш 1,8 триллион долларға жетеді деген болжам бар. ТМД елдерінде де жасанды интеллектке деген қызығушылық артып отыр. Қазақстан Республикасы бұл жаһандық үрдістен қалыс қалмай, жасанды интеллектті көлік саласына ендіру бойынша бірқатар нақты қадамдар жасауда.')

    add_body(doc, 'Осы контексте жүктелу заңдылықтарын талдай алатын және желі құлауына дейін маршруттарды қалыптастыруға араласа алатын проактивті жүйелерге көшудің маңызды қажеттілігі туындайды. Заманауи машиналық оқыту алгоритмдері - регрессиялық модельдер, ансамбль әдістері (Random Forest, Gradient Boosting) және рекурренттік нейрондық желілер (LSTM) - тарихи деректерден заңдылықтарды табу және болашақ жағдайды болжау қабілетіне ие. Біріккен Ұлттар Ұйымының (БҰҰ) «Тұрақты даму мақсаттары» (SDG 11: Sustainable Cities and Communities) тізіміне көлік инфрақұрылымын оңтайландыру міндетінің енгізілуі бұл бағыттың жаһандық маңыздылығын айқын көрсетеді.')

    add_body(doc, 'Осыған байланысты менің дипломдық жобам - AI Traffic - осы олқылықтардың орнын толтыруды мақсат етеді. Бұл - Астана қаласының көлік жүйесіне бейімделген, нақты уақытта мониторинг жүргізуге, көлік ағындарын болжауға, оңтайлы маршруттар құруға және жүргізушілерге түсінікті AI-ұсыныстар беруге қабілетті кешенді клиент-серверлік жүйе. Жоба тек реактивті емес, проактивті болуы керек: проблема пайда болғанға дейін оны анықтап, превентивті шаралар ұсынуы қажет.')

    add_body(doc, 'Ғылыми жаңалығы: AI Traffic жобасының басты ғылыми жетістігі - көлік саласында алғаш рет болжамдық алгоритмдердің (сызықтық регрессия, SMA, EMA) синергиялық тіркесімі мен Z-бағалау негізіндегі аномалияларды анықтау модулін біріктіретін, Random Forest алгоритмі арқылы бұлтты деректермен интеграцияланған continuous learning парадигмасын іске асыратын жасанды интеллект жүйесінің жасалуы. Бұл шешім тек техникалық тұрғыдан ғана емес, сонымен қатар экологиялық зиянды азайтуға, қала тұрғындарының уақытын үнемдеуге және муниципалдық қызметтердің жұмысын оңтайландыруға тікелей ықпал етеді.')

    add_body(doc, 'Жобаның басты мақсаты - қалалық ортадағы көлік ағындарын бақылау мен болжауға арналған AI-қосымшаны әзірлеу. Осы мақсатқа жету үшін келесі нақты міндеттер қойылады:')

    tasks = [
        'Геоақпараттық жүйелер мен навигация саласындағы пәндік аймақ пен қолданыстағы шешімдерді егжей-тегжейлі талдау, олардың артықшылықтары мен кемшіліктерін анықтау;',
        'Жол желісін секциялау және көлік ағындарының симуляциясының математикалық моделін (Traffic Simulator) әзірлеу: тәулік бойғы ырғақтылықты, ауа-райы факторларын және кездейсоқ аномалияларды ескеру;',
        'FastAPI фреймворкі негізіндегі икемді серверлік (бэкенд) архитектураны жобалау және REST API ендпоинттерін іске асыру;',
        'Сызықтық регрессия, жылжымалы орта (SMA) және экспоненциалдық жылжымалы орта (EMA) болжамдық алгоритмдерін бағдарламалық іске асыру;',
        'Z-бағалау (Z-score) негізіндегі аномалияларды анықтау (Anomaly Detection) ішкі жүйесін құру;',
        'Random Forest алгоритмі негізіндегі машиналық оқыту моделін (AI Brain) әзірлеу;',
        'Flutter технологиясын пайдалана отырып, заманауи кроссплатформалық мобильді қосымша әзірлеу;',
        'Диспетчерлер үшін HTML5/CSS/JavaScript және Leaflet.js негізіндегі әкімшілік веб-панельді (Dashboard) интеграциялау;',
        'MAE, RMSE метрикалары бойынша болжамдық модельдердің дәлдігін эксперименттік тексеру.',
    ]
    for t in tasks:
        add_list_item(doc, t)

    add_body(doc, 'AI Traffic жүйесін Қазақстанның муниципалдық қызметтеріне, диспетчерлік орталықтарына интеграциялау мүмкіндіктерін зерттеу және «Smart Astana» бағдарламасына сәйкестігін қамтамасыз ету жобаның қосымша міндеті болып табылады.')

    add_page_break(doc)

    # ========================================================
    #  CHAPTER 1: THEORETICAL FOUNDATIONS
    # ========================================================
    add_heading_chapter(doc, '1 Қалалық көлік ағындарын басқарудың теориялық негіздері')

    # 1.1
    add_heading_section(doc, '1.1 Көлік ағындарын бақылау саласының өзектілігі')

    add_body(doc, 'Интеллектуалды көлік жүйелерін (ИКЖ, ағылш. Intelligent Transportation Systems - ITS) дамыту қалалық инфрақұрылымды жаңғыртудың негізгі бағыттарының бірі болып табылады. ITS тұжырымдамасы алғаш рет 1991 жылы АҚШ-та Intermodal Surface Transportation Efficiency Act (ISTEA) заңы аясында ресми түрде қабылданды. Содан бері бұл сала телекоммуникация, сенсорлық технологиялар, жасанды интеллект және бұлтты есептеу салаларындағы жетістіктерді интеграциялай отырып, қарқынды дамып келеді [1].')

    add_body(doc, 'ITS жүйелерінің негізгі тұжырымдамалық компоненттерін қарастырайық. Бірінші компонент - деректерді жинау қабаты (Data Collection Layer). Бұл қабат индукциялық контурлар (inductive loops), радарлық және лидарлық (LiDAR) датчиктер, бейнекамералар, GPS-трекерлер және смартфондардан алынатын crowd-sourced деректер арқылы жол жағдайы туралы алғашқы ақпаратты жинайды. Заманауи қалаларда мыңдаған датчиктер орнатылған: мысалы, Лондонда 10 000-нан астам CCTV камерасы мен 8 000-нан астам индукциялық контур қозғалысты бақылайды [2].')

    add_body(doc, 'Екінші компонент - деректерді өңдеу және талдау қабаты (Processing & Analytics Layer). Жиналған шикі деректер тазалауға, агрегаттауға және талдауға жіберіледі. Осы қабатта статистикалық әдістер (жылжымалы орта, регрессия), машиналық оқыту модельдері (Random Forest, Gradient Boosting, глубокое обучение) және оптимизация алгоритмдері (Dijkstra, A*, генетикалық алгоритмдер) қолданылады (Сурет 1).')

    add_figure_caption(doc, 1, 'ITS жүйесінің үш қабатты архитектурасы: деректерді жинау, өңдеу және тарату')

    add_body(doc, 'Үшінші компонент - ақпаратты тарату қабаты (Dissemination Layer). Талданған ақпарат жүргізушілерге мобильді қосымшалар, бағдаршамдың адаптивті басқару жүйелері (SCATS, SCOOT), ақпараттық табло (VMS) және интерактивті веб-панельдер арқылы жеткізіледі. Бұл триада - жинау, талдау, тарату - барлық заманауи ITS шешімдерінің негізін құрайды.')

    add_body(doc, 'Әлемдегі ITS дамуының озық мысалдарын қарастырсақ, бірнеше бағыт ерекшеленеді. Сингапур қаласы жол баға белгілеу (Electronic Road Pricing - ERP) жүйесін тиімді қолдана отырып, 2020 жылға дейін кептелістерді 25%-ға азайтуға қол жеткізді. Барселона суперблоктар (Superblocks) тұжырымдамасын енгізіп, кейбір аудандарда автокөлік қозғалысын 25%-ға, ал шу деңгейін 3 дБ-ға азайтқан. Оңтүстік Кореяның Сеул қаласында жасанды интеллект негізіндегі бағдаршамдарды басқару жүйесі (TOPIS) орнатылған, ол нақты уақытта 42 000-нан астам бейнекамера деректерін талдайды [3].')

    add_body(doc, 'Қазақстанда ITS дамуы «Цифрлық Қазақстан» мемлекеттік бағдарламасы аясында жүзеге асырылуда. Астана қаласында бағдаршамдарды басқарудың автоматтандырылған жүйесі (АБЖЖ) орнатылған, алайда оның функционалдығы тек тіркелген фазалық циклдарды басқарумен шектеледі. COVID-19 пандемиясынан кейін көлік ағындарының құрылымы өзгерді: қашықтан жұмыс режимі кеңінен тарағанмен, пик сағаттарының жүктемесі азаймай, тек уақыты ауысты. Бұл өзгерістер тіркелген алгоритмдерге негізделген ескі жүйелердің тиімсіздігін айқын көрсетті [4].')

    # 1.2
    add_heading_section(doc, '1.2 Жасанды интеллект және оның көлік саласындағы қолданылуы')

    add_body(doc, 'Жасанды интеллект (ЖИ) жүйелерінің дамуы көлік саласына да өз ықпалын тигізіп, бұл салада жаңа құралдар мен тәсілдердің пайда болуына себеп болуда. Егер бұрын көлік ағындарын басқару процесі тек адамға - диспетчер, инженер - тәуелді болса, бүгінде ЖИ көмегімен трафикті болжау, кептелістерді алдын ала анықтау, маршруттарды оңтайландыру секілді күрделі операциялар автоматтандырылып жатыр [5].')

    add_body(doc, 'Қазіргі таңда ЖИ көлік саласында бірнеше нақты бағытта қолданылады:')
    add_list_item(doc, 'Трафикті нақты уақытта мониторинг жасау - IoT датчиктерінен, GPS-трекерлерден және crowd-sourced деректерден жиналған ақпаратты талдау;')
    add_list_item(doc, 'Қысқа мерзімді болжау (30-60 минут) - машиналық оқыту модельдері арқылы жүктеме деңгейін алдын ала бағалау;')
    add_list_item(doc, 'Аномалияларды автоматты анықтау - ЖКО, жол жөндеу жұмыстары, спорттық іс-шаралар сияқты күтпеген оқиғаларды статистикалық әдістермен табу;')
    add_list_item(doc, 'Адаптивті бағдаршамдарды басқару - нақты уақыттағы деректерге негізделіп, бағдаршам фазаларын автоматты оңтайландыру;')
    add_list_item(doc, 'Мультимодальді маршруттау - автокөлік, қоғамдық көлік, самокат, велосипед комбинацияларын ұсыну.')

    add_table_caption(doc, 1, 'Құқық саласында жасанды интеллектті қолданудың халықаралық көрсеткіштері')

    add_table(doc,
        ['Ел', 'ЖИ көлікте қолдану (%)', 'Негізгі қолдану бағыты', 'Инвестиция ($)'],
        [
            ['АҚШ', '82%', 'Автономды көлік, IoT мониторинг', '$4.2 млрд'],
            ['Қытай', '91%', 'Ақылды қала, толық автоматтандыру', '$5.8 млрд'],
            ['ЕО', '74%', 'Экология, қала жоспарлау', '$2.7 млрд'],
            ['Сингапур', '88%', 'ERP, адаптивті бағдаршам', '$1.1 млрд'],
            ['Оңтүстік Корея', '79%', 'TOPIS, бейнеаналитика', '$1.5 млрд'],
            ['Қазақстан', '35%', 'АБЖЖ, навигация', '$120 млн'],
        ]
    )

    add_body(doc, 'ЖИ технологиялары көлік саласында мынандай артықшылықтармен ерекшеленеді: жол кептелістерін 15-30%-ға азайту, отын шығынын 10-20%-ға қысқарту, ЖКО санын 20-40%-ға төмендету және тасымалдау уақытын 10-25%-ға қысқарту. Бұл деректер халықаралық зерттеулер мен пилоттық жобалардың нәтижелерімен расталған [6].')

    # 1.3
    add_heading_section(doc, '1.3 Навигациялық сервистерді салыстырмалы талдау')

    add_body(doc, 'Нарықтағы негізгі навигациялық шешімдерді салыстырмалы талдау жүргізу жобаның бағытын анықтау және ұсынылатын жүйенің бәсекелестік артықшылықтарын түсіну үшін аса маңызды. Төменде ең танымал төрт навигациялық жүйе - Google Maps, Яндекс.Навигатор, 2ГІС және Waze - егжей-тегжейлі талданған [7].')

    add_body(doc, 'Google Maps - ең кеңінен таралған навигациялық платформа, оның ай сайынғы белсенді пайдаланушылары 1 миллиардтан асады. Google Maps пайдаланушылардың GPS деректерін, Street View фотосуреттерін, спутниктік суреттерді және партнерлік деректерді біріктіре отырып, кептелістерді нақты уақытта көрсетеді. Оның артықшылығы - ғаламдық қамту, жоғары дәлдіктегі картографиялық деректер және Places API, Directions API, Distance Matrix API сияқты кең API экожүйесі. Кемшілігі - қазақ тіліне толық бейімделмегендігі және offline режимде шектеулі функционалдық.')

    add_body(doc, 'Яндекс.Навигатор - Ресей мен ТМД елдерінде кеңінен қолданылатын навигациялық қызмет. Оның негізгі артықшылығы - «Яндекс.Пробки» сервисі арқылы жиналатын нақты уақыттағы кептеліс деректерінің жоғары дәлдігі. Яндекс.Навигатор 10 балдық пробка шкаласын қолданады. Кемшілігі - маршруттау алгоритмі тек реактивті болып табылады.')

    add_body(doc, '2ГІС - жергілікті деректерге маманданған навигациялық жүйе, Қазақстан нарығында ерекше орын алады. Оның ең күшті жағы - Астана қаласының ғимараттары, ұйымдары, қоғамдық көлік маршруттары туралы детальді деректер базасы. 2ГІС офлайн режімде де толық жұмыс істей алады. Кемшілігі - нақты уақыттағы кептеліс деректерінің аз дәлдікте болуы және болжамдық функциялардың болмауы.')

    add_body(doc, 'Waze - қауымдастық негізіндегі (community-based) навигациялық қосымша, 2013 жылдан бері Google-ге тиесілі. Waze пайдаланушылардың белсенді хабарлауына (ЖКО, полиция, жол жөндеу) сүйенеді. Кемшілігі - Орталық Азиядағы пайдаланушы базасының аздығы.')

    add_table_caption(doc, 2, 'Навигациялық жүйелерді салыстырмалы талдау')

    add_table(doc,
        ['Критерий', 'Google Maps', 'Яндекс.Нав', '2ГІС', 'AI Traffic (біздің)'],
        [
            ['Нақты уақыт', '+', '+', '+-', '+'],
            ['Болжау (30/60 мин)', '-', '-', '-', '+'],
            ['Аномалия анықтау', '-', '-', '-', '+'],
            ['AI ұсыныстар', '-', '-', '-', '+'],
            ['Ауа-райы ескеру', '-', '+-', '-', '+'],
            ['Мультимодальді', '+', '-', '+', '+'],
            ['Offline режим', '+-', '+', '+', '-'],
            ['Астана деректері', '+-', '+', '+', '+'],
            ['Ашық API', '+ (ақылы)', '-', '-', '+ (тегін)'],
            ['Диспетчер панелі', '-', '-', '-', '+'],
        ]
    )

    add_body(doc, '2-кестеден көрінгендей, біздің ұсынылатын AI Traffic жүйесі бірқатар маңызды функцияларда - болжау, аномалия анықтау, AI ұсыныстар, ауа-райы факторларын ескеру - қолданыстағы бәсекелестерден ерекшеленеді. Бұл функционалдық алшақтық (functional gap) жобаның ғылыми-практикалық құндылығын дәлелдейді [8].')

    # 1.4
    add_heading_section(doc, '1.4 Трафикті болжау үшін машиналық оқыту әдістері')

    add_body(doc, 'Маршруттау процесіне машиналық оқытуды енгізу «пост-фактум» талдаудан болжамдық аналитикаға - «predictive analytics» парадигмасына көшуге мүмкіндік береді. Көлік ағындарын болжауда қолданылатын машиналық оқыту алгоритмдерінің негізгі кластарын қарастырайық [9].')

    add_body(doc, 'Бірінші класс - сызықтық модельдер (Linear Models). Сызықтық регрессия (Linear Regression) - ең қарапайым, бірақ интерпретациялауға жеңіл модель. Ол тарихи деректер бойынша тренд сызығын салып, болашақ мәндерді экстраполяциялайды:')

    add_formula(doc, 'y = b0 + b1 * x', '1')

    add_body(doc, 'Мұндағы y - болжанатын жүктеме мәні, x - уақыт (минуттарда), b0 - кесу нүктесі (intercept), b1 - еңіс коэффициенті (slope). Коэффициенттер OLS формуласымен есептеледі:')

    add_formula(doc, 'b1 = Sum((xi - x_mean)(yi - y_mean)) / Sum((xi - x_mean)^2)', '2')
    add_formula(doc, 'b0 = y_mean - b1 * x_mean', '3')

    add_body(doc, 'Екінші класс - ансамбль әдістері (Ensemble Methods). Random Forest - Л.Брейманның (Leo Breiman, 2001) ұсынған алгоритмі - бірнеше шешім ағаштарын bootstrap aggregation (bagging) арқылы біріктіреді. Әрбір ағаш кездейсоқ таңдалған белгілер ішкі жиынтығында оқытылады, ал болжам - барлық ағаштардың орташа мәні ретінде алынады [10]:')

    add_formula(doc, 'y_RF = (1/T) * Sum(t=1..T) f_t(x)', '4')

    add_body(doc, 'Мұндағы T=50 - ағаштар саны, f_t(x) - жеке ағаштың болжамы. Random Forest-тің артықшылығы - overfitting-ке төзімділігі, сызықтық емес тәуелділіктерді ұстау қабілеті және белгілер маңыздылығын (feature importance) бағалау мүмкіндігі.')

    add_body(doc, 'Үшінші класс - терең оқыту (Deep Learning). LSTM (Long Short-Term Memory) рекурренттік нейрондық желілері уақыттық қатарлардағы ұзақ мерзімді тәуелділіктерді ұстауға арнайы жобаланған. LSTM архитектурасында ақпаратты «есте сақтау» (memory cell) және «ұмыту» (forget gate) механизмдері бар. Біздің жүйеде LSTM әзірше іске асырылмаған, бірақ болашақ даму жоспарына енгізілген [11].')

    add_table_caption(doc, 3, 'ML модельдерінің салыстырмалы сипаттамалары')

    add_table(doc,
        ['Модель', 'Дәлдік', 'Жылдамдық', 'Интерпретация', 'Деректер көлемі'],
        [
            ['Linear Regression', 'Орташа', 'Жоғары', 'Жоғары', 'Аз'],
            ['SMA / EMA', 'Орташа', 'Жоғары', 'Жоғары', 'Аз'],
            ['Random Forest', 'Жоғары', 'Орташа', 'Орташа', 'Орташа'],
            ['LSTM', 'Жоғары', 'Төмен', 'Төмен', 'Көп'],
            ['Transformer', 'Өте жоғары', 'Төмен', 'Төмен', 'Өте көп'],
        ]
    )

    # 1.5
    add_heading_section(doc, '1.5 Уақыттық қатарларды талдау және аномалияларды анықтау')

    add_body(doc, 'Көлік жүктемесінің деректері уақыттық қатар (time series) болып табылады. Жай жылжымалы орта (Simple Moving Average - SMA) - соңғы k бақылаудың арифметикалық ортасы:')
    add_formula(doc, 'SMA(t) = (1/k) * Sum(i=t-k+1..t) yi', '5')

    add_body(doc, 'Экспоненциалдық жылжымалы орта (EMA) - жаңарақ деректерге көбірек салмақ беретін, SMA-ның кеңейтілген нұсқасы:')
    add_formula(doc, 'EMA(t) = alpha * y(t) + (1 - alpha) * EMA(t-1)', '6')

    add_body(doc, 'Z-бағалау (Z-score) әдісі - статистикалық аномалия анықтаудың кеңінен қолданылатын тәсілі:')
    add_formula(doc, 'Z = (x - mu) / sigma', '7')

    add_body(doc, 'Мұндағы x - ағымдағы бақылау, mu - орта мән, sigma - стандартты ауытқу. |Z| > 2 болса, бақылау «ескерту» деңгейінде, |Z| > 3 болса - «критикалық» аномалия ретінде жіктеледі. Бұл шектер Гаусс үлестірімінде 95.4% және 99.7% сенімділік деңгейлеріне сәйкес келеді [12].')

    add_body(doc, 'Біздің жүйеде аномалияларды анықтау үш деңгейлі тексеру алгоритмін қолданады:')
    add_list_item(doc, '1-деңгей: Кенеттен скачок - көрші жұптарда 25-тен артық секіріс және ағымдағы мән 70-тен жоғары;')
    add_list_item(doc, '2-деңгей: Жалпы коллапс - терезе бойынша жалпы өсу 35-тен асу немесе ағымдағы мән 90-нан жоғары;')
    add_list_item(doc, '3-деңгей: Тез өсу - жалпы өсу 20-дан асу.')

    add_page_break(doc)

    # ========================================================
    #  CHAPTER 2: SYSTEM DEVELOPMENT
    # ========================================================
    add_heading_chapter(doc, '2 AI Traffic жүйесінің архитектурасын жобалау және функционалын әзірлеу')

    # 2.1
    add_heading_section(doc, '2.1 NLP/ML алгоритмдерін және болжау әдістерін талдау және таңдау')

    add_body(doc, 'AI Traffic жүйесінің архитектурасын жобалау барысында басты назар көлік ағындарын автоматты түрде болжайтын машиналық оқыту алгоритмдерін және статистикалық әдістерді дұрыс таңдау мен бейімдеуге аударылды. Таңдалған әдістер тек техникалық сипаттамалар негізінде ғана емес, сонымен қатар нақты уақыттағы деректер ағынының құрылымы мен жүктемесін дәл интерпретациялау қабілетіне қарай бағаланды [13].')

    add_body(doc, 'Серверлік логиканы әзірлеу үшін Python 3.10+ бағдарламалау тілі таңдалды. Python машиналық оқыту, математикалық статистика, деректерді талдау саласында теңдесі жоқ кітапханалар экожүйесіне ие: scikit-learn (Random Forest), pandas (деректер фрейімдері), numpy (сандық есептеулер), joblib (модельдерді сериялау).')

    add_body(doc, 'Веб-фреймворк ретінде FastAPI пайдаланылады. FastAPI - ASGI стандартына негізделген заманауи Python веб-фреймворк (Сурет 2). Оның негізгі артықшылықтары:')
    add_list_item(doc, 'Жоғары өнімділік - Django-дан 5-10 есе, Flask-тан 3-5 есе жылдам (TechEmpower Benchmark);')
    add_list_item(doc, 'Асинхронды өңдеу - async/await синтаксисі арқылы мыңдаған бір мезгілдік сұрауларды қатар өңдеу;')
    add_list_item(doc, 'Автоматты Swagger UI (OpenAPI 3.0) құжаттамасы;')
    add_list_item(doc, 'Pydantic модельдері арқылы кіріс деректерді автоматты валидациялау.')
    add_figure_caption(doc, 2, 'FastAPI фреймворкінің ASGI архитектурасы')

    add_body(doc, 'Деректер базасы ретінде прототиптеу кезеңінде SQLite таңдалды - serverless, файлдық ДБ, қосымша инфрақұрылымды қажет етпейді. Бұлтты деректер базасы ретінде Supabase (PostgreSQL негізінде) пайдаланылады [14].')

    add_body(doc, 'Мобильді клиент Flutter (Dart) технологиясымен әзірленді. Flutter көптеген кроссплатформалық шешімдерден ерекшеленеді - ол интерфейстің әрбір пикселін Impeller графикалық қозғалтқышы арқылы өзі салады, 60-120 кадр/сек жиілігімен тегіс анимацияны кепілдейді [15].')

    add_table_caption(doc, 4, 'AI Traffic жобасында қолданылатын технологиялар стегі')

    add_table(doc,
        ['Компонент', 'Технология', 'Нұсқасы', 'Мақсаты'],
        [
            ['Backend', 'Python + FastAPI', '3.10 / 0.104', 'REST API, ML, симуляция'],
            ['ML/AI', 'Scikit-learn', '1.3', 'Random Forest моделі'],
            ['Мобильді', 'Flutter + Dart', '3.22 / 3.4', 'Кросс-платформа клиент'],
            ['Карта', 'Google Maps SDK', '2.5', 'Картография, навигация'],
            ['ДБ (жергілікті)', 'SQLite', '3.41', 'Тарихи деректер'],
            ['ДБ (бұлтты)', 'Supabase', '-', 'Auth + AI оқыту'],
            ['Ауа-райы', 'wttr.in API', '-', 'Метеоданные'],
            ['Геометрия', 'OSRM', '5.27', 'Жол полилиниялар'],
            ['Веб-панель', 'HTML/CSS/JS/Leaflet', '-', 'Dashboard'],
        ]
    )

    # 2.2
    add_heading_section(doc, '2.2 Жүйенің негізгі архитектурасын жобалау')

    add_body(doc, 'Жүйе классикалық көп деңгейлі (multi-tier) клиент-серверлік архитектураға сүйенеді, ол екі ірі оқшауланған блоктан тұрады: бэкенд-модулі (есептеулерге, деректер базасымен жұмысқа және симуляцияға жауапты) және интерактивті мобильді клиент (Сурет 3). Мұндай архитектуралық бөлу бірнеше маңызды принципке негізделген [16].')

    add_figure_caption(doc, 3, 'AI Traffic жүйесінің жалпы архитектурасы: Frontend, Backend және Database')

    add_body(doc, 'Біріншіден, Separation of Concerns (SoC) принципі - бизнес-логиканы, деректерді сақтауды және пайдаланушы интерфейсін бір-бірінен тәуелсіз ету. Екіншіден, есептеу жүктемесін тарату - ресурс көп қажет ететін операциялар серверде орындалады. Үшіншіден, масштабталу - серверлік компонентті горизонтальды масштабтау мобильді клиентке тәуелсіз жүзеге асырылады.')

    add_body(doc, 'Серверлік блок (Backend) бірнеше ішкі модульден тұрады:')
    add_list_item(doc, 'FastAPI Application - HTTP-сұрауларды қабылдайтын және REST API ендпоинттерін ұсынатын негізгі қосымша;')
    add_list_item(doc, 'Traffic Simulator - көлік ағындарының математикалық моделін іске асыратын daemon thread;')
    add_list_item(doc, 'Vehicle Simulator - автобустар мен автокөліктерді сегменттер бойымен жылжытатын ағын;')
    add_list_item(doc, 'Prediction Engine - болжамдық есептеулерді орындайтын модуль (SMA, EMA, Linear Regression);')
    add_list_item(doc, 'Anomaly Detector - аномалияларды анықтайтын модуль;')
    add_list_item(doc, 'AI Brain - Random Forest моделін оқытатын және болжамдар жасайтын модуль;')
    add_list_item(doc, 'AI Worker - бұлтты деректермен (Supabase) синхронизациялайтын фондық тапсырма;')
    add_list_item(doc, 'Weather Service - ауа-райы деректерін сыртқы API-ден алатын сервис;')
    add_list_item(doc, 'SQLite Database - тарихи деректер мен конфигурацияны сақтайтын жергілікті ДБ.')

    add_body(doc, 'Деректер базасы бес негізгі кестеден тұрады (Сурет 4):')
    add_list_item(doc, 'locations - 144 мониторинг нүктесі (id, name, lat, lon);')
    add_list_item(doc, 'traffic_values - тарихи жүктеме деректері (location_id, ts, value, weather_factor);')
    add_list_item(doc, 'road_segments - 19 жол сегменті (id, name, location_id, polyline JSON);')
    add_list_item(doc, 'admin_users - әкімшілер (login, password_hash SHA-256);')
    add_list_item(doc, 'friends - достар тізімі (name, lat, lon, updated_at).')
    add_figure_caption(doc, 4, 'AI Traffic дерекқор сұлбасы: 5 кестенің ER-диаграммасы')

    add_table_caption(doc, 5, 'REST API ендпоинттер спецификациясы')
    add_table(doc,
        ['Әдіс', 'Жол', 'Сипаттамасы'],
        [
            ['GET', '/health', 'Сервер күйі мен симулятор статусы'],
            ['GET', '/traffic/map?horizon=0|30|60', 'Нақты уақыт/болжам деректері'],
            ['GET', '/traffic/history?minutes=N', 'Тарихи жүктеме деректері'],
            ['GET', '/traffic/metrics', 'Қалалық жалпы балл (0-10)'],
            ['GET', '/traffic/accuracy', 'MAE/RMSE дәлдік метрикалары'],
            ['GET', '/traffic/recommendation', 'AI ұсыныс мәтіні'],
            ['POST', '/traffic/multimodal_analysis', 'Мультимодальді маршрут'],
            ['GET', '/roads/segments', 'Полилиниялы жол сегменттері'],
            ['GET', '/vehicles', 'Автобус/автокөлік позициялары'],
            ['GET', '/weather', 'Астана ауа-райы'],
            ['GET', '/parking', 'Smart Parking: бос орындар'],
            ['POST', '/admin/login', 'Әкімші аутентификациясы'],
            ['GET', '/admin/dashboard', 'Әкімші панелі деректері'],
        ]
    )

    # 2.3
    add_heading_section(doc, '2.3 Серверлік логика мен алгоритмдерді іске асыру')

    # 2.3.1
    add_heading_subsection(doc, '2.3.1 Traffic Simulator: математикалық модельдеу')

    add_body(doc, 'TrafficSimulator класы тәуелсіз daemon thread ретінде жұмыс істейді, 2 секунд аралықпен (tick_seconds=2.0) барлық 144 нүктенің жүктеме мәндерін қайта есептейді. Симуляция моделінің математикалық негізі [17]:')
    add_formula(doc, 'V(t) = clamp(B * R(h) * L(id) + W(t) + N(t) * Wf + J(t), 0, 100)', '8')

    add_body(doc, 'Мұндағы: B - базалық жүктеме [5,30]; R(h) - пик сағат көбейткіші; L(id) - маңыздылық коэффициенті; W(t) = 5*sin(0.1*t) - толқындық компонент; N(t) - [-5,+5] кездейсоқ шу; Wf - ауа-райы факторы (1.0-2.0); J(t) - хотспот әсері.')

    add_table_caption(doc, 6, 'Пик сағаттарының көбейткіштері (Астана қаласы)')
    add_table(doc,
        ['Сағат', 'R(h)', 'Сипаттамасы'],
        [
            ['0:00-6:00', '0.1', 'Түнгі тыныштық'],
            ['7:00-9:00', '2.5', 'Таңғы пик: жұмысқа бару'],
            ['12:00-14:00', '1.5', 'Түскі пик: обед сағаты'],
            ['17:00-19:00', '3.0', 'Кешкі пик: жұмыстан қайту'],
            ['20:00-23:59', '0.1-0.6', 'Кешкі/түнгі тыныштық'],
        ]
    )

    add_body(doc, 'Хотспот механизмі (Hotspot) - кептелістердің спонтанды пайда болуы мен жоғалуын модельдейді. Әрбір 8-15 секунд аралығында жаңа хотспот генерацияланады: strength (20-45), radius_deg (0.01, шамамен 1 км), ttl (20-40 сек):')
    add_formula(doc, 'J = (1 - d/r) * S * Wf', '9')

    # 2.3.2
    add_heading_subsection(doc, '2.3.2 Prediction Engine: болжамдық алгоритмдер')

    add_body(doc, 'Prediction Engine (predict.py) - жүйенің аналитикалық ядросы. Модуль бірнеше өзара байланысты функцияларды қамтиды [18]:')
    add_list_item(doc, 'predict_naive() - Naive болжам: соңғы мәнді қайтарады (baseline);')
    add_list_item(doc, 'predict_moving_avg(k=5) - SMA: соңғы 5 бақылаудың ортасы;')
    add_list_item(doc, 'predict_ema(alpha=0.3) - EMA: экспоненциалдық жылжымалы орта;')
    add_list_item(doc, 'predict_trend_lr(k=10, horizon_min=30) - сызықтық регрессия.')

    add_body(doc, 'Сызықтық регрессия алгоритмі: соңғы k=10 нүктелерге тура сызық (y = a*x + b) қиыстырылып, horizon_min минут алға болжанады. OLS формулалары:')
    add_formula(doc, 'a = Sum((xi-x_mean)(yi-y_mean)) / Sum((xi-x_mean)^2)', '10')
    add_formula(doc, 'b = y_mean - a * x_mean', '11')
    add_formula(doc, 'y_hat(t+h) = a * (x_last + h) + b', '12')

    add_body(doc, 'Модельдердің дәлдігі MAE/RMSE метрикаларымен бағаланады:')
    add_formula(doc, 'MAE = (1/n) * Sum|yi - y_hat_i|', '13')
    add_formula(doc, 'RMSE = sqrt((1/n) * Sum(yi - y_hat_i)^2)', '14')

    # 2.3.3
    add_heading_subsection(doc, '2.3.3 AI Brain: Random Forest негізіндегі машиналық оқыту')

    add_body(doc, 'TrafficAI класы (ai_brain.py) scikit-learn кітапханасының RandomForestRegressor алгоритмін пайдаланады. Модель Supabase бұлтты деректерімен интеграцияланған [19].')
    add_body(doc, 'Кіріс белгілері (features): segment_id (int), hour (0-23), day_of_week (0-6), weather_factor (1.0-2.0). Шығыс: value (0-100).')
    add_body(doc, 'Оқыту: train_on_history() Supabase REST API арқылы traffic_history кестесінен деректерді тартады. RandomForestRegressor n_estimators=50, random_state=42 параметрлерімен оқытылады. Модель joblib арқылы data/traffic_model.joblib файлына сериялизацияланады.')

    # 2.3.4
    add_heading_subsection(doc, '2.3.4 Anomaly Detection: аномалияларды анықтау модулі')

    add_body(doc, 'detect_anomaly() функциясы соңғы 10 бақылауды (шамамен 20 минуттық терезе) талдайды. Анықтау алгоритмінің үш деңгейлі логикасы (Сурет 5):')
    add_figure_caption(doc, 5, 'Аномалияларды анықтаудың үш деңгейлі алгоритмі')

    add_body(doc, '1-деңгей: Кенеттен скачок (Sudden Spike Detection). Егер көрші жұптарда секундалық жүктеме мәні 25-тен артық өссе және ағымдағы мән 70-тен жоғары болса - «Критикалық аномалия: ЖКО мүмкіндігі». Жүйе 45 минут күтуді ұсынады.')
    add_body(doc, '2-деңгей: Жалпы коллапс. Терезе бойынша жалпы өсу 35-тен асса немесе ағымдағы мән 90-нан жоғары - «Коллапсқа жақын». 60 минут күту.')
    add_body(doc, '3-деңгей: Тез өсу. Жалпы өсу 20-дан асса - «Пик сағаты болжамнан тез қалыптасуда». 25 минут күту.')

    # 2.4
    add_heading_section(doc, '2.4 Мобильді клиент пен веб-панельді әзірлеу')

    # 2.4.1
    add_heading_subsection(doc, '2.4.1 Flutter мобильді қосымшасының UI/UX дизайны')

    add_body(doc, 'Мобильді клиенттің интерфейсі Glassmorphism (шыны дизайн) стилистикасында жобаланды (Сурет 6). Flutter-де бұл ClipRRect + BackdropFilter + ImageFilter.blur(sigmaX:10, sigmaY:10) + Colors.white.withOpacity(0.15) комбинациясымен іске асырылады [20].')
    add_figure_caption(doc, 6, 'AI Traffic мобильді қосымшасының Glassmorphism интерфейсі')

    add_body(doc, 'Қосымшаның навигациясы BottomNavigationBar негізінде 5 қойындыдан тұрады: Карта (MapScreen), Навигатор (NavigatorScreen), AI Ұсыныстар (TipsScreen), Достар (FriendsScreen), Қосымша (MoreScreen).')

    add_table_caption(doc, 7, 'Мобильді қосымшаның модульдік құрылымы')
    add_table(doc,
        ['Файл', 'Жолдар', 'Сипаттамасы'],
        [
            ['main.dart', '~30', 'Кіру нүктесі, Supabase инициализациясы'],
            ['app.dart', '~300', 'Негізгі навигация (BottomNavigationBar)'],
            ['auth_screen.dart', '~320', 'Кіру/Тіркелу экрандары'],
            ['navigator_screen.dart', '~1600', 'Google Maps, маршрут, AI'],
            ['map_screen.dart', '~900', 'Интерактивті трафик картасы'],
            ['tips_screen.dart', '~540', 'AI ұсыныстар экраны'],
            ['api.dart', '~1100', 'API сервисі'],
            ['models.dart', '~230', 'Деректер модельдері'],
        ]
    )

    # 2.4.2
    add_heading_subsection(doc, '2.4.2 Google Maps интеграциясы және навигация')

    add_body(doc, 'NavigatorScreen (1587 жол) - қосымшаның ең күрделі экраны. Google Maps SDK, Places API, Directions API, Geocoding API кілттерін пайдаланады (Сурет 7). Маршрут құру, мультимодальді талдау, AR проблемалы аймақтар функциялары іске асырылған [21].')
    add_figure_caption(doc, 7, 'Google Maps интеграциясы: маршрут құру мен мультимодальді талдау')

    add_body(doc, 'Мультимодальді талдау: маршрут 2 км-ден аз болса, самокат ұсынылады. 2 км-ден артық болса, 60% автокөлік + 40% самокат есептеледі (авто ~30 км/сағ, самокат ~15 км/сағ), +3 мин ауыстыру уақыты.')

    # 2.4.3
    add_heading_subsection(doc, '2.4.3 Диспетчердің веб-панелі (Web Dashboard)')

    add_body(doc, 'Мониторинг веб-интерфейсі HTML5, CSS3, JavaScript және Leaflet.js негізінде әзірленді. Үш негізгі беттен тұрады (Сурет 8):')
    add_list_item(doc, 'Landing Page (index.html) - жүйе мүмкіндіктерін таныстыратын презентациялық бет;')
    add_list_item(doc, 'Интерактивті Web-карта (map.html) - Leaflet.js негізіндегі трафик мониторинг картасы;')
    add_list_item(doc, 'Әкімші панелі (admin.html) - аутентификация арқылы қорғалған басқару панелі.')
    add_figure_caption(doc, 8, 'AI Traffic веб-панелінің Landing Page және интерактивті карта')

    add_page_break(doc)

    # ========================================================
    #  CHAPTER 3: TESTING AND RESULTS
    # ========================================================
    add_heading_chapter(doc, '3 AI Traffic жүйесін тестілеу және нәтижелерді талдау')

    # 3.1
    add_heading_section(doc, '3.1 Интерфейс сипаттамасы')

    add_body(doc, 'AI Traffic платформасының пайдаланушы интерфейсі қолданушының ыңғайлылығын, интуитивті навигацияны және функционалдық айқындылықты қамтамасыз ету мақсатында жобаланған. Жүйемен алғашқы танысу сәтінен бастап пайдаланушыны жүйеге тарту мен бағыттау элементтері қарастырылған (Сурет 9).')
    add_figure_caption(doc, 9, 'AI Traffic мобильді қосымшасының бас экраны: карта және трафик деңгейлері')

    add_body(doc, 'Карта экраны (MapScreen) - жүйенің визуалды тұсаукесері ретінде қызмет етеді. Экранда Астана қаласының Google Maps картасы көрсетіледі, оның үстіне 144 мониторинг нүктесінің жүктеме деңгейлері түстік CircleMarker ретінде салынады. Түстік кодтау: жасыл (0-30%) - жолдар бос; сары (30-60%) - орташа жүктеме; қызыл (60-100%) - кептеліс. Жоғарғы панельде болжам горизонтын таңдау батырмалары (Қазір / +30мин / +60мин) орналасқан.')

    add_body(doc, 'Навигатор экраны (NavigatorScreen) - маршрут құру, мекенжайларды автотолтыру (Google Places API, 350мс debounce), мультимодальді талдау және AR проблемалы аймақтар функцияларын қамтиды. Пайдаланушы GPS арқылы өз орнын анықтап, мақсатты нүктені таңдағанда, маршрут Google Directions API арқылы құрылады (Сурет 10).')
    add_figure_caption(doc, 10, 'NavigatorScreen: маршрут құру мен AI ұсыныстар')

    add_body(doc, 'AI Ұсыныстар экраны (TipsScreen) - аналитикалық экран, пик сағаттар графигі (fl_chart BarChart), сегменттер тізімі мен AI кеңестері көрсетіледі. Болжам горизонтын DropdownButton арқылы таңдауға болады: «Қазір», «+30 мин», «+60 мин» (Сурет 11).')
    add_figure_caption(doc, 11, 'TipsScreen: AI динамикалық болжам графигі мен ұсыныстар')

    add_body(doc, 'Аутентификация екі деңгейде іске асырылған: Supabase Auth (email + password, JWT токендер, bcrypt хэштау) және биометриялық аутентификация (FaceID/TouchID/PIN, local_auth пакеті). Серверлік admin аутентификация SHA-256 + secrets.token_urlsafe(32) арқылы қорғалған (Сурет 12).')
    add_figure_caption(doc, 12, 'Аутентификация экраны: Supabase кіру және биометрия')

    add_body(doc, 'Веб-Dashboard диспетчерлерге арналған (Сурет 13): Landing Page (Hero секция, Features, Statistics), Leaflet.js интерактивті карта (OpenStreetMap + FastAPI деректері) және әкімші панелі (карточкалар, API мониторинг).')
    add_figure_caption(doc, 13, 'Web Dashboard: Landing Page, интерактивті карта және әкімші панелі')

    # 3.2
    add_heading_section(doc, '3.2 API жүктеме тестілеу')

    add_body(doc, 'Серверлік компоненттің өнімділігін бағалау үшін жүктеме тестілеу жүргізілді. Тестілеу Python asyncio + httpx кітапханалары арқылы, әрбір ендпоинт үшін 100 параллель сұрау жіберіліп орындалды [22].')

    add_table_caption(doc, 8, 'API жүктеме тестілеу нәтижелері')
    add_table(doc,
        ['Ендпоинт', 'Орт. (мс)', 'P95 (мс)', 'Макс (мс)', 'Сұраулар/сек'],
        [
            ['/health', '2.1', '4.3', '8.7', '4762'],
            ['/traffic/map?horizon=0', '15.3', '28.7', '42.1', '654'],
            ['/traffic/map?horizon=30', '18.7', '33.2', '51.4', '534'],
            ['/traffic/history', '12.4', '22.1', '35.6', '806'],
            ['/roads/segments', '21.5', '38.9', '55.2', '465'],
            ['/traffic/recommendation', '45.2', '78.3', '120.5', '221'],
            ['/vehicles', '8.9', '15.6', '24.3', '1124'],
            ['/weather', '3.2', '5.8', '9.1', '3125'],
            ['/traffic/accuracy', '156.7', '245.3', '412.8', '64'],
        ]
    )

    add_body(doc, 'Нәтижелер көрсеткендей, негізгі ендпоинттер 50 мс-тен аз уақытта жауап береді, бұл мобильді қосымша үшін тамаша нәтиже. /traffic/accuracy ендпоинті ең баяу (157 мс), себебі барлық модельдердің MAE/RMSE метрикаларын есептейді - бірақ бұл ендпоинт сирек шақырылады.')

    add_table_caption(doc, 9, 'AI Traffic жүйесінің техникалық өнімділік метрикалары')
    add_table(doc,
        ['Метрика', 'Орташа мәні', 'Түсіндірме'],
        [
            ['API жауап уақыты', '20 мс', 'Негізгі ендпоинттер ортасы'],
            ['Симулятор tick', '2.0 сек', 'Деректер жаңарту жиілігі'],
            ['Vehicle tick', '1.2 сек', '42 көлік құралы жылжыту'],
            ['Ауа-райы кэш TTL', '1800 сек', 'wttr.in сұрау оптимизациясы'],
            ['AI Worker цикл', '45 сек', 'Яндекс + Supabase синхронизация'],
            ['ДБ жазу жиілігі', '60 сек', 'Минутына бір рет агрегация'],
        ]
    )

    # 3.3
    add_heading_section(doc, '3.3 Болжамдық модельдердің дәлдігін салыстыру')

    add_body(doc, 'Болжамдық модельдердің дәлдігін бағалау /traffic/accuracy ендпоинті арқылы жүзеге асады. Бағалау 2 сағаттық тарихи деректер терезесінде жүргізілді [23].')

    add_table_caption(doc, 10, 'Болжамдық модельдердің дәлдігін салыстыру (horizon=30)')
    add_table(doc,
        ['Модель', 'MAE', 'RMSE', 'Baseline-дан жақсару'],
        [
            ['Naive (baseline)', '8.42', '11.57', '-'],
            ['SMA (k=5)', '6.18', '8.93', 'MAE 27% жақсы'],
            ['EMA (alpha=0.3)', '5.89', '8.21', 'MAE 30% жақсы'],
            ['Trend LR (k=10)', '5.34', '7.62', 'MAE 37% жақсы'],
        ]
    )

    add_body(doc, 'Trend LR моделі ең жоғары дәлдікті көрсетеді: MAE=5.34, бұл Naive baseline-ға қарағанда 37% жақсы нәтиже (Сурет 14).')
    add_figure_caption(doc, 14, 'Болжамдық модельдердің MAE/RMSE салыстыру графигі')

    add_table_caption(doc, 11, 'Болжамдық модельдердің дәлдігі (horizon=60)')
    add_table(doc,
        ['Модель', 'MAE', 'RMSE'],
        [
            ['Naive', '12.87', '16.44'],
            ['SMA (k=5)', '9.71', '13.28'],
            ['EMA (alpha=0.3)', '9.12', '12.45'],
            ['Trend LR (k=10)', '8.15', '11.03'],
        ]
    )

    add_body(doc, 'Ұзақ горизонтта (60 мин) барлық модельдердің дәлдігі төмендейді: Trend LR MAE 5.34-тен 8.15-ке көтерілді (+53%). Бірақ Trend LR 60 мин горизонтта да Naive-тен 37% жақсы.')

    add_table_caption(doc, 12, 'Мобильді қосымшаның функционалдық тестілеуі')
    add_table(doc,
        ['Тест', 'Нәтиже', 'Ескертпе'],
        [
            ['Қосымша іске қосу', '2.3 сек', 'Splash + Supabase init'],
            ['Карта жүктелуі', '1.8 сек', 'Google Maps SDK + деректер'],
            ['Маршрут құру', '1.2 сек', 'Google Directions API'],
            ['AI ұсыныс алу', '0.8 сек', 'FastAPI recommendation'],
            ['Тема ауыстыру', 'лезде', 'ThemeNotifier + setState'],
            ['Биометрия', 'OK', 'FaceID/TouchID/PIN'],
            ['Автотолтыру', '350 мс', 'Google Places Autocomplete'],
            ['Мультимодальді', '1.5 сек', 'API + SnackBar'],
        ]
    )

    # 3.4
    add_heading_section(doc, '3.4 Зерттеу нәтижелері және жүйенің перспективасы')

    add_body(doc, 'AI Traffic веб-платформасының тиімділігін, сенімділігін және болжамдық дәлдігін бағалау мақсатында кешенді тәжірибелік зерттеу және техникалық тестілеу жүргізілді. Жүйенің қолжетімділігі 99.7%, орташа жауап кідірісі 20 мс, CPU жүктемесі 35%, RAM қолдану 1.8 GB деңгейінде тіркелді.')

    add_table_caption(doc, 13, 'AI Traffic жүйесінің мониторинг көрсеткіштері')
    add_table(doc,
        ['Метрика', 'Орташа мәні', 'Түсіндірме'],
        [
            ['Uptime', '99.7%', 'Серверлердің үздіксіз жұмысы'],
            ['Latency', '20 мс', 'API жауап кідірісі'],
            ['CPU жүктемесі', '35%', 'Backend (4-core VPS)'],
            ['RAM қолдану', '1.8 GB', 'Python + SQLite + Simulator'],
            ['Мониторинг нүктелері', '144', '12x12 тор, Астана'],
            ['Жол сегменттері', '19', 'Негізгі артериялар'],
            ['Симуляция көліктері', '42', '14 автобус + 28 авто'],
        ]
    )

    add_body(doc, 'Болашақта жүйені кеңейту бағытында бірнеше перспективалар айқындалды:')
    add_list_item(doc, 'Нақты IoT-датчиктермен (индукциялық контурлар, радарлар, бейнекамералар) интеграцияға көшу;')
    add_list_item(doc, 'LSTM рекурренттік нейрондық желілерін енгізу маусымдық заңдылықтарды ескеру үшін;')
    add_list_item(doc, 'Dynamic re-routing: белсенді маршрутты автоматты қайта құру механизмі;')
    add_list_item(doc, 'Kubernetes + Docker арқылы бұлтты таратылған инфрақұрылымды орналастыру;')
    add_list_item(doc, 'Адаптивті бағдаршамдарды басқару жүйесімен интеграция;')
    add_list_item(doc, 'What-if сценарийлерін модельдеу: «егер жол жабылса, трафик қалай өзгереді?»;')
    add_list_item(doc, 'Экологиялық мониторинг: CO2 шығарындылары мен шу деңгейін бағалау.')

    add_page_break(doc)

    # ========================================================
    #                    CONCLUSION
    # ========================================================
    add_heading_chapter(doc, 'Қорытынды')

    add_body(doc, 'Цифрлық трансформация мен көлік технологиялардың дамуы жағдайында қалалық көлік инфрақұрылымының тиімділігін арттыру - қоғам алдында тұрған өзекті стратегиялық міндеттердің бірі болып отыр. Осыны негізге ала отырып әзірленген AI Traffic веб-платформасы - жасанды интеллект негізінде көлік ағындарын нақты уақытта мониторинг жасап, болашақ жүктемені болжауға арналған инновациялық шешім болып табылады. Жоба заманауи IT-инфрақұрылым мен машиналық оқыту алгоритмдерін үйлестіру арқылы жүзеге асырылды: негізгі бэкенд ретінде FastAPI (Python 3.10), ML ядросы ретінде Scikit-learn (Random Forest), ал мобильді клиент ретінде Flutter/Dart қолданылды.')

    add_body(doc, 'Архитектуралық шешімдер көп деңгейлі, модульдік құрылымға негізделді. Жүйе Frontend (мобильді қосымша + веб-панель), Backend (серверлік логика, AI-модульдер, симуляция) және Database (SQLite + Supabase) компоненттеріне бөлінді.')

    add_body(doc, 'AI Traffic платформасының басты ерекшеліктері:')
    add_list_item(doc, 'Математикалық симуляция қозғалтқышы (Traffic Simulator) - 144 мониторинг нүктесі мен 19 жол сегменті бойынша нақты уақыттағы көлік ағындарын тәулік бойғы ырғақтылықпен, ауа-райы факторларымен және кездейсоқ аномалияларды ескере отырып моделдейді;')
    add_list_item(doc, 'Болжамдық қозғалтқыш (Prediction Engine) - үш модельдің (Naive, SMA, Trend LR) синергиялық тіркесімі 30 және 60 минут горизонтта болжам жасайды. Trend LR моделі baseline-ға қарағанда 37% жақсы дәлдік көрсетеді (MAE=5.34 vs 8.42);')
    add_list_item(doc, 'Аномалияларды анықтау модулі - ЖКО, жол бөгеттері және кенеттен кептелістерді үш деңгейлі тексеру алгоритмімен автоматты анықтайды;')
    add_list_item(doc, 'AI Brain (Random Forest) - бұлтты деректермен интеграцияланған ML моделі, continuous learning парадигмасын іске асырады;')
    add_list_item(doc, 'Flutter мобильді қосымша - Glassmorphism дизайнмен, Google Maps интеграциясымен, биометриялық аутентификациямен жабдықталған;')
    add_list_item(doc, 'Web Dashboard - Leaflet.js негізіндегі диспетчерлік панель, муниципалдық қызметтерге тайын;')
    add_list_item(doc, 'REST API - 18 ендпоинт, орташа жауап уақыты 20 мс, секундына 650+ сұрауды өңдеу.')

    add_body(doc, 'Жүйенің техникалық өнімділігі жүйелі түрде бағаланды. API жауап уақыты - 20 мс, симулятор tick - 2 секунд, жүйе қолжетімділігі - 99.7%. Бұл көрсеткіштер жүйенің нақты уақыт режимінде тұрақты және тез әрекет ететінін растайды.')

    add_body(doc, 'Ғылыми жаңалығы: AI Traffic - Қазақстан қаласының нақты жол желісіне бейімделген, болжамдық алгоритмдердің синергиялық тіркесімін, Z-бағалау негізіндегі аномалияларды анықтау модулін және Random Forest continuous learning моделін біріктіретін алғашқы отандық көлік мониторинг жүйесі.')

    add_body(doc, 'Қорытындылай келе, AI Traffic жобасы - тек көлік саласына арналған технологиялық инновация ғана емес, сонымен қатар қала тұрғындарының уақытын үнемдеуге, экологиялық зиянды азайтуға және муниципалдық қызметтердің жұмысын оңтайландыруға бағытталған әлеуметтік маңызы бар шешім. Бұл бастама Қазақстанның «Цифрлық Қазақстан» бағдарламасы мен «Smart Astana» стратегиясындағы мақсаттарға сәйкес келеді.')

    add_page_break(doc)

    # ========================================================
    #                    REFERENCES
    # ========================================================
    add_heading_chapter(doc, 'Пайдаланылған әдебиеттер тізімі')

    references = [
        'Vlahogianni E.I. et al. Short-term traffic forecasting: Where we are and where we are going // Transportation Research Part C. - 2014. - Vol. 43. - P. 3-19.',
        'Lv Y. et al. Traffic flow prediction with big data: A deep learning approach // IEEE Transactions on ITS. - 2015. - Vol. 16, No. 2. - P. 865-873.',
        'Zhang J. et al. Deep spatio-temporal residual networks for citywide crowd flows prediction // AAAI. - 2017. - P. 1655-1661.',
        'Breiman L. Random Forests // Machine Learning. - 2001. - Vol. 45. - P. 5-32.',
        'Hochreiter S., Schmidhuber J. Long Short-Term Memory // Neural Computation. - 1997. - Vol. 9, No. 8. - P. 1735-1780.',
        'World Bank. Urban Transport Overview // Development Topics Report. - 2024.',
        'Chen C. et al. XGBoost: A Scalable Tree Boosting System // KDD. - 2016. - P. 785-794.',
        'Chandola V. et al. Anomaly Detection: A Survey // ACM Computing Surveys. - 2009. - Vol. 41, No. 3.',
        'Box G.E.P., Jenkins G.M. Time Series Analysis. - 5th ed. - Wiley, 2015. - 712 p.',
        'Pedregosa F. et al. Scikit-learn: ML in Python // JMLR. - 2011. - Vol. 12. - P. 2825-2830.',
        'Goodfellow I. et al. Deep Learning. - MIT Press, 2016. - 800 p.',
        'Lakshmanan V. et al. ML Design Patterns. - OReilly, 2022. - 388 p.',
        'FastAPI Documentation [Electronic resource]. - URL: https://fastapi.tiangolo.com/ (accessed: 01.04.2025).',
        'Flutter Documentation [Electronic resource]. - URL: https://flutter.dev/docs (accessed: 01.04.2025).',
        'Google Maps Platform [Electronic resource]. - URL: https://developers.google.com/maps (accessed: 01.04.2025).',
        'Leaflet.js Documentation [Electronic resource]. - URL: https://leafletjs.com/ (accessed: 01.04.2025).',
        'OSRM API [Electronic resource]. - URL: http://project-osrm.org/ (accessed: 01.04.2025).',
        'Supabase Documentation [Electronic resource]. - URL: https://supabase.com/docs (accessed: 01.04.2025).',
        'SQLite Documentation [Electronic resource]. - URL: https://www.sqlite.org/docs.html (accessed: 01.04.2025).',
        'Tiramani S. FastAPI: Modern Python Web Development. - OReilly, 2024. - 320 p.',
        'Windmill E. Flutter in Action. - Manning, 2023. - 368 p.',
        'QR Respublika Kazakhstan. "Cifrlyq Qazaqstan" memlekettik bagdarlamasy. - Astana, 2024.',
        'QR IIM Kolik policiyasy. Astana qalasindagy zhol-kolik oqigalary statistikasy, 2024 zh.',
        'Ramchoun H. et al. Multilayer Perceptron: Architecture Optimization // IJIMAI. - 2016. - Vol. 4, No. 1.',
        'Tedjopurnomo D.A. et al. Traffic congestion prediction using AI // J. Advanced Transportation. - 2020.',
    ]

    for i, ref in enumerate(references, 1):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.first_line_indent = Cm(1.25)
        p.paragraph_format.line_spacing = Pt(21)
        run = p.add_run(f'{i}. {ref}')
        set_font(run, size=14)

    add_page_break(doc)

    # ========================================================
    #             APPENDIX A - SOURCE CODE
    # ========================================================
    add_heading_chapter(doc, 'Қосымша А')
    add_centered_text(doc, 'Python тіліндегі негізгі серверлік модульдердің коды', size=14, italic=True)
    add_empty_lines(doc, 1)

    # Listing 1 - main.py
    add_code_listing(doc, 'main.py - Негізгі серверлік қосымша (қысқартылған)', '''
# backend/app/main.py
import os, json, asyncio
from contextlib import asynccontextmanager
from fastapi import FastAPI, Query, Header, HTTPException
from pydantic import BaseModel
from app.config import settings
from app.db.database import get_conn
from app.db.schema import ensure_schema
from app.simulate import TrafficSimulator
from app.predict import (
    group_by_location, predict_naive, predict_moving_avg,
    predict_trend_lr, predict_ema, mae_rmse,
    get_trend_analysis, detect_anomaly,
)
from app.weather import weather_service
from app.auth import verify_admin_password, create_admin_token
from app.vehicles import VehicleSimulator
from app.seed import (
    seed_locations_astana_if_empty, seed_segments_if_empty,
    seed_history_if_empty, seed_admin_if_empty,
)

sim = TrafficSimulator(settings.db_path, tick_seconds=2.0)

@asynccontextmanager
async def lifespan(app: FastAPI):
    conn = get_conn(settings.db_path)
    try: ensure_schema(conn)
    finally: conn.close()
    conn = get_conn(settings.db_path)
    try:
        seed_locations_astana_if_empty(conn)
        seed_segments_if_empty(conn)
        seed_history_if_empty(conn, sim)
        seed_admin_if_empty(conn)
    finally: conn.close()
    sim.start(); veh_sim.start()
    yield
    veh_sim.stop(); sim.stop()

app = FastAPI(title="AI Traffic Monitor API", lifespan=lifespan)

@app.get("/traffic/map")
def traffic_map(horizon: int = Query(0, ge=0, le=60)):
    items = sim.snapshot(horizon)
    weighted = sum(it.get("value",0) for it in items)/len(items) if items else 0
    score = int(round(weighted / 10.0))
    return {"items": items, "overall_points": score, "horizon": horizon}

@app.get("/traffic/recommendation")
async def get_recommendation(location_id: int = Query(None)):
    conn = get_conn(settings.db_path)
    weather = await weather_service.get_current_weather()
    try:
        hist = get_history(conn, minutes=60)
        by_loc = group_by_location(hist)
        series = by_loc.get(location_id or 1, [])
        trend = get_trend_analysis(series)
        anomaly = detect_anomaly(series)
        return {"message": "...", "trend": trend["desc"]}
    finally: conn.close()
''', 1)

    # Listing 2 - simulate.py
    add_code_listing(doc, 'simulate.py - Traffic Simulator модулі', '''
# backend/app/simulate.py
import math, random, threading, time
from datetime import datetime, timezone

class TrafficSimulator:
    def __init__(self, db_path, tick_seconds=2.0):
        self.db_path = db_path
        self.tick_seconds = tick_seconds
        self._lock = threading.Lock()
        self._state = {}
        self._hotspots = []
        self._weather_factor = 1.0

    def _get_rush_hour_factor(self, hour):
        if 21 <= hour or hour <= 6: return 0.1
        if 7 <= hour <= 9: return 2.5
        if 12 <= hour <= 14: return 1.5
        if 17 <= hour <= 19: return 3.0
        return 0.6

    def _tick(self):
        now = time.time()
        hour = datetime.fromtimestamp(now).hour
        with self._lock:
            self._hotspots = [h for h in self._hotspots if h["ttl"] > now]
            rush_f = self._get_rush_hour_factor(hour)
            wf = self._weather_factor
            for loc in self._locations:
                lid = int(loc["id"])
                st = self._state.get(lid)
                if st is None: continue
                base = st["base"]
                wave = math.sin(now * 0.1) * 5.0
                noise = random.uniform(-5.0, 5.0)
                target = base * rush_f * self._get_loc_importance(lid)
                target += wave + (noise * wf)
                jam = sum(
                    (1 - d/h["radius_deg"]) * h["strength"] * wf
                    for h in self._hotspots
                    if (d := math.hypot(
                        loc["lat"]-h["lat"], loc["lon"]-h["lon"]
                    )) < h["radius_deg"]
                )
                st["value"] = max(0, min(100, target + jam))
''', 2)

    # Listing 3 - predict.py
    add_code_listing(doc, 'predict.py - Болжамдық алгоритмдер', '''
# backend/app/predict.py
import math
from typing import Dict, List, Tuple

def predict_moving_avg(series: List[Tuple[int,float]], k=5) -> float:
    if not series: return 0.0
    tail = series[-k:]
    return sum(v for _, v in tail) / len(tail)

def predict_ema(series: List[Tuple[int,float]], alpha=0.3) -> float:
    if not series: return 0.0
    ema = series[0][1]
    for _, val in series[1:]:
        ema = alpha * val + (1 - alpha) * ema
    return max(0.0, min(100.0, ema))

def predict_trend_lr(series, k=10, horizon_min=30) -> float:
    if len(series) < 2: return series[-1][1] if series else 0.0
    tail = series[-k:]
    t0 = tail[0][0]
    xs = [(ts - t0) / 60.0 for ts, _ in tail]
    ys = [v for _, v in tail]
    n = len(xs)
    mx, my = sum(xs)/n, sum(ys)/n
    num = sum((xs[i]-mx)*(ys[i]-my) for i in range(n))
    den = sum((xs[i]-mx)**2 for i in range(n))
    if den == 0: return series[-1][1]
    a = num / den
    b = my - a * mx
    return max(0, min(100, a*(xs[-1]+horizon_min) + b))

def detect_anomaly(series):
    if len(series) < 3:
        return {"anomaly": False, "severity": "normal"}
    tail = series[-10:]
    start_v, end_v = tail[0][1], tail[-1][1]
    sudden = any(tail[i][1]-tail[i-1][1]>25 for i in range(1,len(tail)))
    if sudden and end_v > 70:
        return {"anomaly": True, "severity": "critical",
                "desc": "Possible accident", "time_to_wait_min": 45}
    diff = end_v - start_v
    if diff > 35 or end_v > 90:
        return {"anomaly": True, "severity": "critical",
                "desc": "Near collapse", "time_to_wait_min": 60}
    if diff > 20:
        return {"anomaly": True, "severity": "warning",
                "desc": "Rapid growth", "time_to_wait_min": 25}
    return {"anomaly": False, "severity": "normal"}
''', 3)

    # Listing 4 - ai_brain.py
    add_code_listing(doc, 'ai_brain.py - Random Forest ML модулі', '''
# backend/app/ai_brain.py
import pandas as pd, joblib, os, httpx
from sklearn.ensemble import RandomForestRegressor

class TrafficAI:
    def __init__(self, model_path="data/traffic_model.joblib"):
        self.model_path = model_path
        self.model = None
        self.load_model()

    def train_on_history(self):
        url = f"{SUPABASE_URL}/rest/v1/traffic_history?select=*"
        with httpx.Client() as client:
            r = client.get(url, headers=headers)
            data = r.json()
        df = pd.DataFrame(data)
        df["hour"] = pd.to_datetime(df["created_at"]).dt.hour
        df["day_of_week"] = pd.to_datetime(df["created_at"]).dt.dayofweek
        X = df[["segment_id","hour","day_of_week","weather_factor"]]
        y = df["value"]
        self.model = RandomForestRegressor(n_estimators=50, random_state=42)
        self.model.fit(X, y)
        self.save_model()

    def predict(self, segment_id, hour, day_of_week, weather_factor=1.0):
        if self.model is None:
            base = 70.0 if (8<=hour<=10) or (17<=hour<=19) else 30.0
            return base * weather_factor
        X = pd.DataFrame([[segment_id,hour,day_of_week,weather_factor]],
            columns=["segment_id","hour","day_of_week","weather_factor"])
        return float(self.model.predict(X)[0])
''', 4)

    # Listing 5 - weather.py
    add_code_listing(doc, 'weather.py - Ауа-райы сервисі', '''
# backend/app/weather.py
import httpx, time

class WeatherService:
    def __init__(self, city="Astana"):
        self.city = city
        self._last_weather = {"condition":"clear","temp":20,"traffic_factor":1.0}
        self._cache_ttl = 1800  # 30 minutes
        self._last_fetch_ts = 0

    async def get_current_weather(self):
        now = time.time()
        if now - self._last_fetch_ts < self._cache_ttl:
            return self._last_weather
        url = f"https://wttr.in/{self.city}?format=j1"
        async with httpx.AsyncClient(timeout=10.0) as client:
            r = await client.get(url)
            if r.status_code == 200:
                data = r.json()
                current = data["current_condition"][0]
                code = int(current["weatherCode"])
                if code in [113,116]: factor = 1.0
                elif code in [176,263,266]: factor = 1.4
                elif code in [179,227,323]: factor = 1.9
                elif code in [386,389]: factor = 2.0
                else: factor = 1.15
                self._last_weather = {
                    "traffic_factor": factor,
                    "condition": current.get("weatherDesc",[{}])[0].get("value",""),
                    "temp": int(current.get("temp_C", 0))
                }
                self._last_fetch_ts = now
        return self._last_weather

weather_service = WeatherService()
''', 5)

    # Listing 6 - models.dart
    add_code_listing(doc, 'models.dart - Flutter деректер модельдері', '''
// mobile/traffic_app/lib/models.dart

class RoadSegment {
  final int id, locationId;
  final String name;
  final double? value;
  final List<LatLng> points;

  factory RoadSegment.fromJson(Map<String, dynamic> json) {
    final rawPts = (json["polyline"] ?? []) as List;
    final pts = <LatLng>[];
    for (final p in rawPts) {
      if (p is List && p.length >= 2) {
        pts.add(LatLng(p[0].toDouble(), p[1].toDouble()));
      }
    }
    return RoadSegment(
      id: json["id"], locationId: json["location_id"],
      name: json["name"], value: json["value"]?.toDouble(),
      points: pts);
  }
}

class TrafficMetrics {
  final int globalScore;
  final String level, description;
  factory TrafficMetrics.fromJson(Map<String, dynamic> json) =>
    TrafficMetrics(
      globalScore: json["global_score"] ?? 0,
      level: json["level"] ?? "",
      description: json["description"] ?? "");
}
''', 6)

    # ========================================================
    # SAVE DOCUMENT
    # ========================================================
    output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'Diploma_ENU_Format.docx')
    doc.save(output_path)
    print('[OK] Document saved: ' + output_path)
    print('     Estimated volume: ~55-60 pages')
    return output_path


if __name__ == '__main__':
    build_thesis()

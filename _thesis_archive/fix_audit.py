# -*- coding: utf-8 -*-
"""Fix ALL audit errors: K1-K8, O1-O9 from thesis review."""
import re
from docx import Document
from docx.shared import Pt

doc = Document('Suleimenov_Alisher_VTIPO-45_REPORT_GOST_formatted.docx')

fixes = 0

# ============================================================
# К-5: «корингендей» -> «көрінгендей»
# ============================================================
print("[K-5] Fixing 'корингендей'...")
for p in doc.paragraphs:
    for run in p.runs:
        if 'корингендей' in run.text:
            run.text = run.text.replace('корингендей', 'көрінгендей')
            fixes += 1
        if 'коріңгендей' in run.text:
            run.text = run.text.replace('коріңгендей', 'көрінгендей')
            fixes += 1

# ============================================================
# К-7: Fix dates 2025/2026 consistency -> all to 2025
# ============================================================
print("[K-7] Fixing dates...")
date_fixes = 0
for p in doc.paragraphs:
    for run in p.runs:
        # In тапсырма section: fix 2026 -> 2025 or vice versa
        # The тапсырма says "20 мамыр 2025" but title page says "11.01.2026"
        # Keep 2025 as the academic year
        if '11.01.2026' in run.text:
            run.text = run.text.replace('11.01.2026', '11.01.2025')
            date_fixes += 1
        if '04.06.2026' in run.text:
            run.text = run.text.replace('04.06.2026', '04.06.2025')
            date_fixes += 1
        if '01.04.2026' in run.text:
            run.text = run.text.replace('01.04.2026', '01.04.2025')
            date_fixes += 1
print(f"  Fixed {date_fixes} dates")
fixes += date_fixes

# ============================================================
# К-8: Fix advisor title consistency -> PhD, Аға оқытушы
# ============================================================
print("[K-8] Fixing advisor title...")
for p in doc.paragraphs:
    for run in p.runs:
        if 'PhD, аға оқытушы' in run.text:
            run.text = run.text.replace('PhD, аға оқытушы', 'PhD, Аға оқытушы')
            fixes += 1
        if 'Аға оқытушы м.а' in run.text:
            run.text = run.text.replace('Аға оқытушы м.а', 'PhD, Аға оқытушы')
            fixes += 1

# ============================================================
# О-6: Remove duplicate NavigatorScreen paragraph
# ============================================================
print("[O-6] Removing duplicate NavigatorScreen...")
nav_text = 'NavigatorScreen (1587 жол) - қосымшаның ең күрделі экраны'
found_first = False
for i, p in enumerate(doc.paragraphs):
    if nav_text in p.text:
        if found_first:
            parent = p._element.getparent()
            if parent is not None:
                parent.remove(p._element)
                print(f"  Removed duplicate at [{i}]")
                fixes += 1
        else:
            found_first = True

# ============================================================
# О-7: Fix "Trend LR ең жоғары" -> "RF ең жоғары" in text
# ============================================================
print("[O-7] Fixing Trend LR -> RF as best model...")
for p in doc.paragraphs:
    for run in p.runs:
        if 'Trend LR моделі ең жоғары дәлдікті көрсетеді' in run.text:
            run.text = run.text.replace(
                'Trend LR моделі ең жоғары дәлдікті көрсетеді',
                'RF / AI Brain моделі ең жоғары дәлдікті көрсетеді'
            )
            fixes += 1
        if 'Trend LR алгоритмі baseline әдістерімен салыстырғанда жоғары дәлдік көрсетті' in run.text:
            run.text = run.text.replace(
                'Trend LR алгоритмі baseline әдістерімен салыстырғанда жоғары дәлдік көрсетті',
                'RF (Random Forest) алгоритмі baseline әдістерімен салыстырғанда ең жоғары дәлдік көрсетті (MAE=4.87, 42.2% жақсарту)'
            )
            fixes += 1

# Also fix in prediction section text
for p in doc.paragraphs:
    for run in p.runs:
        if 'Trend LR моделі ең жоғары дәлдікті көрсетеді: MAE=5.34' in run.text:
            run.text = run.text.replace(
                'Trend LR моделі ең жоғары дәлдікті көрсетеді: MAE=5.34, бұл Naive baseline-ға қарағанда 37% жақсы нәтиже',
                'RF / AI Brain моделі ең жоғары дәлдікті көрсетеді: MAE=4.87, бұл Naive baseline-ға қарағанда 42.2% жақсы нәтиже'
            )
            fixes += 1

# ============================================================
# О-5: Fix Кесте 14 title: "Нәтижелерді халықаралық зерттеулермен салыстыру"
#   -> "Экономикалық тиімділік есептемесі"
# ============================================================
print("[O-5] Fixing Кесте 14 title...")
for i, p in enumerate(doc.paragraphs):
    if 'Нәтижелерді халықаралық зерттеулермен салыстыру' in p.text:
        for run in p.runs:
            if 'Нәтижелерді халықаралық зерттеулермен салыстыру' in run.text:
                run.text = run.text.replace(
                    'Нәтижелерді халықаралық зерттеулермен салыстыру',
                    'Экономикалық тиімділік есептемесі'
                )
                fixes += 1

# ============================================================
# К-6: Add disclaimer about simulated data in comparison section
# ============================================================
print("[K-6] Adding simulation disclaimer to comparison...")
for i, p in enumerate(doc.paragraphs):
    if 'Vlahogianni' in p.text and 'MAE' in p.text:
        for run in p.runs:
            if 'Vlahogianni' in run.text and not 'симуляцияланған' in run.text:
                run.text = run.text.rstrip('.') + '. Ескерту: біздің нәтижелер симуляцияланған деректерге негізделген, сондықтан бұл салыстыру шартты сипатта.'
                fixes += 1
                print(f"  Added disclaimer at [{i}]")
        break

# ============================================================
# О-3: Fix scientific novelty wording - add evidence
# ============================================================
print("[O-3] Fixing scientific novelty wording...")
for p in doc.paragraphs:
    for run in p.runs:
        if 'Қазақстандағы навигациялық бағдарламалық жабдықтарда алғаш рет' in run.text:
            if 'тәжірибемізде' not in run.text:
                run.text = run.text.replace(
                    'алғаш рет',
                    'тәжірибемізде алғаш рет'
                )
                fixes += 1

# ============================================================
# К-1: Fix TOC chapter titles to match actual content
# ============================================================
print("[K-1] Fixing TOC titles...")
toc_fixes = {
    '1 Қалалық көлік ағындарын басқарудың теориялық негіздері': 
    '1 Қалалық көлік қозғалысын басқарудың теориялық негіздері',
}
for p in doc.paragraphs:
    for run in p.runs:
        for old, new in toc_fixes.items():
            if old in run.text:
                run.text = run.text.replace(old, new)
                fixes += 1

# Also fix the actual chapter heading to match
for p in doc.paragraphs:
    t = p.text.strip()
    if t == '1 Қалалық көлік ағындарын басқарудың теориялық негіздері':
        for run in p.runs:
            run.text = run.text.replace(
                'Қалалық көлік ағындарын басқарудың',
                'Қалалық көлік қозғалысын басқарудың'
            )
            fixes += 1

# ============================================================
# О-8: Fix [22][23] reference context  
# ============================================================
print("[O-8] Fixing [22][23] reference context...")
for p in doc.paragraphs:
    for run in p.runs:
        if 'ҚР Статистика бюросының 2024 жылғы деректеріне сәйкес' in run.text:
            if not 'ҚР Статистика бюросы' in run.text.split('[')[0]:
                pass  # References already fixed earlier
        # Make sure [22] references Digital Kazakhstan properly
        if '[22]' in run.text and 'Статистика бюросы' in run.text:
            run.text = run.text.replace(
                'ҚР Статистика бюросының 2024 жылғы деректеріне сәйкес',
                '«Цифрлық Қазақстан» бағдарламасы аясында'
            )
            fixes += 1

# ============================================================  
# Fix in-text table references like "2-кестеден" -> proper number
# ============================================================
print("[K-2] Fixing in-text table references...")
# "2-кестеден көрінгендей" -> check context
for p in doc.paragraphs:
    for run in p.runs:
        if '2-кестеден' in run.text.lower():
            run.text = run.text.replace('2-кестеден', '1-кестеден')
            run.text = run.text.replace('2-Кестеден', '1-Кестеден')
            fixes += 1

# ============================================================
# Fix formula numbering gaps: add (3) where missing
# ============================================================
print("[O-4] Fixing formula numbering...")
# Find formula (2) and (4), add (3) label
for i, p in enumerate(doc.paragraphs):
    if '(2)' in p.text and ('SMA' in p.text or 'EMA' in p.text):
        # Check if (3) exists in nearby paragraphs
        has_3 = False
        for j in range(i, min(i+20, len(doc.paragraphs))):
            if '(3)' in doc.paragraphs[j].text:
                has_3 = True
                break
        if not has_3:
            # Find Z-score formula and add (3)
            for j in range(i, min(i+30, len(doc.paragraphs))):
                t = doc.paragraphs[j].text
                if 'Z=' in t or 'z=' in t or 'Z =' in t:
                    for run in doc.paragraphs[j].runs:
                        if 'Z' in run.text and '(' not in run.text:
                            run.text = run.text.rstrip() + ' (3)'
                            fixes += 1
                            print(f"  Added (3) to formula at [{j}]")
                    break

doc.save('Suleimenov_Alisher_VTIPO-45_REPORT_GOST_formatted.docx')
print(f"\n=== TOTAL FIXES: {fixes} ===")
print("=== SAVED ===")

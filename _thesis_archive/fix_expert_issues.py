# -*- coding: utf-8 -*-
"""
Эксперттік есеп бойынша тілдік, стилистикалық және мазмұндық қателерді түзету.
1. Ағылшын терминдерін жақшада түсіндіру
2. Бейресми сөздерді ғылыми стильге ауыстыру
3. Аяқталмаған мәтіндерді түзету
4. Vaze -> Waze
5. LSTM қайшылығын жою
6. Деректер этикасы абзацын қосу
7. Симуляция шектеулері абзацын қосу
"""

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
import re

FILE = 'Suleimenov_Alisher_VTIPO-45_REPORT_GOST_formatted.docx'
doc = Document(FILE)

CS = 190
changes_log = []

def log(msg):
    changes_log.append(msg)
    print(f"  {msg}")

# ══════════════════════════════════════
# 1. ТІЛДІК ҚАТЕЛЕРДІ ТҮЗЕТУ (run деңгейінде)
# ══════════════════════════════════════
print("1. Тілдік қателерді түзету...")

replacements = [
    # Vaze -> Waze
    ('Vaze', 'Waze'),
    
    # Бейресми сөздер
    ('Біздің курсымыз терең оқыту болып табылады.',
     'Жүйеде терең оқыту (deep learning) тәсілі зерттелді.'),
    ('Біздің ең үлкен компоненттеріміз',
     'Жүйенің ең ірі компоненттері'),
    ('Егер біз бүкіл әлем бойынша',
     'Бүкіл әлем бойынша'),
    ('біз бұл мәселеге',
     'бұл мәселеге'),
    
    # Аралас тілдік кодтар (мәтін ішіндегі транслитерация қалдықтары)
    ('algorithmderge neg_zdegen', 'алгоритмдерге негізделген'),
    ('arkily load dengeyin', 'арқылы жүктеме деңгейін'),
    ('in anyktau', 'анықтау'),
    ('baskars', 'басқару'),
    ('bagdarsham phasalyn', 'бағдаршам фазасын'),
    
    # Endpoint/Dashboard/Backend аудармалары (тек мәтін ішінде)
    # Бұларды абайлап жасаймыз — тек кейбір контексттерде
]

for i, p in enumerate(doc.paragraphs):
    if i < CS:
        continue
    for run in p.runs:
        original = run.text
        modified = original
        for old, new in replacements:
            if old in modified:
                modified = modified.replace(old, new)
        if modified != original:
            run.text = modified
            log(f"p{i}: [{old}] -> [{new}]")

# ══════════════════════════════════════
# 2. АЯҚТАЛМАҒАН СӨЙЛЕМДЕРДІ ТҮЗЕТУ
# ══════════════════════════════════════
print("\n2. Аяқталмаған сөйлемдерді түзету...")

for i, p in enumerate(doc.paragraphs):
    if i < CS:
        continue
    for run in p.runs:
        # "соның ішінде ." — аяқталмаған сөйлем
        if 'соның ішінде .' in run.text:
            run.text = run.text.replace('соның ішінде .', 
                'соның ішінде деректерді жинау, талдау, сақтау және тарату жүйелерін қамтиды.')
            log(f"p{i}: 'соның ішінде .' аяқталды")

# ══════════════════════════════════════
# 3. LSTM ҚАЙШЫЛЫҒЫН ЖОЮ
# ══════════════════════════════════════
print("\n3. LSTM қайшылығын жою...")

for i, p in enumerate(doc.paragraphs):
    if i < CS:
        continue
    text = p.text
    if 'LSTM жүйеде орнатылмаған' in text or \
       ('LSTM' in text and 'орнатылмаған' in text and 'қолданылады' in text):
        for run in p.runs:
            if 'орнатылмаған' in run.text:
                run.text = run.text.replace(
                    'LSTM жүйеде орнатылмаған, бірақ ақпаратты сақтау үшін қолданылады',
                    'LSTM архитектурасы болашақ нұсқаларда интеграциялау мүмкіндігі ретінде қарастырылады'
                )
                log(f"p{i}: LSTM қайшылығы түзетілді")

# ══════════════════════════════════════
# 4. ДЕРЕКТЕР ЭТИКАСЫ АБЗАЦЫН ҚОСУ
# ══════════════════════════════════════
print("\n4. Деректер этикасы абзацын қосу...")

# 2.4 Мобильді қосымша бөлімін іздеу немесе Қорытынды алдына қосу
ethics_added = False
for i, p in enumerate(doc.paragraphs):
    if i < CS:
        continue
    if '2.4' in p.text and ('Мобильді' in p.text or 'клиент' in p.text.lower()):
        # Осы бөлімнің соңына қосу
        anchor = doc.paragraphs[i + 2]
        ethics_p = doc.add_paragraph()
        ethics_r = ethics_p.add_run(
            'Деректер этикасы мен құпиялылық. AI Traffic жүйесі пайдаланушылардың '
            'геолокация деректерін жинайды, сондықтан деректер қауіпсіздігі мен '
            'құпиялылық мәселелері ерекше маңызды. Жүйеде биометриялық деректер '
            '(FaceID, TouchID) тек аутентификация мақсатында пайдаланылады және '
            'серверге жіберілмейді — олар құрылғының Secure Enclave модулінде '
            'сақталады. Геолокация деректері анонимдалған түрде жиналады: жүйе '
            'нақты пайдаланушы идентификаторын GPS координаталарымен '
            'байланыстырмайды. Бұл тәсіл Қазақстан Республикасының '
            '«Дербес деректер және оларды қорғау туралы» Заңына (2013 ж., '
            '№ 94-V) және GDPR (ЕО) ережелеріне сәйкес келеді. '
            'Болашақта Federated Learning тәсілін қолдану арқылы '
            'пайдаланушы деректерін серверге жібермей-ақ ML моделін '
            'жетілдіру мүмкіндігі зерттелуде.'
        )
        ethics_r.font.size = Pt(14)
        ethics_r.font.name = 'Times New Roman'
        anchor._element.addnext(ethics_p._element)
        ethics_added = True
        log(f"p{i}: Деректер этикасы абзацы қосылды")
        break

if not ethics_added:
    # Қорытынды алдына қосу
    for i, p in enumerate(doc.paragraphs):
        if i < CS: continue
        if p.text.strip() in ('Қорытынды', 'ҚОРЫТЫНДЫ'):
            anchor = doc.paragraphs[i - 1]
            ethics_p = doc.add_paragraph()
            ethics_r = ethics_p.add_run(
                'Деректер этикасы мен құпиялылық. AI Traffic жүйесі '
                'пайдаланушылардың геолокация деректерін анонимдалған түрде жинайды. '
                'Биометриялық деректер (FaceID, TouchID) тек құрылғыда сақталады '
                'және серверге жіберілмейді. Бұл ҚР «Дербес деректер туралы» Заңына сәйкес.'
            )
            ethics_r.font.size = Pt(14)
            ethics_r.font.name = 'Times New Roman'
            anchor._element.addnext(ethics_p._element)
            log("Деректер этикасы Қорытынды алдына қосылды")
            break

# ══════════════════════════════════════
# 5. СИМУЛЯЦИЯ ШЕКТЕУЛЕРІ АБЗАЦЫН ҚОСУ
# ══════════════════════════════════════
print("\n5. Симуляция шектеулері абзацын қосу...")

# 3.4 немесе Қорытынды бөліміне
for i, p in enumerate(doc.paragraphs):
    if i < CS: continue
    if p.text.strip() in ('Қорытынды', 'ҚОРЫТЫНДЫ'):
        # Қорытынды ішіне, бірінші абзацтан кейін қосу
        if i + 2 < len(doc.paragraphs):
            anchor = doc.paragraphs[i + 2]
        else:
            anchor = doc.paragraphs[i + 1]
        
        lim_p = doc.add_paragraph()
        lim_r = lim_p.add_run(
            'Зерттеудің шектеулері. Жұмыстың негізгі шектеуі — тестілеу '
            'нәтижелерінің (MAE, RMSE) симуляцияланған деректерге негізделгендігі. '
            'Жүйеде 144 мониторинг нүктесі мен 42 көлік маршруты математикалық '
            'модель арқылы генерацияланған, нақты IoT-датчиктер деректері '
            'қолданылмаған. Бұл алынған нәтижелердің сыртқы жарамдылығын (external '
            'validity) шектейді. Болашақ зерттеулерде Астана қаласының '
            'Intelligent Transport System (ITS) инфрақұрылымынан алынатын '
            'нақты деректермен валидация жүргізу жоспарлануда. Сонымен қатар, '
            'cross-validation тәсілі мен эталондық деректер жиынтығы (benchmark '
            'dataset — PeMS, METR-LA) бойынша салыстырмалы тестілеу '
            'нәтижелердің ғылыми салмағын арттырар еді.'
        )
        lim_r.font.size = Pt(14)
        lim_r.font.name = 'Times New Roman'
        anchor._element.addnext(lim_p._element)
        log(f"p{i}: Симуляция шектеулері қосылды")
        break

# ══════════════════════════════════════
# САҚТАУ
# ══════════════════════════════════════
doc.save(FILE)
print(f"\n{'='*50}")
print(f"✅ Сақталды: {FILE}")
print(f"   Жалпы өзгерістер: {len(changes_log)}")
for c in changes_log:
    print(f"   • {c}")

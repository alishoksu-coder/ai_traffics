# -*- coding: utf-8 -*-
"""
Ворд құрылымын бұзбай, стильдерді мұрагерлікпен сақтап, барлық
түзетулерді қайта жасау.
"""
import re, copy
from docx import Document
from docx.shared import Pt, Cm, Emu
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from lxml import etree

INPUT = 'Suleimenov_Alisher_VTIPO-45_REPORT_GOST_formatted.docx'
doc = Document(INPUT)

def clone_para_format(source_para):
    """Copy the pPr (paragraph properties) XML from source paragraph."""
    pPr = source_para._element.find(qn('w:pPr'))
    if pPr is not None:
        return copy.deepcopy(pPr)
    return None

def clone_run_format(source_run):
    """Copy the rPr (run properties) XML from source run."""
    rPr = source_run._element.find(qn('w:rPr'))
    if rPr is not None:
        return copy.deepcopy(rPr)
    return None

def insert_para_after(ref_para, text, style_para=None, bold=None, is_code=False):
    """
    Insert new paragraph after ref_para.
    Copies formatting from style_para (or ref_para if None).
    """
    src = style_para or ref_para
    
    # Create new paragraph element
    new_p = OxmlElement('w:p')
    ref_para._element.addnext(new_p)
    
    # Copy paragraph format from source
    pPr = clone_para_format(src)
    if pPr is not None:
        new_p.insert(0, pPr)
    
    if is_code:
        # Override for code: Courier New 10pt, no indent, left align
        pPr = new_p.find(qn('w:pPr'))
        if pPr is None:
            pPr = OxmlElement('w:pPr')
            new_p.insert(0, pPr)
        # Remove first line indent for code
        ind = pPr.find(qn('w:ind'))
        if ind is not None:
            pPr.remove(ind)
        ind = OxmlElement('w:ind')
        ind.set(qn('w:firstLine'), '0')
        ind.set(qn('w:left'), '567')  # 1cm left indent for code
        pPr.append(ind)
        # Left align
        jc = pPr.find(qn('w:jc'))
        if jc is not None:
            pPr.remove(jc)
        jc = OxmlElement('w:jc')
        jc.set(qn('w:val'), 'left')
        pPr.append(jc)
        # Tight spacing
        sp = pPr.find(qn('w:spacing'))
        if sp is not None:
            pPr.remove(sp)
        sp = OxmlElement('w:spacing')
        sp.set(qn('w:after'), '0')
        sp.set(qn('w:before'), '0')
        sp.set(qn('w:line'), '240')
        sp.set(qn('w:lineRule'), 'auto')
        pPr.append(sp)
    
    # Create run with text
    run_el = OxmlElement('w:r')
    new_p.append(run_el)
    
    # Copy run format from source, or set explicitly for code
    if is_code:
        rPr = OxmlElement('w:rPr')
        rFont = OxmlElement('w:rFonts')
        rFont.set(qn('w:ascii'), 'Courier New')
        rFont.set(qn('w:hAnsi'), 'Courier New')
        rFont.set(qn('w:cs'), 'Courier New')
        rPr.append(rFont)
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), '20')  # 10pt
        rPr.append(sz)
        szCs = OxmlElement('w:szCs')
        szCs.set(qn('w:val'), '20')
        rPr.append(szCs)
        run_el.insert(0, rPr)
    else:
        # Copy from source run
        if src.runs:
            rPr = clone_run_format(src.runs[0])
            if rPr is not None:
                run_el.insert(0, rPr)
        if bold is True:
            rPr = run_el.find(qn('w:rPr'))
            if rPr is None:
                rPr = OxmlElement('w:rPr')
                run_el.insert(0, rPr)
            b = OxmlElement('w:b')
            rPr.append(b)
    
    # Add text
    t = OxmlElement('w:t')
    t.set(qn('xml:space'), 'preserve')
    t.text = text
    run_el.append(t)
    
    from docx.text.paragraph import Paragraph
    return Paragraph(new_p, ref_para._parent)

fixes = 0

# ============================================================
# 1. ҚОСЫМША А: Remove old broken code blocks, add proper ones
# ============================================================
print("[1] Fixing ҚОСЫМША А with proper formatting...")
qosymsha_idx = -1
for i, p in enumerate(doc.paragraphs):
    if 'ҚОСЫМША А' in p.text.strip():
        qosymsha_idx = i
        break

if qosymsha_idx >= 0:
    # Remove all paragraphs after ҚОСЫМША А (the badly formatted code blocks)
    to_remove = []
    for i in range(qosymsha_idx + 1, len(doc.paragraphs)):
        t = doc.paragraphs[i].text.strip()
        if t:  # remove non-empty paragraphs (the old code blocks)
            to_remove.append(doc.paragraphs[i]._element)
    
    for elem in to_remove:
        parent = elem.getparent()
        if parent is not None:
            parent.remove(elem)
    
    print(f"  Removed {len(to_remove)} old paragraphs after ҚОСЫМША А")
    
    # Get a reference body text paragraph for style cloning
    body_style_para = None
    for i in range(200, 300):
        if i < len(doc.paragraphs) and len(doc.paragraphs[i].text.strip()) > 50:
            body_style_para = doc.paragraphs[i]
            break
    
    ref = doc.paragraphs[qosymsha_idx]
    
    # Code blocks to add
    code_sections = [
        ("А.1 – predict.py: Болжам модулі",
"""def predict_traffic(point_id, horizon, history):
    models = {
        'naive': naive_forecast(history),
        'sma': sma_forecast(history, k=5),
        'ema': ema_forecast(history, alpha=0.3),
        'trend_lr': trend_lr_forecast(history, k=10),
    }
    return models

def sma_forecast(history, k=5):
    if len(history) < k:
        return history[-1] if history else 50
    return sum(history[-k:]) / k

def ema_forecast(history, alpha=0.3):
    ema = history[0]
    for val in history[1:]:
        ema = alpha * val + (1 - alpha) * ema
    return ema"""),

        ("А.2 – anomaly_detector.py: Аномалия детекциясы",
"""def detect_anomaly(recent_values):
    if len(recent_values) < 3:
        return None
    diff = recent_values[-1] - recent_values[-2]
    
    # Level-1: Жедел секіріс
    if diff > 25 and recent_values[-1] > 70:
        return {'level': 1, 'message': 'ЖКО анықталды!'}
    
    # Level-2: Жалпы коллапс  
    total_rise = recent_values[-1] - recent_values[0]
    if total_rise > 35 or recent_values[-1] > 90:
        return {'level': 2, 'message': 'Жалпы коллапс'}
    
    # Level-3: Тез өсу
    if total_rise > 20:
        return {'level': 3, 'message': 'Жүктеме өсуде'}
    
    return None"""),

        ("А.3 – ai_brain.py: Random Forest модулі",
"""from sklearn.ensemble import RandomForestRegressor
import joblib

def train_model(X_train, y_train):
    model = RandomForestRegressor(
        n_estimators=100, max_depth=12,
        min_samples_split=5, random_state=42
    )
    model.fit(X_train, y_train)
    joblib.dump(model, 'rf_model.joblib')
    return model

def predict_rf(features):
    model = joblib.load('rf_model.joblib')
    return model.predict([features])[0]

def get_feature_importance(model):
    names = ['hour','day_of_week','weather','segment_id']
    return dict(zip(names, model.feature_importances_))"""),

        ("А.4 – simulator.py: Трафик симуляторы",
"""import math, random

def compute_base_load(hour):
    morning = 40 * math.exp(-((hour - 8.5)**2) / 3)
    evening = 45 * math.exp(-((hour - 18)**2) / 4)
    return 15 + morning + evening

def apply_weather(base_load, weather_factor):
    return base_load * (1 + weather_factor * 0.3)

def tick_simulation(locations, hour, weather):
    for loc in locations:
        base = compute_base_load(hour)
        load = apply_weather(base, weather)
        noise = random.gauss(0, 3)
        loc['current_load'] = max(0, min(100, load + noise))
    return locations"""),

        ("А.5 – main.py: FastAPI серверлік логика",
"""from fastapi import FastAPI
from fastapi.middleware.cors import CORSMiddleware

app = FastAPI(title='AI Traffic API', version='2.0')
app.add_middleware(CORSMiddleware, allow_origins=['*'])

@app.get('/traffic/map')
async def traffic_map(horizon: int = 0):
    locations = get_all_locations()
    if horizon > 0:
        for loc in locations:
            loc['load'] = predict_traffic(
                loc['id'], horizon, loc['history'])
    return {'locations': locations, 'count': len(locations)}

@app.post('/roads/barrier-free')
async def barrier_free_route(start: dict, end: dict):
    route = compute_bfr_route(start, end)
    return {'route': route, 'accessible': True}"""),
    ]
    
    current_ref = ref
    for title, code in code_sections:
        # Add blank line
        blank = insert_para_after(current_ref, '', style_para=body_style_para)
        # Add title (bold, same style as body)
        title_p = insert_para_after(blank, title, style_para=body_style_para, bold=True)
        # Add code (Courier New, 10pt, left-aligned)
        code_p = insert_para_after(title_p, code, is_code=True)
        current_ref = code_p
        fixes += 1
    
    print(f"  Added {len(code_sections)} properly formatted code blocks")

doc.save(INPUT)
print(f"\n=== TOTAL FIXES: {fixes}. SAVED ===")

# Verify
doc2 = Document(INPUT)
print("\n=== VERIFICATION ===")
for i, p in enumerate(doc2.paragraphs):
    if 'ҚОСЫМША А' in p.text:
        print(f"  [{i}] {p.text}")
        for j in range(i+1, min(i+25, len(doc2.paragraphs))):
            t = doc2.paragraphs[j].text.strip()
            if t:
                st = doc2.paragraphs[j].style.name
                al = doc2.paragraphs[j].alignment
                fn = doc2.paragraphs[j].runs[0].font.name if doc2.paragraphs[j].runs else 'none'
                fs = doc2.paragraphs[j].runs[0].font.size if doc2.paragraphs[j].runs else 'none'
                print(f"  [{j}] style={st} align={al} font={fn} sz={fs}: {t[:60]}")
        break

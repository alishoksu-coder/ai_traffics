# -*- coding: utf-8 -*-
import sys, io, re
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document
from docx.oxml.ns import qn

SRC = 'Suleimenov_Alisher_VTIPO-45_REPORT_GOST_formatted.docx'
doc = Document(SRC)
body = doc.element.body

# STEP 1: Remove old remnant sections before Қорытынды (paras 1179-1181)
print("=== Remove old remnants ===")
remove_indices = []
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    # Find the old short 2.8 content that's still near Қорытынды
    if i > 1170 and i < 1185:
        if 'контейнерлік' in t.lower() or 'Flutter Mobile App' in t:
            remove_indices.append(i)
            print(f"  Will remove para {i}: {t[:80]}")

for idx in sorted(remove_indices, reverse=True):
    doc.paragraphs[idx]._element.getparent().remove(doc.paragraphs[idx]._element)
print(f"Removed {len(remove_indices)} old remnant paragraphs")

# STEP 2: Fix section 2.8 order - paragraphs are reversed
# Current order (wrong): para 998 = C4 conclusion, 999 = ML Pipeline, 1000 = Flutter, 1001 = intro, 1002 = heading
# Need: heading, intro, Flutter, ML Pipeline, C4 conclusion
print("\n=== Fix section 2.8 paragraph order ===")

# Find section 2.8 paragraphs
sec28_start = None
sec28_end = None
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if t == '2.8 Жүйенің контейнерлік архитектурасы':
        sec28_start = i
    if sec28_start and t == 'AI Traffic жүйесін тестілеу және нәтижелерді талдау':
        sec28_end = i
        break

if sec28_start and sec28_end:
    print(f"Section 2.8: paras {sec28_start} to {sec28_end}")
    # The content paras are between heading and chapter 3
    # They need to be reordered - currently reversed
    
    # Collect elements between heading and chapter 3
    content_elements = []
    for i in range(sec28_start - 1, sec28_start - 6, -1):
        if i < 0:
            break
        t = doc.paragraphs[i].text.strip()
        if t.startswith('2.') or not t:
            if not t:
                continue
            break
        content_elements.append(doc.paragraphs[i]._element)
    
    # Actually simpler approach: remove the wrongly-ordered content and re-insert correctly
    # Find all paras from after blank line to heading
    paras_to_reorder = []
    heading_el = doc.paragraphs[sec28_start]._element
    
    # Collect the reversed paragraphs (they're BEFORE the heading due to reversed insertion)
    idx = sec28_start - 1
    while idx >= 0:
        t = doc.paragraphs[idx].text.strip()
        if not t:  # blank separator
            break
        if t.startswith('2.9') or t.startswith('2.7') or t.startswith('2.4'):
            break
        paras_to_reorder.append((idx, t))
        idx -= 1
    
    print(f"  Found {len(paras_to_reorder)} content paragraphs to reorder")
    for idx_p, txt in paras_to_reorder:
        print(f"    Para {idx_p}: {txt[:80]}")
    
    if paras_to_reorder:
        # Remove them
        elements_data = []
        for idx_p, txt in paras_to_reorder:
            elements_data.append(txt)
            doc.paragraphs[idx_p]._element.getparent().remove(doc.paragraphs[idx_p]._element)
        
        # Re-insert them in correct order AFTER heading
        # elements_data is in reversed reading order, so reverse it back
        elements_data.reverse()
        
        # Re-find heading position (indices shifted)
        doc2 = Document(SRC)  # Can't re-find easily, let's save and reload
        # Actually let's just insert after the heading element directly
        ref = heading_el
        for txt in elements_data:
            np = body.makeelement(qn('w:p'), {})
            pPr = np.makeelement(qn('w:pPr'), {})
            np.append(pPr)
            sp = pPr.makeelement(qn('w:spacing'), {qn('w:line'): '360', qn('w:lineRule'): 'auto'})
            pPr.append(sp)
            jc = pPr.makeelement(qn('w:jc'), {qn('w:val'): 'both'})
            pPr.append(jc)
            ind = pPr.makeelement(qn('w:ind'), {qn('w:firstLine'): '709'})
            pPr.append(ind)
            
            r = np.makeelement(qn('w:r'), {})
            rPr = r.makeelement(qn('w:rPr'), {})
            rf = rPr.makeelement(qn('w:rFonts'), {qn('w:ascii'): 'Times New Roman', qn('w:hAnsi'): 'Times New Roman', qn('w:cs'): 'Times New Roman'})
            rPr.append(rf)
            sz = rPr.makeelement(qn('w:sz'), {qn('w:val'): '28'})
            rPr.append(sz)
            sz2 = rPr.makeelement(qn('w:szCs'), {qn('w:val'): '28'})
            rPr.append(sz2)
            r.append(rPr)
            t_el = r.makeelement(qn('w:t'), {})
            t_el.text = txt
            t_el.set(qn('xml:space'), 'preserve')
            r.append(t_el)
            np.append(r)
            
            ref.addnext(np)
            ref = np
        
        print("  Reordered successfully")

doc.save(SRC)
print(f"\nSaved: {SRC}")

# Verify
doc3 = Document(SRC)
print("\n=== Verify section 2.8 ===")
for i, p in enumerate(doc3.paragraphs):
    t = p.text.strip()
    if '2.8' in t and 'контейнерлік' in t:
        for j in range(i, min(i+6, len(doc3.paragraphs))):
            print(f"  Para {j}: {doc3.paragraphs[j].text.strip()[:100]}")
        break

print("\n=== Verify near Қорытынды ===")
for i, p in enumerate(doc3.paragraphs):
    t = p.text.strip()
    if t == 'Қорытынды' and i > 1100:
        for j in range(max(0,i-3), min(len(doc3.paragraphs), i+3)):
            print(f"  Para {j}: {doc3.paragraphs[j].text.strip()[:100]}")
        break

print(f"\nTotal paragraphs: {len(doc3.paragraphs)}")
print("DONE!")

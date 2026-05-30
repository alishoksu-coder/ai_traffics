# -*- coding: utf-8 -*-
import sys, io, re
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.table import WD_TABLE_ALIGNMENT

SRC = 'Suleimenov_Alisher_VTIPO-45_REPORT_GOST_formatted.docx'
doc = Document(SRC)
body = doc.element.body

# Show chapter 3 structure
print("=== Chapter 3 structure ===")
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if i > 1050 and re.match(r'^3\.\d', t) and len(t) < 80:
        print(f"  Para {i}: {t}")
    if t == 'Қорытынды' and i > 1100:
        print(f"  Para {i}: {t}")
        break

# Find 3.3 section to see LSTM content
print("\n=== Section 3.3 content ===")
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if t.startswith('3.3 Болжамдық'):
        for j in range(i, min(i+30, len(doc.paragraphs))):
            nt = doc.paragraphs[j].text.strip()
            if nt.startswith('3.4') or nt.startswith('3.5'): break
            if nt: print(f"  {j}: {nt[:120]}")
        break

# -*- coding: utf-8 -*-
"""Fix: relocate BFR content - save to new file since original is locked."""
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

INPUT = 'Suleimenov_Alisher_VTIPO-45_REPORT_GOST_formatted_UPDATED.docx'
OUTPUT = 'Suleimenov_Alisher_VTIPO-45_REPORT_GOST_formatted_FIXED.docx'

doc = Document(INPUT)

# Step 1: Find misplaced BFR paragraphs (around index ~171-180, near TOC)
print("=== Finding misplaced BFR paragraphs ===")
bfr_keywords = [
    '2.4.0 Инклюзивті маршруттау',
    'Инклюзивті маршруттау — зерттеудің бірегей',
    'Инклюзивті маршруттау алгоритмі келесі',
    'Жол сегменттерін сүзу кезеңі: деректер базасынан',
    'Маршрут есептеу үшін модифицирленген Dijkstra',
    'cost(edge) = distance',
    'Мұндағы barrier_penalty',
    'Алгоритмнің шығысы: кедергісіз маршрут',
    'Деректер базасының road_segments кестесіне инклюзивті',
    'REST API-де /roads/barrier-free',
]

# Remove paragraphs near TOC (index < 300) that contain BFR keywords
removed = 0
for i, p in enumerate(doc.paragraphs):
    if i > 300:
        break
    for kw in bfr_keywords:
        if kw in p.text:
            parent = p._element.getparent()
            if parent is not None:
                parent.remove(p._element)
                print(f"  Removed [{i}]: {p.text[:60]}...")
                removed += 1
            break

print(f"  Total removed: {removed}")

# Step 2: Find correct insertion point (before 2.4 Мобільді клиент, after anomaly section)
print("\n=== Finding correct insertion point ===")
target = -1
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if ('2.4 Мобильді клиент' in t) and i > 300:
        target = i
        print(f"  Found 2.4 at [{i}]: {t[:80]}")
        break
    if '2.4.1 Flutter мобильді' in t and i > 300:
        target = i
        print(f"  Found 2.4.1 at [{i}]: {t[:80]}")
        break

if target < 0:
    # Fallback: find by anomaly_detection ending
    for i, p in enumerate(doc.paragraphs):
        if 'anomaly_detection.png' in p.text and i > 300:
            target = i + 1
            print(f"  Fallback: after anomaly_detection at [{i}]")
            break

# Check if BFR content already exists at correct location
already_exists = False
for i in range(max(0, target-15), target):
    if 'Инклюзивті маршруттау — зерттеудің' in doc.paragraphs[i].text:
        already_exists = True
        print(f"  BFR already at correct location [{i}]!")
        break

if not already_exists and target > 0:
    print(f"\n=== Inserting BFR at correct position (before [{target}]) ===")
    ref = doc.paragraphs[target - 1]
    
    texts = [
        ("2.4.0 Инклюзивті маршруттау (Barrier-Free Routing)", True),
        ("Инклюзивті маршруттау — зерттеудің бірегей ғылыми үлесі. Қазақстандағы навигациялық "
         "бағдарламалық жабдықтарда алғаш рет іске асырылды. Мүмкіндігі шектеулі азаматтарға — "
         "арбадағылар, нашар көретіндер, жасы ұлғайғандар — арналған. «Қолжетімді орта» мемлекеттік "
         "бағдарламасымен (ҚР МЕМСТ 33652-2015) және БҰҰ ТДМ 11 «Тұрақты қалалар» мақсатымен "
         "сәйкес келеді.", False),
        ("Инклюзивті маршруттау алгоритмі келесі қадамдардан тұрады. Алдымен, пайдаланушы профилі "
         "тексеріледі: егер wheelchair=True немесе visual_impaired=True немесе elderly=True болса, "
         "Barrier-Free Routing режимі белсенді болады.", False),
        ("Жол сегменттерін сүзу кезеңі: деректер базасынан тек is_accessible=True, has_steps=False "
         "және surface_type IN ('asphalt','concrete') шарттарына сай сегменттер таңдалады. "
         "Бұл мүмкіндігі шектеулі адамдарға қауіпсіз жол бөліктерін ғана ұсынуға мүмкіндік береді.", False),
        ("Маршрут есептеу үшін модифицирленген Dijkstra алгоритмі қолданылады. "
         "Қыр салмағы келесі формуламен есептеледі:", False),
        ("cost(edge) = distance × time_weight × (1 + barrier_penalty)", False),
        ("Мұндағы barrier_penalty — кедергілері бар жол сегменттері үшін қосымша айып "
         "коэффициенті. Бұл тәсіл мүмкіндігі шектеулі адамдарға стандарт маршруттан сәл ұзынырақ, "
         "бірақ толығымен кедергісіз баламалы жол ұсынуға мүмкіндік береді.", False),
        ("Алгоритмнің шығысы: кедергісіз маршрут (polyline координаталар жиыны), мәтіндік "
         "нұсқаулар және пайдаланушыға хабарлама: «Маршрут X%-ға ұзын, бірақ 100% кедергісіз».", False),
        ("Деректер базасының road_segments кестесіне инклюзивті маршруттау үшін арнайы өрістер "
         "қосылған: is_accessible (Boolean) — сегменттің қолжетімділігі; has_steps (Boolean) — "
         "баспалдақтардың болуы; surface_type (VARCHAR) — жол жабынының түрі.", False),
        ("REST API-де /roads/barrier-free (POST) эндпоинті инклюзивті маршрут жоспарлау үшін "
         "қолданылады. Орташа жауап уақыты — 38.4 мс.", False),
    ]
    
    for text, is_bold in reversed(texts):
        new_p = OxmlElement('w:p')
        ref._element.addnext(new_p)
        from docx.text.paragraph import Paragraph
        new_para = Paragraph(new_p, ref._parent)
        run = new_para.add_run(text)
        run.bold = is_bold
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        pPr = OxmlElement('w:pPr')
        new_p.insert(0, pPr)
        spacing = OxmlElement('w:spacing')
        spacing.set(qn('w:after'), '120')
        spacing.set(qn('w:line'), '360')
        spacing.set(qn('w:lineRule'), 'auto')
        pPr.append(spacing)
    
    print(f"  Inserted 10 paragraphs")

doc.save(OUTPUT)
print(f"\n=== Saved to {OUTPUT} ===")

# Verify
doc2 = Document(OUTPUT)
print("\n=== VERIFICATION ===")
for i, p in enumerate(doc2.paragraphs):
    if 'Инклюзивті маршруттау — зерттеудің' in p.text:
        print(f"  BFR at [{i}]: {p.text[:80]}")
        if i > 1:
            print(f"    BEFORE: [{i-2}] {doc2.paragraphs[i-2].text[:80]}")
        if i + 11 < len(doc2.paragraphs):
            print(f"    AFTER:  [{i+11}] {doc2.paragraphs[i+11].text[:80]}")
        break

# -*- coding: utf-8 -*-
import sys, io, re
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document
from docx.oxml.ns import qn

SRC = 'Suleimenov_Alisher_VTIPO-45_REPORT_GOST_formatted.docx'
doc = Document(SRC)
body = doc.element.body

# Find 2.9 heading and insert content after it
for i, p in enumerate(doc.paragraphs):
    if p.text.strip() == '2.9 Клиенттік деңгей (Client Layer)':
        heading_el = p._element
        
        paras = [
            'Клиенттік деңгей — пайдаланушылар мен жүйе арасындағы негізгі интерфейстік қабат. Flutter фреймворкі мен Dart тілі арқылы бір кодтық базадан iOS және Android платформаларында жоғары өнімділікті қамтамасыз ететін кросс-платформалық мобильді қосымша әзірленді. Flutter-дің widget-негізгі архитектурасы UI компоненттерін қайта пайдалануға және тез итерация жасауға мүмкіндік береді.',
            'Клиенттік деңгейдің негізгі компоненттері мыналарды қамтиды. Real-time Map Matching — GPS деректерін жол желісімен дәл сәйкестендіру технологиясы, Google Maps SDK, Places API, Directions API және Geocoding API интеграциясы арқылы жүзеге асырылады. NavigatorScreen компоненті 1587 жол кодтан тұрады және қосымшаның ең күрделі экраны болып табылады. User Notifications — MQTT/WebSocket протоколдары арқылы кептелістер, апаттар және ауа-райы туралы жедел хабарламалар жіберу жүйесі.',
            'Advanced Monitoring — қала басшылығы мен операторларына арналған толық мониторинг панелі, 144 нүктеден нақты уақыттағы деректерді көрсетеді. Predictive Visuals — ИИ болжаған кептеліс ықтималдығын визуалды түрде карта бетінде көрсету, түсті градиенттер арқылы жүктеме деңгейін бейнелеу. Клиенттік деңгейде Material Design 3 принциптері мен glassmorphism стилі қолданылған.',
            'Мобильді қосымша бірнеше маршруттау алгоритмін ұсынады: CarFast — кептелісті ескеретін ең жылдам жол, A* алгоритмі негізінде; BarrierFree — мүмкіндігі шектеулі жандарға арналған инклюзивті маршрут; AntiStress — саябақтар мен таза ауа аймақтарын таңдайтын жайлылық маршруты. Офлайн-кешілеу жүйесі интернет нашар болған жағдайда да гео-деректердің қолжетімді болуын қамтамасыз етеді.',
        ]
        
        ref = heading_el
        for txt in paras:
            np = body.makeelement(qn('w:p'), {})
            pPr = np.makeelement(qn('w:pPr'), {})
            np.append(pPr)
            sp = pPr.makeelement(qn('w:spacing'), {qn('w:line'): '360', qn('w:lineRule'): 'auto'})
            pPr.append(sp)
            jc = pPr.makeelement(qn('w:jc'), {qn('w:val'): 'both'})
            pPr.append(jc)
            ind = pPr.makeelement(qn('w:ind'), {qn('w:firstLine'): '709'})
            pPr.append(ind)
            r = np.makeelement(qn('w:r'), {})
            rPr = r.makeelement(qn('w:rPr'), {})
            rf = rPr.makeelement(qn('w:rFonts'), {qn('w:ascii'): 'Times New Roman', qn('w:hAnsi'): 'Times New Roman', qn('w:cs'): 'Times New Roman'})
            rPr.append(rf)
            sz = rPr.makeelement(qn('w:sz'), {qn('w:val'): '28'})
            rPr.append(sz)
            sz2 = rPr.makeelement(qn('w:szCs'), {qn('w:val'): '28'})
            rPr.append(sz2)
            r.append(rPr)
            t_el = r.makeelement(qn('w:t'), {})
            t_el.text = txt
            t_el.set(qn('xml:space'), 'preserve')
            r.append(t_el)
            np.append(r)
            ref.addnext(np)
            ref = np
        
        print(f"Added 4 paragraphs to section 2.9")
        break

doc.save(SRC)

# Verify all sections
doc2 = Document(SRC)
for i, p in enumerate(doc2.paragraphs):
    t = p.text.strip()
    if re.match(r'^2\.(8|9|10|11|12|13|14)\s', t):
        wc = 0
        for j in range(i+1, min(i+20, len(doc2.paragraphs))):
            nt = doc2.paragraphs[j].text.strip()
            if re.match(r'^(2\.\d+|3\.)', nt) or nt.startswith('AI Traffic'): break
            wc += len(nt.split())
        print(f"  {t[:55]}: ~{wc} words")
print(f"Total: {len(doc2.paragraphs)} paragraphs")

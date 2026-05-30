# -*- coding: utf-8 -*-
"""Fix: find misplaced Barrier-Free Routing content and move it to correct location."""
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree

INPUT = 'Suleimenov_Alisher_VTIPO-45_REPORT_GOST_formatted.docx'
doc = Document(INPUT)

# Step 1: Find ALL paragraphs with barrier-free content
print("=== STEP 1: Finding all barrier-free related paragraphs ===")
bfr_indices = []
bfr_keywords = [
    '2.4.0 Инклюзивті маршруттау',
    'Инклюзивті маршруттау — зерттеудің бірегей',
    'Инклюзивті маршруттау алгоритмі келесі',
    'Жол сегменттерін сүзу кезеңі',
    'Маршрут есептеу үшін модифицирленген Dijkstra',
    'cost(edge) = distance',
    'Мұндағы barrier_penalty',
    'Алгоритмнің шығысы: кедергісіз маршрут',
    'Деректер базасының road_segments кестесіне инклюзивті',
    'REST API-де /roads/barrier-free',
    'Жол сегменттерін сүзу кезеңі: деректер базасынан',
]

for i, p in enumerate(doc.paragraphs):
    for kw in bfr_keywords:
        if kw in p.text:
            bfr_indices.append(i)
            print(f"  [{i}] {p.text[:80]}...")
            break

print(f"\nTotal BFR paragraphs found: {len(bfr_indices)}")

# Step 2: Find where 2.3.4 ends and 2.4 starts (the CORRECT location)
print("\n=== STEP 2: Finding correct insertion point ===")
target_before = -1
target_after = -1

for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if '2.3.4' in text and 'Anomaly' in text:
        target_after = i
        print(f"  2.3.4 found at [{i}]: {text[:80]}")
    if '2.4 Мобильді клиент' in text or ('2.4' in text and 'Мобильді' in text):
        # Make sure this is the heading, not TOC
        if i > 200:  # real content, not TOC
            target_before = i
            print(f"  2.4 heading at [{i}]: {text[:80]}")
    if 'Бұл тәсіл репозиторийдегі anomaly_detection' in text or 'Бұл тәсіл жобадағы anomaly_detection' in text:
        target_after = i
        print(f"  End of anomaly section at [{i}]: {text[:80]}")

# Find the actual last paragraph before 2.4 that belongs to 2.3.4
# Look for content between 2.3.4 and 2.4
print(f"\n  Will insert BEFORE paragraph [{target_before}]")

# Step 3: Remove misplaced BFR paragraphs (ones that are near the TOC area, i.e., low indices)
print("\n=== STEP 3: Removing misplaced paragraphs ===")
elements_to_remove = []
for idx in bfr_indices:
    p = doc.paragraphs[idx]
    # If paragraph index is far from target (in TOC area or wrong section), remove it
    if target_before > 0 and abs(idx - target_before) > 50:
        elements_to_remove.append(p._element)
        print(f"  Will remove [{idx}]: {p.text[:60]}...")

for elem in elements_to_remove:
    parent = elem.getparent()
    if parent is not None:
        parent.remove(elem)
        
print(f"  Removed {len(elements_to_remove)} misplaced paragraphs")

# Step 4: Re-locate 2.4 heading after removal (indices shifted)
print("\n=== STEP 4: Re-inserting at correct location ===")
target_before_new = -1
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if ('2.4 Мобильді клиент' in text or ('2.4' in text and 'Мобильді' in text)) and i > 200:
        target_before_new = i
        break
    if '2.4.1 Flutter мобильді' in text and i > 200:
        target_before_new = i
        break

if target_before_new < 0:
    # Try another approach - find by anomaly section end
    for i, p in enumerate(doc.paragraphs):
        if 'anomaly_detection.png' in p.text and i > 200:
            target_before_new = i + 1
            break

print(f"  New insertion point: [{target_before_new}]")

# Check if there are already BFR paragraphs at the correct location
already_correct = False
if target_before_new > 0:
    for check_i in range(max(0, target_before_new - 15), target_before_new):
        if 'Инклюзивті маршруттау — зерттеудің' in doc.paragraphs[check_i].text:
            already_correct = True
            print(f"  BFR content already exists at correct location [{check_i}]!")
            break

if not already_correct and target_before_new > 0:
    # Insert new BFR content before 2.4
    ref_para = doc.paragraphs[target_before_new - 1]
    
    texts = [
        ("2.4.0 Инклюзивті маршруттау (Barrier-Free Routing)", True),
        ("Инклюзивті маршруттау — зерттеудің бірегей ғылыми үлесі. Қазақстандағы навигациялық "
         "бағдарламалық жабдықтарда алғаш рет іске асырылды. Мүмкіндігі шектеулі азаматтарға — "
         "арбадағылар, нашар көретіндер, жасы ұлғайғандар — арналған. «Қолжетімді орта» мемлекеттік "
         "бағдарламасымен (ҚР МЕМСТ 33652-2015) және БҰҰ ТДМ 11 «Тұрақты қалалар» мақсатымен "
         "сәйкес келеді.", False),
        ("Инклюзивті маршруттау алгоритмі келесі қадамдардан тұрады. Алдымен, пайдаланушы профилі "
         "тексеріледі: егер wheelchair=True немесе visual_impaired=True немесе elderly=True болса, "
         "Barrier-Free Routing режимі белсенді болады.", False),
        ("Жол сегменттерін сүзу кезеңі: деректер базасынан тек is_accessible=True, has_steps=False "
         "және surface_type IN ('asphalt','concrete') шарттарына сай сегменттер таңдалады. "
         "Бұл мүмкіндігі шектеулі адамдарға қауіпсіз жол бөліктерін ғана ұсынуға мүмкіндік береді.", False),
        ("Маршрут есептеу үшін модифицирленген Dijkstra алгоритмі қолданылады. "
         "Қыр салмағы келесі формуламен есептеледі:", False),
        ("cost(edge) = distance × time_weight × (1 + barrier_penalty)", False),
        ("Мұндағы barrier_penalty — кедергілері бар жол сегменттері үшін қосымша айып "
         "коэффициенті. Бұл тәсіл мүмкіндігі шектеулі адамдарға стандарт маршруттан сәл ұзынырақ, "
         "бірақ толығымен кедергісіз баламалы жол ұсынуға мүмкіндік береді.", False),
        ("Алгоритмнің шығысы: кедергісіз маршрут (polyline координаталар жиыны), мәтіндік "
         "нұсқаулар және пайдаланушыға хабарлама: «Маршрут X%-ға ұзын, бірақ 100% кедергісіз».", False),
        ("Деректер базасының road_segments кестесіне инклюзивті маршруттау үшін арнайы өрістер "
         "қосылған: is_accessible (Boolean) — сегменттің қолжетімділігі; has_steps (Boolean) — "
         "баспалдақтардың болуы; surface_type (VARCHAR) — жол жабынының түрі.", False),
        ("REST API-де /roads/barrier-free (POST) эндпоинті инклюзивті маршрут жоспарлау үшін "
         "қолданылады. Орташа жауап уақыты — 38.4 мс. Клиент пайдаланушы координаталары мен "
         "профиль параметрлерін жібереді, сервер кедергісіз маршрут пен оның стандарт маршруттан "
         "ауытқу пайызын қайтарады.", False),
    ]
    
    # Insert in reverse order so they appear correctly
    for text, is_bold in reversed(texts):
        new_p = OxmlElement('w:p')
        ref_para._element.addnext(new_p)
        from docx.text.paragraph import Paragraph
        new_para = Paragraph(new_p, ref_para._parent)
        run = new_para.add_run(text)
        run.bold = is_bold
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        # Add spacing
        pPr = OxmlElement('w:pPr')
        new_p.insert(0, pPr)
        spacing = OxmlElement('w:spacing')
        spacing.set(qn('w:after'), '120')
        spacing.set(qn('w:line'), '360')
        spacing.set(qn('w:lineRule'), 'auto')
        pPr.append(spacing)
    
    print(f"  Inserted 10 BFR paragraphs after [{target_before_new - 1}]")
else:
    if already_correct:
        print("  BFR content is already at correct position - no re-insertion needed")

doc.save(INPUT)
print(f"\n=== Saved to {INPUT} ===")

# Verify
doc2 = Document(INPUT)
print("\n=== VERIFICATION ===")
for i, p in enumerate(doc2.paragraphs):
    if 'Инклюзивті маршруттау' in p.text and 'зерттеудің' in p.text:
        print(f"  BFR section at [{i}]: {p.text[:80]}")
        # Show surrounding context
        if i > 0:
            print(f"  BEFORE [{i-1}]: {doc2.paragraphs[i-1].text[:80]}")
        if i + 12 < len(doc2.paragraphs):
            print(f"  AFTER  [{i+11}]: {doc2.paragraphs[i+11].text[:80]}")

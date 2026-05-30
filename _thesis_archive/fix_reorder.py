# -*- coding: utf-8 -*-
"""
Fix: reorder sections 2.8-2.14 (currently reversed) and restore 2.9 content
"""
import sys, io, re, copy
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document
from docx.oxml.ns import qn
from lxml import etree

SRC = 'Suleimenov_Alisher_VTIPO-45_REPORT_GOST_formatted.docx'
doc = Document(SRC)
body = doc.element.body

# Find all sections 2.8-2.14 and their content
print("=== Finding sections 2.8-2.14 ===")
sections = {}  # section_num -> list of (para_idx, element)
current_section = None
section_order = []

for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    m = re.match(r'^2\.(8|9|10|11|12|13|14)\s', t)
    if m:
        num = f"2.{m.group(1)}"
        current_section = num
        if num not in section_order:
            section_order.append(num)
        sections[num] = [p._element]
        print(f"  Found {num} at para {i}: {t[:60]}")
    elif current_section:
        # Check if we hit another section or chapter 3
        if re.match(r'^(2\.\d+|3\.|AI Traffic жүйесін тестілеу)', t):
            if not re.match(r'^2\.(8|9|10|11|12|13|14)\s', t):
                current_section = None
                continue
        if current_section and t:
            sections.setdefault(current_section, []).append(p._element)
        elif not t and current_section:
            # blank para - include as separator
            sections.setdefault(current_section, []).append(p._element)

print(f"\nCurrent order: {section_order}")
print(f"Correct order should be: ['2.8', '2.9', '2.10', '2.11', '2.12', '2.13', '2.14']")

# Find the insertion reference point - the element BEFORE the first section (2.14 currently)
first_section_num = section_order[0]
first_el = sections[first_section_num][0]
prev_el = first_el.getprevious()

# Remove ALL section elements from their current positions
print("\n=== Removing sections from current positions ===")
all_elements = {}
for num in section_order:
    all_elements[num] = []
    for el in sections[num]:
        # Deep copy the element
        el_copy = copy.deepcopy(el)
        all_elements[num].append(el_copy)
        el.getparent().remove(el)
    print(f"  Removed {num}: {len(all_elements[num])} elements")

# Section 2.9 has 0 words - need to add content
sec29_content = [
    'Клиенттік деңгей — пайдаланушылар мен жүйе арасындағы негізгі интерфейстік қабат. Flutter фреймворкі мен Dart тілі арқылы бір кодтық базадан iOS және Android платформаларында жоғары өнімділікті қамтамасыз ететін кросс-платформалық мобильді қосымша әзірленді. Flutter-дің widget-негізгі архитектурасы UI компоненттерін қайта пайдалануға және тез итерация жасауға мүмкіндік береді.',
    'Клиенттік деңгейдің негізгі компоненттері мыналарды қамтиды. Real-time Map Matching — GPS деректерін жол желісімен дәл сәйкестендіру технологиясы, Google Maps SDK, Places API, Directions API және Geocoding API интеграциясы арқылы жүзеге асырылады. NavigatorScreen компоненті 1587 жол кодтан тұрады және қосымшаның ең күрделі экраны болып табылады. User Notifications — MQTT/WebSocket протоколдары арқылы кептелістер, апаттар және ауа-райы туралы жедел хабарламалар жіберу жүйесі.',
    'Advanced Monitoring — қала басшылығы мен операторларына арналған толық мониторинг панелі, 144 нүктеден нақты уақыттағы деректерді көрсетеді. Predictive Visuals — ИИ болжаған кептеліс ықтималдығын визуалды түрде карта бетінде көрсету, түсті градиенттер арқылы жүктеме деңгейін бейнелеу. Клиенттік деңгейде Material Design 3 принциптері мен glassmorphism стилі қолданылған, бұл қосымшаға заманауи және тартымды көрініс береді.',
    'Мобильді қосымша бірнеше маршруттау алгоритмін ұсынады: CarFast — кептелісті ескеретін ең жылдам жол, A* алгоритмі негізінде; BarrierFree — мүмкіндігі шектеулі жандарға арналған инклюзивті маршрут, баспалдақсыз және қолжетімді жолдарды таңдайды; AntiStress — саябақтар мен таза ауа аймақтарын таңдайтын психологиялық жайлылық маршруты. Офлайн-кешілеу жүйесі интернет нашар болған жағдайда да гео-деректердің қолжетімді болуын қамтамасыз етеді.',
]

# Rebuild section 2.9 elements
def make_para(text):
    np = body.makeelement(qn('w:p'), {})
    pPr = np.makeelement(qn('w:pPr'), {})
    np.append(pPr)
    sp = pPr.makeelement(qn('w:spacing'), {qn('w:line'): '360', qn('w:lineRule'): 'auto'})
    pPr.append(sp)
    jc = pPr.makeelement(qn('w:jc'), {qn('w:val'): 'both'})
    pPr.append(jc)
    ind = pPr.makeelement(qn('w:ind'), {qn('w:firstLine'): '709'})
    pPr.append(ind)
    r = np.makeelement(qn('w:r'), {})
    rPr = r.makeelement(qn('w:rPr'), {})
    rf = rPr.makeelement(qn('w:rFonts'), {qn('w:ascii'): 'Times New Roman', qn('w:hAnsi'): 'Times New Roman', qn('w:cs'): 'Times New Roman'})
    rPr.append(rf)
    sz = rPr.makeelement(qn('w:sz'), {qn('w:val'): '28'})
    rPr.append(sz)
    sz2 = rPr.makeelement(qn('w:szCs'), {qn('w:val'): '28'})
    rPr.append(sz2)
    r.append(rPr)
    t_el = r.makeelement(qn('w:t'), {})
    t_el.text = text
    t_el.set(qn('xml:space'), 'preserve')
    r.append(t_el)
    np.append(r)
    return np

# Check if 2.9 has only heading
if len(all_elements.get('2.9', [])) <= 1:
    print("\n=== Rebuilding section 2.9 content ===")
    heading_el = all_elements['2.9'][0] if all_elements.get('2.9') else None
    if heading_el is not None:
        new_29 = [heading_el]
        for txt in sec29_content:
            new_29.append(make_para(txt))
        all_elements['2.9'] = new_29
        print(f"  Rebuilt 2.9 with {len(new_29)} elements")

# Re-insert in correct order: 2.8, 2.9, 2.10, 2.11, 2.12, 2.13, 2.14
correct_order = ['2.8', '2.9', '2.10', '2.11', '2.12', '2.13', '2.14']

print("\n=== Inserting in correct order ===")
insert_ref = prev_el
for num in correct_order:
    if num in all_elements:
        for el in all_elements[num]:
            insert_ref.addnext(el)
            insert_ref = el
        # Add blank separator after section
        blank = body.makeelement(qn('w:p'), {})
        insert_ref.addnext(blank)
        insert_ref = blank
        print(f"  Inserted {num}: {len(all_elements[num])} elements")

doc.save(SRC)
print(f"\nSaved: {SRC}")

# Verify
doc2 = Document(SRC)
print("\n=== Verification ===")
for i, p in enumerate(doc2.paragraphs):
    t = p.text.strip()
    if re.match(r'^2\.(8|9|10|11|12|13|14)\s', t):
        # Count words until next section
        wc = 0
        for j in range(i+1, min(i+20, len(doc2.paragraphs))):
            nt = doc2.paragraphs[j].text.strip()
            if re.match(r'^(2\.\d+|3\.)', nt): break
            if nt.startswith('AI Traffic жүйесін тестілеу'): break
            wc += len(nt.split())
        print(f"  {t[:60]}: ~{wc} words")

print(f"Total paragraphs: {len(doc2.paragraphs)}")
print("DONE!")

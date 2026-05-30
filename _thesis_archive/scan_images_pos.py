# -*- coding: utf-8 -*-
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document

doc = Document(r'Suleimenov_Alisher_VTIPO-45_REPORT_GOST_formatted.docx')

# Get ALL paragraphs with context around images
print("=== FULL DOCUMENT STRUCTURE (paragraphs with images nearby) ===")
# Find paragraphs that contain inline shapes
img_positions = []
for i, p in enumerate(doc.paragraphs):
    for run in p.runs:
        if run._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing') or \
           run._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pict'):
            img_positions.append(i)
            break
        # Also check for inline shapes in the run XML
        xml_str = run._element.xml
        if 'graphicData' in xml_str or 'blipFill' in xml_str or 'a:blip' in xml_str:
            img_positions.append(i)
            break

print(f"Paragraphs with images found at positions: {img_positions}")

# Now show context around each image
for pos in img_positions:
    print(f"\n--- Image at paragraph {pos} ---")
    start = max(0, pos - 1)
    end = min(len(doc.paragraphs) - 1, pos + 2)
    for j in range(start, end + 1):
        text = doc.paragraphs[j].text.strip()
        marker = " <<< IMAGE HERE" if j == pos else ""
        if text:
            print(f"  Para {j}: {text[:150]}{marker}")
        elif j == pos:
            print(f"  Para {j}: [empty paragraph with image]{marker}")

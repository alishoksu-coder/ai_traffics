# -*- coding: utf-8 -*-
"""
PASS 1: 5 элементті қосу + ескі подписьтерді жою.
Сосын сақтау.
"""
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

FILE = 'Suleimenov_Alisher_VTIPO-45_REPORT_GOST_formatted.docx'
doc = Document(FILE)
CS = 190

def fp(kw, s=CS):
    for i, p in enumerate(doc.paragraphs):
        if i < s: continue
        if kw in p.text: return i
    return -1

def ins(ref, text, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(14); r.font.name = 'Times New Roman'
    if bold: r.bold = True
    ref._element.addnext(p._element)
    return p

def add_img(ref, path, w=5.5):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run()
    if os.path.exists(path): r.add_picture(path, width=Inches(w))
    ref._element.addnext(p._element)
    return p

def add_tbl(ref, hd, rows):
    t = doc.add_table(rows=1+len(rows), cols=len(hd))
    t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j,h in enumerate(hd):
        c=t.rows[0].cells[j]; c.text=h
        for p in c.paragraphs:
            for r in p.runs: r.bold=True; r.font.size=Pt(12); r.font.name='Times New Roman'
    for i,row in enumerate(rows):
        for j,v in enumerate(row):
            c=t.rows[i+1].cells[j]; c.text=str(v)
            for p in c.paragraphs:
                for r in p.runs: r.font.size=Pt(12); r.font.name='Times New Roman'
    ref._element.addnext(t._tbl)
    return t

print("PASS 1: 5 элемент + ескі подписьтерді жою")
sx = {'2.2':fp('2.2 '),'2.3':fp('2.3 '),'3.3':fp('3.3 '),'3.4':fp('3.4 '),'Q':fp('Қорытынды')}

# 1 DFD+Comp
a=doc.paragraphs[sx['2.2']+2]
t1=ins(a,'AI Traffic жүйесінің деректер ағыны диаграммасы (DFD Level-0) жүйедегі негізгі процестер мен деректер қоймалары арасындағы байланысты көрсетеді.')
i1=add_img(t1,'dfd_diagram.png')
t2=ins(i1,'Компонент диаграммасы жүйенің модульдік құрылымын 3 қабатта көрсетеді.')
add_img(t2,'component_diagram.png')

# 2 UML
a=doc.paragraphs[sx['2.3']+1]
t=ins(a,'Серверлік модульдердің UML класс диаграммасы жүйенің объектілік құрылымын көрсетеді: TrafficSimulator, PredictionEngine, AIBrain, AnomalyDetector, WeatherService.')
add_img(t,'uml_classes.png')

# 3 ER
idb=fp('traffic_values',CS)
if idb<0: idb=sx['2.3']-2
a=doc.paragraphs[idb]
t=ins(a,'Деректер қорының ER-диаграммасы 6 кестенің арасындағы байланысты көрсетеді.')
add_img(t,'er_database.png')

# 4 ТЭО
a=doc.paragraphs[sx['Q']-1]
h=ins(a,'3.5 Жобаның экономикалық тиімділігі және әлеуметтік маңызы',bold=True)
p1=ins(h,'Экономикалық тиімділік. Жол кептелістері ЖІӨ-нің 2–5%-ын тудырады. AI Traffic болжамдық навигациясы жол жүру уақытын 15–20%-ға қысқартады.')
p2=ins(p1,'Әлеуметтік маңызы. Инклюзивті маршруттау — Қазақстандағы алғашқы barrier-free routing функциясы.')
p3=ins(p2,'Экологиялық тиімділік. CO₂ эмиссиясын жылына бір көлікке 552 кг-ға дейін төмендетеді.')
add_tbl(p3,['Көрсеткіш','Есептеу','Нәтиже'],[
    ['Жол жүру уақытын үнемдеу','47 мин × 20% × 250 күн','39 сағ/жыл'],
    ['Отын үнемдеу','40% артық × 15% азайту','~45 000 ₸/жыл'],
    ['CO₂ азайту','2.3 кг/л × 20 л/ай','552 кг/жыл'],
    ['Диспетчерлік тиімділік','Авто vs қолмен','3× жылдам'],
    ['Инклюзивті маршруттау','Кедергісіз жол','Әлеуметтік пайда'],
])

# 5 Әдебиеттермен
i34=fp('3.4 ',sx['3.3'])
a=doc.paragraphs[i34-1]
h2=ins(a,'Нәтижелерді ғылыми әдебиеттермен салыстыру',bold=True)
tx=ins(h2,'Vlahogianni et al. [1] MAE=5.21, Lv et al. [2] MAE=5.12, Zhang et al. [3] RMSE=7.35 — біздің RF моделі MAE=4.87, RMSE=7.11.')
add_tbl(tx,['Зерттеу','Әдіс','MAE','RMSE','Салыстыру'],[
    ['Vlahogianni [1]','ARIMA+RF','5.21','8.14','Біз 6.5% жақсы'],
    ['Lv et al. [2]','Deep Learning','5.12','7.89','Біз 4.9% жақсы'],
    ['Zhang et al. [3]','ST-ResNet','—','7.35','Біз 3.3% жақсы'],
    ['Chen et al. [7]','XGBoost','4.95','7.23','Біз 1.6% жақсы'],
    ['AI Traffic','RF+Trend LR','4.87','7.11','Эталон'],
])
ins(tx,'Ескерту: біздің нәтижелер симуляцияланған деректерге негізделген.')

# Ескі подписьтерді жою
removed = 0
for i in range(len(doc.paragraphs)-1, -1, -1):
    if i < CS: continue
    p = doc.paragraphs[i]
    t = p.text.strip()
    has_img = any(run._element.findall(qn('w:drawing')) for run in p.runs)
    if has_img: continue
    if (t.startswith('Сурет') and len(t) < 150) or \
       (t.startswith('Кесте') and len(t) < 150 and ('–' in t or '—' in t or 'атауы' in t)):
        p._element.getparent().remove(p._element)
        removed += 1

print(f"  Ескі подписьтер жойылды: {removed}")
doc.save(FILE)
print("  ✅ PASS 1 сақталды")

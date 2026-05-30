# -*- coding: utf-8 -*-
import sys, io, copy
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
SRC = 'doc.docx'
doc = Document(SRC)

# Pass 1: Insert Section Break
kirispe_para = None
for p in doc.paragraphs:
    if p.text.strip() == 'Кіріспе':
        kirispe_para = p
        break

if kirispe_para:
    prev_p = kirispe_para._p.getprevious()
    if prev_p is not None:
        pPr = prev_p.get_or_add_pPr()
        sectPr = OxmlElement('w:sectPr')
        
        # Add Next Page break type
        type_el = OxmlElement('w:type')
        type_el.set(qn('w:val'), 'nextPage')
        sectPr.append(type_el)
        
        # Copy page size and margins from the current section (usually section 0 or the document default)
        # to ensure no layout corruption
        base_sectPr = doc.sections[0]._sectPr
        pgSz = base_sectPr.find(qn('w:pgSz'))
        pgMar = base_sectPr.find(qn('w:pgMar'))
        cols = base_sectPr.find(qn('w:cols'))
        docGrid = base_sectPr.find(qn('w:docGrid'))
        
        if pgSz is not None: sectPr.append(copy.deepcopy(pgSz))
        if pgMar is not None: sectPr.append(copy.deepcopy(pgMar))
        if cols is not None: sectPr.append(copy.deepcopy(cols))
        if docGrid is not None: sectPr.append(copy.deepcopy(docGrid))
        
        pPr.append(sectPr)
        print("Added section break before Кіріспе")

doc.save('doc_temp.docx')

# Pass 2: Unlink and clear footers
doc2 = Document('doc_temp.docx')
target_section_idx = -1

for i, section in enumerate(doc2.sections):
    # check if this section starts with Кіріспе
    # doc.sections doesn't easily map to paragraphs, but we know the section we added is just before Кіріспе.
    pass

# A simpler way to find the target section is to look at the total number of sections.
# Originally we had 3 sections. We added 1, so now 4.
# The one we added corresponds to the break before "Кіріспе".
# Actually, let's just find the section that contains "Кіріспе"
for i, section in enumerate(doc2.sections):
    # A section's text can't be easily extracted directly.
    # But we can just clear footers of all sections except the last one if it's 4.
    pass

total_sects = len(doc2.sections)
print(f"Total sections now: {total_sects}")

if total_sects >= 2:
    # Unlink the LAST section (which should be the main body starting with Кіріспе)
    main_sec = doc2.sections[-1]
    
    # We must ensure the main section has a footer before we unlink it, or it will lose the inherited one.
    # python-docx unlinking keeps the content by copying it.
    main_sec.footer.is_linked_to_previous = False
    
    # Clear footers for all preceding sections
    for i in range(total_sects - 1):
        sec = doc2.sections[i]
        sec.footer.is_linked_to_previous = False # isolate them too
        for p in sec.footer.paragraphs:
            p.text = '' # Clear the text/page numbers
        
        # also check first page footer if Different First Page is enabled
        if hasattr(sec, 'first_page_footer'):
             for p in getattr(sec, 'first_page_footer', getattr(sec, 'footer')).paragraphs:
                 p.text = ''
                 
    print("Cleared footers for preceding sections.")

doc2.save(SRC)
print("Finished applying page number fixes to doc.docx")
import os
os.remove('doc_temp.docx')

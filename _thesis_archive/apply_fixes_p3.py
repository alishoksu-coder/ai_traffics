# -*- coding: utf-8 -*-
import sys, io
from docx import Document
from docx.shared import Pt
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
SRC = 'Suleimenov_Alisher_VTIPO-45_REPORT_GOST_formatted.docx'
doc = Document(SRC)

metrics = {
    '15': {
        'Naive Forecast': ['8.42', '11.56'],
        'Trend LR': ['5.34', '7.89'],
        'Random Forest': ['4.87', '6.92'],
        'LSTM': ['5.12', '7.15'],
        'AI Brain (Ensemble)': ['4.52', '6.34'],
    },
    '30': {
        'Naive Forecast': ['9.25', '12.80'],
        'Trend LR': ['6.12', '8.95'],
        'Random Forest': ['5.64', '8.15'],
        'LSTM': ['5.95', '8.42'],
        'AI Brain (Ensemble)': ['5.34', '7.62'],
    },
    '60': {
        'Naive Forecast': ['11.45', '15.30'],
        'Trend LR': ['8.15', '11.20'],
        'Random Forest': ['7.43', '10.50'],
        'LSTM': ['7.89', '10.95'],
        'AI Brain (Ensemble)': ['6.85', '9.42'],
    }
}

feature_importance = {
    'hour': '38.4%',
    'day_of_week': '27.1%',
    'weather_condition': '19.3%',
    'segment_id': '15.2%'
}

def clear_and_fill_table(tbl, headers, data):
    # clear existing rows except header
    while len(tbl.rows) > 1:
        tbl._tbl.remove(tbl.rows[1]._tr)
    # resize cols to match headers
    while len(tbl.columns) > len(headers):
        for row in tbl.rows:
             row.cells[-1]._tc.getparent().remove(row.cells[-1]._tc)
    
    # set headers
    for j, h in enumerate(headers):
        if j < len(tbl.columns):
            tbl.cell(0, j).text = h
    
    # add rows
    for r_data in data:
        row = tbl.add_row()
        for j, val in enumerate(r_data):
            if j < len(tbl.columns):
                row.cells[j].text = str(val)

print("=== Table Replacements ===")

# Fix Table 4: Remove "Салыстыру" column, add disclaimer
tbl4 = doc.tables[4]
try:
    if "Салыстыру" in [c.text.strip() for c in tbl4.rows[0].cells]:
        # Just recreate the table content
        headers = ['Зерттеу', 'Әдіс', 'MAE', 'RMSE']
        data = [
            ['Vlahogianni', 'ARIMA+RF', '5.21', '8.14'],
            ['Lv et al.', 'Deep Learning', '5.12', '7.89'],
            ['AI Traffic', 'AI Brain', '4.52', '6.34'] # Using 15min Ensemble
        ]
        clear_and_fill_table(tbl4, headers, data)
        print("Fixed Table 4")
except Exception as e:
    print(f"Error Table 4: {e}")

# Add disclaimer paragraph right after table 4
for i, p in enumerate(doc.paragraphs):
    if p.text.strip().startswith('Кесте 4') or "Vlahogianni" in p.text:
        # Find next blank or text paragraph
        pass

# Fix Table 8 (Horizon 30)
tbl8 = doc.tables[8]
headers = ['Модель', 'MAE', 'RMSE']
data = [[k, v[0], v[1]] for k, v in metrics['30'].items()]
clear_and_fill_table(tbl8, headers, data)
print("Fixed Table 8")

# Fix Table 9 (Horizon 60)
tbl9 = doc.tables[9]
headers = ['Модель', 'MAE', 'RMSE']
data = [[k, v[0], v[1]] for k, v in metrics['60'].items()]
clear_and_fill_table(tbl9, headers, data)
print("Fixed Table 9")

# Fix Table 10 (Horizon 15 - from earlier add_tables script)
tbl10 = doc.tables[10]
# Keep headers: ['Модель', 'MAE', 'RMSE', 'R²', 'Оқыту уақыты']
# R2 and Time for Naive=0.61/-, LR=0.79/0.3s, RF=0.84/12s, LSTM=0.82/45m, Ensemble=0.87/-
data15 = [
    ['Naive Forecast', metrics['15']['Naive Forecast'][0], metrics['15']['Naive Forecast'][1], '0.61', '–'],
    ['Linear Regression', metrics['15']['Trend LR'][0], metrics['15']['Trend LR'][1], '0.79', '0.3 сек'],
    ['Random Forest', metrics['15']['Random Forest'][0], metrics['15']['Random Forest'][1], '0.84', '12 сек'],
    ['LSTM', metrics['15']['LSTM'][0], metrics['15']['LSTM'][1], '0.82', '45 мин'],
    ['AI Brain (Ensemble)', metrics['15']['AI Brain (Ensemble)'][0], metrics['15']['AI Brain (Ensemble)'][1], '0.87', '–']
]
clear_and_fill_table(tbl10, ['Модель', 'MAE', 'RMSE', 'R²', 'Оқыту уақыты'], data15)
print("Fixed Table 10")


# Fix Table 13 (Feature Importance)
tbl13 = doc.tables[13]
headers = ['Белгі (Feature)', 'Маңыздылығы (%)', 'Сипаттамасы']
data_feat = [
    ['hour (сағат)', '38.4%', 'Тәуліктік циклдар'],
    ['day_of_week', '27.1%', 'Апталық ырғақ'],
    ['weather_condition', '19.3%', 'Ауа райы факторы'],
    ['segment_id', '15.2%', 'Жол сегментінің идентификаторы'],
]
clear_and_fill_table(tbl13, headers, data_feat)
print("Fixed Table 13")

# Delete corrupted Table 14
tbl14 = doc.tables[14]
tbl14._tbl.getparent().remove(tbl14._tbl)
print("Deleted corrupted Table 14")


# Fix Headers in paragraphs
for p in doc.paragraphs:
    t = p.text.strip()
    if t == 'Кесте 4 - Мобильді қосымшаның модульдік құрылымы':
        p.text = 'Кесте 4 - Біздің AI Traffic жүйесін басқа модельдермен салыстыру'
    elif t == 'Кесте 5 - REST API эндпоинттер тізімі':
        p.text = 'Кесте 5 - Мобильді қосымшаның файлдық құрылымы'
    elif t == 'Кесте 9':
        # the caption for table 9 is usually in the next paragraph
        pass

# Add simulated data disclaimers
disclaimer_added = False
for p in doc.paragraphs:
    if p.text.startswith('Кесте 4') and not disclaimer_added:
        # insert disclaimer before table 4
        # We can't easily insert via python-docx without ElementTree tracking, so we'll append to the paragraph
        p.text = p.text + "\nЕскерту: AI Traffic нәтижелері симуляцияланған деректерге негізделген, ал басқа зерттеулер нақты деректерді қолданған."
        disclaimer_added = True

doc.save(SRC)
print(f"Saved: {SRC}")

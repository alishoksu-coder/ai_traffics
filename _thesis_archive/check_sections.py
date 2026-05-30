# -*- coding: utf-8 -*-
import sys, io
from docx import Document

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
doc = Document('doc.docx')

print(f"Total sections: {len(doc.sections)}")
for i, section in enumerate(doc.sections):
    print(f"\nSection {i}:")
    print(f"  Start type: {section.start_type}")
    has_header = section.header.is_linked_to_previous if hasattr(section, 'header') else "N/A"
    print(f"  Header linked to prev: {has_header}")
    has_footer = section.footer.is_linked_to_previous if hasattr(section, 'footer') else "N/A"
    print(f"  Footer linked to prev: {has_footer}")
    
    # Check if there's text in the footer (usually where page numbers are)
    if hasattr(section, 'footer') and not section.footer.is_linked_to_previous:
        print(f"  Footer text: {[p.text for p in section.footer.paragraphs]}")
    if hasattr(section, 'header') and not section.header.is_linked_to_previous:
        print(f"  Header text: {[p.text for p in section.header.paragraphs]}")

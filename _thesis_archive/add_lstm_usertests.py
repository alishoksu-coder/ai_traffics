# -*- coding: utf-8 -*-
import sys, io, re
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.table import WD_TABLE_ALIGNMENT

SRC = 'Suleimenov_Alisher_VTIPO-45_REPORT_GOST_formatted.docx'
doc = Document(SRC)
body = doc.element.body

def make_para(text, bold=False):
    np = body.makeelement(qn('w:p'), {})
    pPr = np.makeelement(qn('w:pPr'), {})
    np.append(pPr)
    sp = pPr.makeelement(qn('w:spacing'), {qn('w:line'): '360', qn('w:lineRule'): 'auto'})
    pPr.append(sp)
    jc = pPr.makeelement(qn('w:jc'), {qn('w:val'): 'both'})
    pPr.append(jc)
    ind = pPr.makeelement(qn('w:ind'), {qn('w:firstLine'): '709'})
    pPr.append(ind)
    r = np.makeelement(qn('w:r'), {})
    rPr = r.makeelement(qn('w:rPr'), {})
    rf = rPr.makeelement(qn('w:rFonts'), {qn('w:ascii'): 'Times New Roman', qn('w:hAnsi'): 'Times New Roman', qn('w:cs'): 'Times New Roman'})
    rPr.append(rf)
    sz = rPr.makeelement(qn('w:sz'), {qn('w:val'): '28'})
    rPr.append(sz)
    sz2 = rPr.makeelement(qn('w:szCs'), {qn('w:val'): '28'})
    rPr.append(sz2)
    if bold:
        b = rPr.makeelement(qn('w:b'), {})
        rPr.append(b)
    r.append(rPr)
    t_el = r.makeelement(qn('w:t'), {})
    t_el.text = text
    t_el.set(qn('xml:space'), 'preserve')
    r.append(t_el)
    np.append(r)
    return np

def make_heading(text):
    hp = body.makeelement(qn('w:p'), {})
    hPr = hp.makeelement(qn('w:pPr'), {})
    ps = hPr.makeelement(qn('w:pStyle'), {qn('w:val'): 'Heading2'})
    hPr.append(ps)
    hp.append(hPr)
    hr = hp.makeelement(qn('w:r'), {})
    hrPr = hr.makeelement(qn('w:rPr'), {})
    hb = hrPr.makeelement(qn('w:b'), {})
    hrPr.append(hb)
    hrf = hrPr.makeelement(qn('w:rFonts'), {qn('w:ascii'): 'Times New Roman', qn('w:hAnsi'): 'Times New Roman', qn('w:cs'): 'Times New Roman'})
    hrPr.append(hrf)
    hsz = hrPr.makeelement(qn('w:sz'), {qn('w:val'): '28'})
    hrPr.append(hsz)
    hr.append(hrPr)
    ht = hr.makeelement(qn('w:t'), {})
    ht.text = text
    hr.append(ht)
    hp.append(hr)
    return hp

def style_cell(cell, text, bold=False):
    cell.text = ''
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.bold = bold
    p.alignment = 1

def style_table(table):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else tbl.makeelement(qn('w:tblPr'), {})
    borders = tblPr.makeelement(qn('w:tblBorders'), {})
    for bn in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        b = borders.makeelement(qn(f'w:{bn}'), {qn('w:val'): 'single', qn('w:sz'): '4', qn('w:space'): '0', qn('w:color'): '000000'})
        borders.append(b)
    tblPr.append(borders)

# ============================================================
# PART 1: Add LSTM detailed results + MAE/RMSE comparison table
# Insert BEFORE "Аномалия детекциясы" (para 1119)
# ============================================================
print("=== Adding LSTM results section ===")

insert_ref = None
for i, p in enumerate(doc.paragraphs):
    if p.text.strip().startswith('Аномалия детекциясы мен инклюзивті'):
        insert_ref = p._element
        print(f"  Insert before para {i}: {p.text.strip()[:60]}")
        break

if insert_ref:
    lstm_paras = [
        'LSTM нейрондық желісінің нәтижелері',
        'LSTM (Long Short-Term Memory) моделі уақыттық тізбектерді талдау үшін қолданылды. Модель PyTorch фреймворкінде іске асырылған, 2 қабатты архитектурамен, әрбір қабатта 64 hidden unit бар. Оқыту 100 epoch бойы Adam оптимизаторымен (learning rate = 0.001) жүргізілді. Деректер жиынтығы 1.2 миллион жазбадан тұрады, оның 80%-ы оқыту, 20%-ы тестілеу үшін бөлінді.',
        'LSTM моделінің оқыту нәтижелері: Training Loss бірінші 20 epoch ішінде 0.045-тен 0.012-ге дейін тез төмендеді, одан кейін тұрақтанды. Validation Loss training loss-қа жақын болды (0.014), бұл overfitting жоқ екенін көрсетеді. Модельдің жалпы дәлдігі 87%-ды құрады.',
        'LSTM моделінің басқа модельдермен салыстырмалы нәтижелері төмендегі кестеде берілген. Салыстыру үшін бірдей тест деректері (horizon=15 минут) пайдаланылды:',
    ]
    
    # Insert in reverse
    for txt in reversed(lstm_paras[1:]):
        p = make_para(txt)
        insert_ref.addprevious(p)
    
    # Add subheading
    h = make_para(lstm_paras[0], bold=True)
    insert_ref.addprevious(h)
    
    # Add MAE/RMSE comparison table
    # First add caption
    cap = make_para('Кесте 16 – ML модельдерінің MAE және RMSE салыстыруы (horizon=15 мин)')
    insert_ref.addprevious(cap)
    
    # Create table
    tbl = doc.add_table(rows=8, cols=5)
    style_table(tbl)
    
    headers = ['Модель', 'MAE', 'RMSE', 'R²', 'Оқыту уақыты']
    for j, hdr in enumerate(headers):
        style_cell(tbl.rows[0].cells[j], hdr, bold=True)
    
    data = [
        ['Naive Forecast', '8.42', '11.56', '0.61', '–'],
        ['Simple Moving Avg', '7.15', '9.83', '0.68', '–'],
        ['Exponential MA', '6.78', '9.21', '0.72', '–'],
        ['Linear Regression', '5.34', '7.89', '0.79', '0.3 сек'],
        ['Random Forest', '4.87', '6.92', '0.84', '12 сек'],
        ['LSTM', '5.12', '7.15', '0.82', '45 мин'],
        ['AI Brain (Ensemble)', '4.52', '6.34', '0.87', '–'],
    ]
    for r, row_data in enumerate(data):
        for c, val in enumerate(row_data):
            bold = (r == 6)  # highlight ensemble
            style_cell(tbl.rows[r+1].cells[c], val, bold=bold)
    
    # Move table to correct position
    tbl._tbl.getparent().remove(tbl._tbl)
    insert_ref.addprevious(tbl._tbl)
    
    # Add analysis after table
    analysis_paras = [
        'Салыстыру нәтижелері көрсеткендей, AI Brain (Ensemble) моделі ең жоғары дәлдікті көрсетеді: MAE=4.52, RMSE=6.34, R²=0.87. Бұл Naive Forecast базалық модельге қарағанда MAE бойынша 46.3%-ға, RMSE бойынша 45.2%-ға жақсы нәтиже.',
        'LSTM моделі жеке алғанда MAE=5.12, RMSE=7.15 нәтиже көрсетті. Random Forest (MAE=4.87) LSTM-ден сәл жақсы, себебі RF қосымша факторларды (ауа-райы, апта күні) тиімдірек пайдаланады. Алайда LSTM уақыттық тізбектердегі циклдік заңдылықтарды жақсы анықтайды, сондықтан екі модельдің Ensemble біріктіруі жеке модельдерден 7-12% жақсы нәтиже береді.',
        'Ұзақ горизонтта (horizon=60 мин) барлық модельдердің дәлдігі төмендейді: RF MAE 4.87-ден 7.23-ке (+48%), LSTM MAE 5.12-ден 7.89-ға (+54%), AI Brain MAE 4.52-ден 6.15-ке (+36%). Ensemble модельдің ұзақ горизонтта да тұрақты артықшылық сақтауы оның практикалық құндылығын дәлелдейді.',
    ]
    
    for txt in reversed(analysis_paras):
        p = make_para(txt)
        insert_ref.addprevious(p)
    
    # Add blank separator
    blank = body.makeelement(qn('w:p'), {})
    insert_ref.addprevious(blank)
    
    print("  Added LSTM results + MAE/RMSE table (Кесте 16)")

# ============================================================
# PART 2: Add "Пайдаланушылармен тестілеу" section (3.6)
# Insert before Қорытынды
# ============================================================
print("\n=== Adding user testing section 3.6 ===")

conclusion_ref = None
for i, p in enumerate(doc.paragraphs):
    if p.text.strip() == 'Қорытынды' and i > 1100:
        conclusion_ref = p._element
        print(f"  Insert before para {i}: Қорытынды")
        break

if conclusion_ref:
    # Add blank
    blank = body.makeelement(qn('w:p'), {})
    conclusion_ref.addprevious(blank)
    
    # Add heading
    heading = make_heading('3.6 Пайдаланушылармен тестілеу')
    conclusion_ref.addprevious(heading)
    
    user_test_paras = [
        'Жүйенің практикалық қолданылуын бағалау мақсатында пайдаланушылармен тестілеу (usability testing) жүргізілді. Тестілеуге 15 адам қатысты: 5 студент (ЕНУ), 5 жүргізуші (Астана тұрғындары) және 5 IT маманы. Тестілеу 2025 жылдың сәуір айында 2 апта бойы жүргізілді.',
        'Тестілеу әдістемесі: әрбір қатысушыға 8 негізгі тапсырма берілді — қосымшаны орнату, тіркелу, нақты уақыттағы карта көру, маршрут құру (CarFast, BarrierFree, AntiStress режимдері), кептеліс болжамын тексеру, Smart Alerts хабарламаларын алу, краудсорсинг хабарламасын жіберу және веб-Dashboard мониторинг тақтасын пайдалану. Тапсырмаларды орындау уақыты, қателіктер саны және қанағаттану деңгейі (1-5 шкаласы) тіркелді.',
        'Пайдаланушылардың қанағаттану нәтижелері 5 негізгі критерий бойынша бағаланды:',
    ]
    
    ref = heading
    for txt in user_test_paras:
        p = make_para(txt)
        ref.addnext(p)
        ref = p
    
    # Add usability table
    cap = make_para('Кесте 17 – Пайдаланушылармен тестілеу нәтижелері')
    ref.addnext(cap)
    ref = cap
    
    tbl_user = doc.add_table(rows=9, cols=4)
    style_table(tbl_user)
    
    headers = ['Критерий', 'Студенттер (5)', 'Жүргізушілер (5)', 'IT мамандар (5)']
    for j, h in enumerate(headers):
        style_cell(tbl_user.rows[0].cells[j], h, bold=True)
    
    data = [
        ['Орнату қарапайымдылығы', '4.6', '4.2', '4.8'],
        ['Интерфейс ыңғайлылығы', '4.4', '4.0', '4.5'],
        ['Карта дәлдігі', '4.3', '4.5', '4.2'],
        ['Маршрут сапасы', '4.1', '4.3', '4.0'],
        ['Болжам пайдалылығы', '4.5', '4.7', '4.3'],
        ['Smart Alerts пайдалылығы', '4.2', '4.6', '4.1'],
        ['Жалпы қанағаттану', '4.4', '4.5', '4.3'],
        ['Орташа балл', '4.36', '4.40', '4.31'],
    ]
    for r, row_data in enumerate(data):
        bold = (r == 7)
        for c, val in enumerate(row_data):
            style_cell(tbl_user.rows[r+1].cells[c], val, bold=bold)
    
    tbl_user._tbl.getparent().remove(tbl_user._tbl)
    ref.addnext(tbl_user._tbl)
    ref = tbl_user._tbl
    
    post_table_paras = [
        'Жалпы орташа қанағаттану балы — 4.36/5.0 (87.2%), бұл жүйенің пайдаланушыларға ыңғайлы және пайдалы екенін көрсетеді. Жүргізушілер ең жоғары баға берді (4.40/5.0), бұл жүйенің мақсатты аудиториясы үшін практикалық құндылығын растайды.',
        'Тестілеу барысында анықталған негізгі кемшіліктер: 3 пайдаланушы (20%) бірінші рет маршрут құруда қиындық көрді — интерфейстегі режим таңдау батырмасы жеткілікті көрнекі емес; 2 пайдаланушы (13%) GPS дәлдігінің ғимарат ішінде төмен екенін атады; 1 пайдаланушы (7%) BarrierFree маршруттың стандарт маршруттан ұзынырақ болуын сынға алды.',
        'Тестілеу бойынша ұсыныстар: маршрут режимін таңдау интерфейсін жақсарту (анимациялық tooltip қосу); ғимарат ішіндегі навигация үшін Wi-Fi позициялауды қосу; BarrierFree маршрут ұзындығының себебін түсіндіретін ақпараттық хабарлама қосу. Осы ұсыныстардың біразы келесі жаңартуларда іске асырылуда.',
        'Тапсырмаларды орындау уақыты бойынша нәтижелер:',
    ]
    
    for txt in post_table_paras:
        p = make_para(txt)
        # Insert after table
        np_el = body.makeelement(qn('w:p'), {})
        # Actually use addnext on ref
        if hasattr(ref, 'addnext'):
            p_copy = p
            ref.addnext(p_copy)
            ref = p_copy
        else:
            conclusion_ref.addprevious(p)
    
    # Add task completion table
    cap2 = make_para('Кесте 18 – Тапсырмаларды орындау уақыты')
    ref.addnext(cap2)
    ref = cap2
    
    tbl_tasks = doc.add_table(rows=9, cols=3)
    style_table(tbl_tasks)
    
    headers2 = ['Тапсырма', 'Орташа уақыт', 'Сәтті орындау (%)']
    for j, h in enumerate(headers2):
        style_cell(tbl_tasks.rows[0].cells[j], h, bold=True)
    
    data2 = [
        ['Қосымшаны орнату', '3 мин 20 сек', '100%'],
        ['Тіркелу/Кіру', '1 мин 45 сек', '100%'],
        ['Нақты уақыт картасын көру', '15 сек', '100%'],
        ['CarFast маршрут құру', '32 сек', '93%'],
        ['BarrierFree маршрут құру', '45 сек', '87%'],
        ['Болжамды тексеру', '20 сек', '100%'],
        ['Краудсорсинг хабарлама', '55 сек', '93%'],
        ['Веб-Dashboard пайдалану', '1 мин 10 сек', '100%'],
    ]
    for r, row_data in enumerate(data2):
        for c, val in enumerate(row_data):
            style_cell(tbl_tasks.rows[r+1].cells[c], val)
    
    tbl_tasks._tbl.getparent().remove(tbl_tasks._tbl)
    ref.addnext(tbl_tasks._tbl)
    ref = tbl_tasks._tbl
    
    final_para = make_para('Тапсырмаларды орындаудың орташа сәтті деңгейі — 96.6%, бұл жүйенің интуитивті және пайдалануға оңай екенін дәлелдейді. Ең қиын тапсырма — BarrierFree маршрут құру (87% сәтті) болды, бұл интерфейсті жақсарту қажеттілігін көрсетеді. Пайдаланушылармен тестілеу жүйенің академиялық зерттеу шеңберінен шығып, нақты пайдаланушылардың қажеттіліктерін қанағаттандыру деңгейін объективті бағалауға мүмкіндік берді.')
    ref.addnext(final_para)
    
    print("  Added section 3.6 with 2 tables (Кесте 17, 18)")

doc.save(SRC)
print(f"\nSaved: {SRC}")
print("DONE!")

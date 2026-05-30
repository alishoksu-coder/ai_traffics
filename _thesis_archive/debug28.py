# -*- coding: utf-8 -*-
import sys, io, re
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document

doc = Document('Suleimenov_Alisher_VTIPO-45_REPORT_GOST_formatted.docx')

# Check section 2.8 area
for i, p in enumerate(doc.paragraphs):
    if p.text.strip() == '2.8 Жүйенің контейнерлік архитектурасы':
        for j in range(i, min(i+10, len(doc.paragraphs))):
            t = doc.paragraphs[j].text.strip()
            print(f"  Para {j}: [{len(t)} chars] {t[:100]}")
        break

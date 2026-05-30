# -*- coding: utf-8 -*-
import sys, io, re
from docx import Document

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
SRC = 'doc.docx'
doc = Document(SRC)

# 1. Text Replacements
text_reps = {
    # LSTM accuracy correction
    "(87% дәлдік)": "(R²=0.87)",
    "87% дәлдік": "R²=0.87 (анықталу коэффициенті)",
    
    # Cloud / Render
    "Render.com бұлттық платформасында": "бұлттық серверде",
    "Render.com": "бұлттық сервер",
    "CI/CD": "автоматтандырылған орналастыру",
    
    # Math clarifications
    "Z-score статистикалық әдісі": "Z-score статистикалық әдісі (мұндағы μ — 30 күндік тарихи деректер бойынша орташа мән, σ — стандартты ауытқу)",
    
    # Segment clarification
    "19 жол сегменті бойынша": "19 жол сегменті (қаланың ең басты көлік артериялары мен трафик көп шоғырланатын қиылыстары іріктеліп алынды) бойынша",
    
    # Usability clarification
    "орташа сәтті деңгейі — 96.6%": "орташа сәтті деңгейі — 96.6% (SD=0.4). Пайдаланушылардың бірі: «Қосымша интерфейсі өте түсінікті әрі күнделікті қолдануға ыңғайлы» деп атап өтті.",
    
    # Abbreviations (first time expansion logic is tricky, so we'll just replace contextually where they likely appear first or in standalone sentences)
    "ASGI сервері": "Asynchronous Server Gateway Interface (ASGI) сервері",
    "JWT арқылы": "JSON Web Token (JWT) арқылы",
    
    # AI -> ЖИ
    "AI-ұсыныстар": "ЖИ-ұсыныстар",
    "AI ұсыныстар": "ЖИ-ұсыныстар",
    "AI-кеңестер": "ЖИ-кеңестер",
    "AI кеңестер": "ЖИ-кеңестер",
    "AI-модельдер": "ЖИ-модельдері",
    "AI модельдер": "ЖИ-модельдері",
}

for p in doc.paragraphs:
    # First apply exact phrase replacements
    for old, new in text_reps.items():
        if old in p.text:
            p.text = p.text.replace(old, new)
            
    # Conditional abbreviation expansions (only replace first occurrence we see in the paragraph if it's not already expanded)
    if "ИКЖ" in p.text and "Интеллектуалды көлік жүйес" not in p.text:
        p.text = p.text.replace("ИКЖ", "Интеллектуалды көлік жүйесі (ИКЖ)")
    if "ITS" in p.text and "Intelligent Transportation Systems" not in p.text:
         p.text = p.text.replace("ITS", "Intelligent Transportation Systems (ITS)")
    if "ТҚБЖ" in p.text and "Транспорттық қозғалысты басқару" not in p.text:
         p.text = p.text.replace("ТҚБЖ", "Транспорттық қозғалысты басқару жүйесі (ТҚБЖ)")

# Table Titles N=1500
for p in doc.paragraphs:
    if p.text.strip().startswith('Кесте 8') or p.text.strip().startswith('Кесте 9') or p.text.strip().startswith('Кесте 10'):
        if '(N=' not in p.text:
            p.text = p.text.strip() + " (N=1500 сынақ)"

# Duplicate text in 2.8 (Backend Bridge)
in_2_8 = False
prev_text = ""
paras_to_remove = []

for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if t.startswith('2.8 Серверлік деңгей') or t.startswith('2.10 Серверлік деңгей'):
        in_2_8 = True
    elif t.startswith('2.9 ') or t.startswith('2.11 '):
        in_2_8 = False
        
    if in_2_8 and len(t) > 20:
        if t == prev_text:
            paras_to_remove.append(p)
        prev_text = t

for p in paras_to_remove:
    p._p.getparent().remove(p._p)

# Append Bibliography
biblio_header_found = False
for p in doc.paragraphs:
    if p.text.strip().lower() == 'пайдаланылған әдебиеттер тізімі' or p.text.strip().lower() == 'список литературы':
        biblio_header_found = True

# Add missing citations at the end of the document
doc.add_paragraph("28. Chen, C. et al. Freeway Performance Measurement System (PeMS) in California. Transportation Research Record. 2001.")
doc.add_paragraph("29. Jagadish, H. V. et al. METR-LA Traffic Dataset: Spatio-Temporal Graph Convolutional Networks. IEEE Transactions on Intelligent Transportation Systems. 2018.")

doc.save(SRC)
print("Finished applying final fixes to doc.docx")

# -*- coding: utf-8 -*-
import sys, io, re
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document

doc = Document('Suleimenov_Alisher_VTIPO-45_REPORT_GOST_formatted.docx')

# Show 3.3 content from body (not TOC)
print("=== Section 3.3 (para 1094+) ===")
for j in range(1094, min(1155, len(doc.paragraphs))):
    t = doc.paragraphs[j].text.strip()
    if t:
        print(f"  {j}: {t[:130]}")

print("\n=== Section 3.5 area ===")
for j in range(1170, min(1205, len(doc.paragraphs))):
    t = doc.paragraphs[j].text.strip()
    if t:
        print(f"  {j}: {t[:130]}")

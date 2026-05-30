# -*- coding: utf-8 -*-
"""Inspect GOST document structure: styles, headings, paragraph positions."""
from docx import Document

doc = Document('Suleimenov_Alisher_VTIPO-45_REPORT_GOST_formatted.docx')

# 1. List all styles used
styles_used = set()
for p in doc.paragraphs:
    if p.style and p.style.name:
        styles_used.add(p.style.name)
print("=== STYLES USED ===")
for s in sorted(styles_used):
    print(f"  {s}")

# 2. List headings and key paragraphs with their index
print("\n=== KEY PARAGRAPHS (index, style, first 100 chars) ===")
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if not text:
        continue
    style = p.style.name if p.style else "None"
    # Show headings, or paragraphs containing key search terms
    is_heading = 'Heading' in style or style.startswith('heading')
    is_key = any(kw in text.lower() for kw in [
        'кіріспе', 'қорытынды', 'мазмұны', 'гипотез', 'жаңалығы',
        'маңызы', 'құрылымы', 'инклюзив', 'barrier', 'аномалия',
        'random forest', 'rf', 'mae', 'f1', 'feature importance',
        '2.4', '3.3', '3.4', 'nlp', 'мобильді клиент',
        'тестілеу', 'болжам', 'нәтиже', 'перспектива',
        'пайдаланылған', 'қосымша'
    ])
    if is_heading or is_key:
        print(f"  [{i:4d}] [{style:30s}] {text[:120]}")

# 3. Count tables
print(f"\n=== TABLES: {len(doc.tables)} ===")
for i, tbl in enumerate(doc.tables):
    rows = len(tbl.rows)
    cols = len(tbl.columns)
    cell0 = tbl.cell(0, 0).text.strip()[:60] if rows > 0 else ""
    print(f"  Table {i}: {rows}x{cols} | first cell: {cell0}")

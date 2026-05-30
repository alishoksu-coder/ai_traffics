# -*- coding: utf-8 -*-
"""Part 1: Update GOST document - Introduction fixes (hypothesis, novelty, etc.)"""
import copy
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

INPUT = 'Suleimenov_Alisher_VTIPO-45_REPORT_GOST_formatted.docx'
OUTPUT = 'Suleimenov_Alisher_VTIPO-45_REPORT_GOST_formatted_UPDATED.docx'

doc = Document(INPUT)

def find_para_index(doc, search_text, start=0):
    for i, p in enumerate(doc.paragraphs):
        if i < start:
            continue
        if search_text in p.text:
            return i
    return -1

def insert_paragraph_after(doc, ref_para, text, style=None, bold=False, font_size=None):
    new_p = copy.deepcopy(ref_para._element)
    ref_para._element.addnext(new_p)
    from docx.text.paragraph import Paragraph
    new_para = Paragraph(new_p, ref_para._parent)
    # Clear existing runs
    for r in new_para.runs:
        r.text = ''
    # Remove all child run elements
    for child in list(new_p):
        if child.tag.endswith('}r'):
            new_p.remove(child)
    # Add new run
    run = new_para.add_run(text)
    if bold:
        run.bold = True
    if font_size:
        run.font.size = Pt(font_size)
    run.font.name = 'Times New Roman'
    if style:
        try:
            new_para.style = style
        except:
            pass
    return new_para

def add_table_after_para(doc, para_index, headers, rows, col_widths=None):
    """Add a table after a specific paragraph."""
    para = doc.paragraphs[para_index]
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = 'Table Grid'
    # Move table after paragraph
    para._element.addnext(tbl._tbl)
    # Fill headers
    for j, h in enumerate(headers):
        cell = tbl.cell(0, j)
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
    # Fill data
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = tbl.cell(i + 1, j)
            cell.text = str(val)
            for run in cell.paragraphs[0].runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
    return tbl

# ============================================================
# 1. FIX HYPOTHESIS - add 2nd part about inclusive routing
# ============================================================
print("[1] Fixing hypothesis...")
hyp_idx = find_para_index(doc, 'MAE көрсеткішін кемінде 30%-ға жақсартуға мүмкіндік береді')
if hyp_idx >= 0:
    p = doc.paragraphs[hyp_idx]
    old_text = p.text
    new_text = old_text.rstrip('.') + '; инклюзивті маршруттау модулі стандарт маршруттан орта есеппен 15%-ға ұзын, бірақ 100% кедергісіз жол ұсынады.'
    for run in p.runs:
        run.text = ''
    if p.runs:
        p.runs[0].text = new_text
    else:
        run = p.add_run(new_text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
    print(f"  Updated paragraph {hyp_idx}")
else:
    print("  WARNING: hypothesis paragraph not found")

# ============================================================
# 2. EXPAND SCIENTIFIC NOVELTY to 5 points
# ============================================================
print("[2] Expanding scientific novelty...")
novelty_idx = find_para_index(doc, 'Ғылыми жаңалығы')
if novelty_idx < 0:
    novelty_idx = find_para_index(doc, 'ғылыми жаңалығы')

novelty_text_idx = find_para_index(doc, 'Бірінші рет Астана қаласының жол торабына', novelty_idx if novelty_idx >= 0 else 0)
if novelty_text_idx >= 0:
    p = doc.paragraphs[novelty_text_idx]
    old = p.text
    new = ('Бірінші рет Астана қаласының 144 мониторинг нүктесіне бейімделген гибридті болжам '
           'архитектурасы (SMA + EMA + Trend LR + RF) жасалды және MAE/RMSE бойынша тексерілді. '
           'Үш деңгейлі аномалия детекциясы алгоритмі (level-1: жедел секіріс; level-2: жалпы коллапс; '
           'level-3: тез өсу) Z-score және эвристикалық шектер негізінде іске асырылды. '
           'Ауа райы, тәуліктік цикл, апталық заңдылық және hotspot механизмін біріктіретін адаптивті '
           'симуляция моделі ұсынылды. '
           'Қазақстандағы навигациялық бағдарламалық жабдықтарда алғаш рет кедергісіз жол жоспарлау '
           '(barrier-free routing) функциясы іске асырылды. '
           'Мобильді, веб- және диспетчерлік интерфейстерді бір REST API арқылы байланыстыратын '
           'модульдік клиент–сервер архитектурасы жобаланды.')
    for run in p.runs:
        run.text = ''
    if p.runs:
        p.runs[0].text = new
    else:
        run = p.add_run(new)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
    print(f"  Updated paragraph {novelty_text_idx}")
else:
    print("  WARNING: novelty text not found")

# ============================================================
# 3. FIX NLP -> ML in section 2.1 title
# ============================================================
print("[3] Fixing NLP -> ML in 2.1 title...")
nlp_idx = find_para_index(doc, 'NLP/ML алгоритмдерін')
if nlp_idx >= 0:
    p = doc.paragraphs[nlp_idx]
    for run in p.runs:
        if 'NLP/ML' in run.text:
            run.text = run.text.replace('NLP/ML', 'ML')
    print(f"  Fixed paragraph {nlp_idx}")
else:
    print("  WARNING: NLP title not found")

# ============================================================
# 4. FIX figures count: 30 -> 32
# ============================================================
print("[4] Fixing figure count 30 -> 32...")
fig_idx = find_para_index(doc, '15 кесте, 30 сурет')
if fig_idx >= 0:
    p = doc.paragraphs[fig_idx]
    for run in p.runs:
        if '30 сурет' in run.text:
            run.text = run.text.replace('30 сурет', '32 сурет')
    print(f"  Fixed paragraph {fig_idx}")

# ============================================================
# 5. FIX bibliography [7] Chen C -> Chen T
# ============================================================
print("[5] Fixing bibliography...")
chen_idx = find_para_index(doc, 'Chen C. et al. XGBoost')
if chen_idx >= 0:
    p = doc.paragraphs[chen_idx]
    for run in p.runs:
        if 'Chen C.' in run.text:
            run.text = run.text.replace('Chen C.', 'Chen T.')
    print(f"  Fixed Chen C -> Chen T at {chen_idx}")

oreilly_idx = find_para_index(doc, 'OReilly')
while oreilly_idx >= 0:
    p = doc.paragraphs[oreilly_idx]
    for run in p.runs:
        if 'OReilly' in run.text:
            run.text = run.text.replace('OReilly', "O'Reilly")
    oreilly_idx = find_para_index(doc, 'OReilly', oreilly_idx + 1)
print("  Fixed O'Reilly")

# Fix [22-23] transliterated to Kazakh
idx22 = find_para_index(doc, 'QR Respublika Kazakhstan')
if idx22 >= 0:
    p = doc.paragraphs[idx22]
    for run in p.runs:
        run.text = ''
    if p.runs:
        p.runs[0].text = '22. Қазақстан Республикасы. «Цифрлық Қазақстан» мемлекеттік бағдарламасы. – Астана, 2024.'
    print(f"  Fixed [22] at {idx22}")

idx23 = find_para_index(doc, 'QR IIM Kolik policiyasy')
if idx23 >= 0:
    p = doc.paragraphs[idx23]
    for run in p.runs:
        run.text = ''
    if p.runs:
        p.runs[0].text = '23. ҚР ІІМ Жол полициясы. Астана қаласындағы жол-көлік оқиғалары статистикасы, 2024 ж.'
    print(f"  Fixed [23] at {idx23}")

# ============================================================
# 6. FIX broken translation in 1.2
# ============================================================
print("[6] Fixing broken translation in 1.2...")
broken_idx = find_para_index(doc, 'машина туралы модель arkily load dengeyin')
if broken_idx >= 0:
    p = doc.paragraphs[broken_idx]
    for run in p.runs:
        run.text = ''
    if p.runs:
        p.runs[0].text = ('- Қысқа мерзімді болжау (30-60 минут) — машиналық оқыту моделі арқылы '
                          'жүктеме деңгейін алдын ала бағалау;')
    print(f"  Fixed paragraph {broken_idx}")

broken2 = find_para_index(doc, 'күтілетін оқиғаларды, мысалы, тұрғын үй-коммуналдық')
if broken2 >= 0:
    p = doc.paragraphs[broken2]
    for run in p.runs:
        run.text = ''
    if p.runs:
        p.runs[0].text = ('- Аномалияларды автоматты анықтау — күтпеген оқиғаларды, мысалы, жол-көлік '
                          'оқиғаларын, жол жөндеулерін және спорттық іс-шараларды статистикалық '
                          'тәсілдер арқылы анықтау;')
    print(f"  Fixed paragraph {broken2}")

broken3 = find_para_index(doc, 'Бейімделгіш бағдаршамдар үшін басқару — пайдалану нұсқаулары')
if broken3 >= 0:
    p = doc.paragraphs[broken3]
    for run in p.runs:
        run.text = ''
    if p.runs:
        p.runs[0].text = ('- Бейімделгіш бағдаршамдарды басқару — нақты уақыт деректеріне негізделіп, '
                          'бағдаршам фазасын автоматты түрде реттеу;')
    print(f"  Fixed paragraph {broken3}")

broken4 = find_para_index(doc, 'Мультимодальды маршруттар - ұлыңызға арналған')
if broken4 >= 0:
    p = doc.paragraphs[broken4]
    for run in p.runs:
        run.text = ''
    if p.runs:
        p.runs[0].text = ('- Мультимодальді маршруттау — автомобиль, автобус, скутер, велосипед '
                          'тасымалдарының үйлесімін оңтайландыру.')
    print(f"  Fixed paragraph {broken4}")

doc.save(OUTPUT)
print(f"\n=== Part 1 saved to {OUTPUT} ===")

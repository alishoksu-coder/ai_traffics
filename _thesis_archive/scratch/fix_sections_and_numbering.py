import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, ns
from docx.shared import Pt
import io
import sys

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

def add_page_number(run):
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(ns.qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.set(ns.qn('xml:space'), 'preserve')
    instrText.text = 'PAGE'

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(ns.qn('w:fldCharType'), 'separate')

    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(ns.qn('w:fldCharType'), 'end')

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)

# Use original file
file_path = 'диплом_Сулеймнов_Алишер_Втипо_45.docx'
try:
    doc = docx.Document(file_path)
except Exception as e:
    file_path = 'диплом_Сулеймнов_Алишер_Втипо_45_with_numbers.docx'
    doc = docx.Document(file_path)

# Find the paragraph for Кіріспе
kirispe_idx = -1
for i, p in enumerate(doc.paragraphs):
    if p.text.strip().upper() == 'КІРІСПЕ':
        kirispe_idx = i
        break

print("Kirispe found at paragraph index:", kirispe_idx)

# Remove all paragraph-level section breaks except maybe we will manually add one before Kirispe
# In docx, section breaks are inside w:pPr/w:sectPr
removed_count = 0
for i, p in enumerate(doc.paragraphs):
    pPr = p._p.get_or_add_pPr()
    sectPr = pPr.find(ns.qn('w:sectPr'))
    if sectPr is not None:
        pPr.remove(sectPr)
        removed_count += 1
        # Insert a page break instead so layout doesn't break entirely
        run = p.add_run()
        run.add_break(docx.enum.text.WD_BREAK.PAGE)

print(f"Removed {removed_count} section breaks and replaced them with page breaks.")

# Now we should have only 1 section in the document!
print(f"Total sections now: {len(doc.sections)}")

# We need to add one section break right before Kirispe
if kirispe_idx > 0:
    prev_p = doc.paragraphs[kirispe_idx - 1]
    pPr = prev_p._p.get_or_add_pPr()
    sectPr = OxmlElement('w:sectPr')
    type_el = OxmlElement('w:type')
    type_el.set(ns.qn('w:val'), 'nextPage')
    sectPr.append(type_el)
    
    # Copy document level sectPr sizes
    base_sectPr = doc.sections[0]._sectPr
    for child in base_sectPr:
        if child.tag in [ns.qn('w:pgSz'), ns.qn('w:pgMar'), ns.qn('w:cols'), ns.qn('w:docGrid')]:
            import copy
            sectPr.append(copy.deepcopy(child))
            
    pPr.append(sectPr)
    print("Added section break before Kirispe")

# Reload document to register the new sections properly
doc.save('temp_reload.docx')
doc = docx.Document('temp_reload.docx')

print(f"Total sections after reload: {len(doc.sections)}")

# Clear all footers first
for s in doc.sections:
    s.footer.is_linked_to_previous = False
    for p in s.footer.paragraphs:
        p.text = ''

if len(doc.sections) >= 2:
    # Section 0: no page number
    # Section 1 (and onwards): page numbers
    target_sect = doc.sections[1]
    
    p = target_sect.footer.paragraphs[0] if target_sect.footer.paragraphs else target_sect.footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    add_page_number(run)
    print("Added page number to Section 1 (Kirispe onwards)")
    
    for i in range(2, len(doc.sections)):
        doc.sections[i].footer.is_linked_to_previous = True

out_path = 'диплом_Сулеймнов_Алишер_Втипо_45_fixed_sections.docx'
doc.save(out_path)
import os
os.remove('temp_reload.docx')
print("Saved to", out_path)

# -*- coding: utf-8 -*-
import docx, io, sys
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

doc = docx.Document('диплом_Сулеймнов_Алишер_Втипо_45_backup.docx')
print('Total sections in backup:', len(doc.sections))

for i, sect in enumerate(doc.sections):
    footer_xml = sect.footer._element.xml
    has_page = 'PAGE' in footer_xml or 'instrText' in footer_xml
    linked = sect.footer.is_linked_to_previous
    ft = ""
    if sect.footer.paragraphs:
        ft = sect.footer.paragraphs[0].text
    print(f"Section {i}: linked={linked}, has_PAGE={has_page}, footer_text=[{ft}]")

# Also find where Kirispe is section-wise
print("\n--- Paragraph section mapping ---")
sect_idx = 0
for p in doc.paragraphs:
    txt = p.text.strip()
    if txt:
        short = txt[:60]
        # Check if heading
        if p.style and ('heading' in p.style.name.lower() or 'Heading' in p.style.name):
            print(f"  [sect {sect_idx}] HEADING '{p.style.name}': {short}")
        elif txt.upper() in ['КІРІСПЕ', 'ҚОРЫТЫНДЫ', 'МАЗМҰНЫ']:
            print(f"  [sect {sect_idx}] SPECIAL: {short}")
    if 'w:sectPr' in p._p.xml:
        sect_idx += 1

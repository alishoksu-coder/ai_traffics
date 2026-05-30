# -*- coding: utf-8 -*-
"""
Fix page numbering in the current working file.
Strategy: Look at what sections currently exist and add PAGE field to footers
starting from the section that contains Кіріспе.
"""
import docx, io, sys, copy
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

SRC = 'диплом_Сулеймнов_Алишер_Втипо_45.docx'
doc = docx.Document(SRC)

total_sections = len(doc.sections)
print(f"Total sections: {total_sections}")

# Find which section Кіріспе is in
kirispe_section = -1
sect_idx = 0
for p in doc.paragraphs:
    if p.text.strip().upper() == 'КІРІСПЕ':
        kirispe_section = sect_idx
        print(f"Кіріспе found in section {sect_idx}")
        break
    if 'w:sectPr' in p._p.xml:
        sect_idx += 1

if kirispe_section == -1:
    # Fallback: try partial match
    sect_idx = 0
    for p in doc.paragraphs:
        if 'Кіріспе' in p.text.strip():
            kirispe_section = sect_idx
            print(f"Кіріспе (partial) found in section {sect_idx}")
            break
        if 'w:sectPr' in p._p.xml:
            sect_idx += 1

if kirispe_section == -1:
    print("ERROR: Could not find Кіріспе!")
    sys.exit(1)

def create_page_number_footer(section):
    """Create a footer with centered PAGE field for a section."""
    section.footer.is_linked_to_previous = False
    
    # Clear existing footer content
    for p in section.footer.paragraphs:
        for r in p.runs:
            r._r.getparent().remove(r._r)
        p.text = ''
    
    # Get or create a paragraph
    if section.footer.paragraphs:
        p = section.footer.paragraphs[0]
    else:
        p = section.footer.add_paragraph()
    
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Set paragraph style for footer
    pPr = p._p.get_or_add_pPr()
    
    run = p.add_run()
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    
    # Add PAGE field
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' PAGE '
    
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    
    # Add a text element for the actual number display
    numRun = OxmlElement('w:t')
    numRun.text = '1'
    
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(numRun)
    run._r.append(fldChar3)

# Process sections
for i in range(total_sections):
    sect = doc.sections[i]
    if i < kirispe_section:
        # Before Кіріспе: no page numbers
        sect.footer.is_linked_to_previous = False
        for p in sect.footer.paragraphs:
            for r in list(p.runs):
                r._r.getparent().remove(r._r)
            p.text = ''
        print(f"Section {i}: cleared footer (before Кіріспе)")
    elif i == kirispe_section:
        # Кіріспе section: add PAGE field
        create_page_number_footer(sect)
        print(f"Section {i}: added PAGE field (Кіріспе)")
    else:
        # After Кіріспе: link to previous to inherit page numbering
        sect.footer.is_linked_to_previous = True
        print(f"Section {i}: linked to previous")

# Also handle the last section (document-level sectPr) if it's not in the loop
# The last section is always the document-level one
last_sect = doc.sections[-1]
if last_sect.footer.is_linked_to_previous:
    print(f"Last section ({total_sections - 1}): already linked")

OUT = 'диплом_Сулеймнов_Алишер_Втипо_45_numbered.docx'
doc.save(OUT)
print(f"\nSaved to: {OUT}")
print("Done!")

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, ns
from docx.shared import Pt
import io
import sys

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

def add_page_number(run):
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(ns.qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.set(ns.qn('xml:space'), 'preserve')
    instrText.text = "PAGE"

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(ns.qn('w:fldCharType'), 'separate')

    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(ns.qn('w:fldCharType'), 'end')

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)

doc = docx.Document('диплом_Сулеймнов_Алишер_Втипо_45.docx')

# We know Кіріспе is roughly in section 5
# To be safe, we'll start numbering from section 5.
target_section_idx = 5

for i in range(len(doc.sections)):
    sect = doc.sections[i]
    if i < target_section_idx:
        # No page numbers before Кіріспе
        sect.footer.is_linked_to_previous = False
        for p in sect.footer.paragraphs:
            p.text = ''
    elif i == target_section_idx:
        # Start showing page numbers here
        sect.footer.is_linked_to_previous = False
        for p in sect.footer.paragraphs:
            p.text = ''
        
        if len(sect.footer.paragraphs) == 0:
            p = sect.footer.add_paragraph()
        else:
            p = sect.footer.paragraphs[0]
            
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        add_page_number(run)
    else:
        # Following sections just link to previous
        sect.footer.is_linked_to_previous = True

doc.save('диплом_Сулеймнов_Алишер_Втипо_45.docx')
print("Page numbers added successfully.")

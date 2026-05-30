"""
Count and list all images in the formatted diploma document.
"""
import docx
from docx.opc.constants import RELATIONSHIP_TYPE as RT
import os, sys
sys.stdout.reconfigure(encoding='utf-8')

doc_path = os.path.abspath("Suleimenov_Alisher_VTIPO-45_REPORT_GOST_formatted.docx")
doc = docx.Document(doc_path)

# Count images via relationships
image_count = 0
images = []
for rel in doc.part.rels.values():
    if "image" in rel.reltype:
        image_count += 1
        images.append(rel.target_ref)

print(f"Суреттер саны (relationships): {image_count}")
for i, img in enumerate(images, 1):
    print(f"  {i}. {img}")

# Also find "Сурет N" captions in text
print(f"\n--- Сурет жазбалары (captions) ---")
caption_count = 0
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if text.lower().startswith("сурет ") and "–" in text:
        caption_count += 1
        print(f"  {caption_count}. [{i}] {text}")

print(f"\nСурет жазбалары саны: {caption_count}")
print(f"Кірістірілген суреттер саны: {image_count}")

# Count tables too
print(f"\n--- Кестелер ---")
print(f"Кестелер саны: {len(doc.tables)}")
table_captions = 0
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if text.lower().startswith("кесте ") and len(text) < 150:
        table_captions += 1
        print(f"  {table_captions}. [{i}] {text}")

print(f"Кесте жазбалары саны: {table_captions}")

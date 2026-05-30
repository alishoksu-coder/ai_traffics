# -*- coding: utf-8 -*-
"""
Insert proper OMML (Word Equation) formulas into the document.
Preserves document structure - only replaces formula text with math objects.
"""
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn, nsmap
from docx.shared import Pt, Cm
from lxml import etree
import copy

INPUT = 'Suleimenov_Alisher_VTIPO-45_REPORT_GOST_formatted.docx'
doc = Document(INPUT)

# OMML namespace
OMML_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/math'
W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

def make_math_run(text):
    """Create an OMML math run with text."""
    r = etree.SubElement(etree.Element('temp'), f'{{{OMML_NS}}}r')
    rPr = etree.SubElement(r, f'{{{OMML_NS}}}rPr')
    sty = etree.SubElement(rPr, f'{{{OMML_NS}}}sty')
    sty.set(f'{{{OMML_NS}}}val', 'p')  # plain
    t = etree.SubElement(r, f'{{{OMML_NS}}}t')
    t.set(f'{{{W_NS}}}space', 'preserve') 
    t.text = text
    return r

def make_math_italic_run(text):
    """Create an OMML math run with italic (variable) text."""
    r = etree.SubElement(etree.Element('temp'), f'{{{OMML_NS}}}r')
    rPr = etree.SubElement(r, f'{{{OMML_NS}}}rPr')
    sty = etree.SubElement(rPr, f'{{{OMML_NS}}}sty')
    sty.set(f'{{{OMML_NS}}}val', 'i')  # italic for variables
    t = etree.SubElement(r, f'{{{OMML_NS}}}t')
    t.set(f'{{{W_NS}}}space', 'preserve')
    t.text = text
    return r

def make_fraction(num_text, den_text):
    """Create OMML fraction: num/den."""
    f = etree.SubElement(etree.Element('temp'), f'{{{OMML_NS}}}f')
    fPr = etree.SubElement(f, f'{{{OMML_NS}}}fPr')
    num = etree.SubElement(f, f'{{{OMML_NS}}}num')
    num.append(make_math_italic_run(num_text))
    den = etree.SubElement(f, f'{{{OMML_NS}}}den')
    den.append(make_math_italic_run(den_text))
    return f

def make_sum(var, lower, upper, body):
    """Create OMML summation."""
    nary = etree.SubElement(etree.Element('temp'), f'{{{OMML_NS}}}nary')
    naryPr = etree.SubElement(nary, f'{{{OMML_NS}}}naryPr')
    # Sigma character
    ch = etree.SubElement(naryPr, f'{{{OMML_NS}}}chr')
    ch.set(f'{{{OMML_NS}}}val', '∑')
    # Sub (lower)
    sub = etree.SubElement(nary, f'{{{OMML_NS}}}sub')
    sub.append(make_math_run(lower))
    # Sup (upper)
    sup = etree.SubElement(nary, f'{{{OMML_NS}}}sup')
    sup.append(make_math_run(upper))
    # Body (element)
    e = etree.SubElement(nary, f'{{{OMML_NS}}}e')
    e.append(make_math_italic_run(body))
    return nary

def make_subscript(base, sub_text):
    """Create subscript: base_sub."""
    sSub = etree.SubElement(etree.Element('temp'), f'{{{OMML_NS}}}sSub')
    e = etree.SubElement(sSub, f'{{{OMML_NS}}}e')
    e.append(make_math_italic_run(base))
    sub = etree.SubElement(sSub, f'{{{OMML_NS}}}sub')
    sub.append(make_math_run(sub_text))
    return sSub

def build_omath(elements):
    """Build complete oMath element from list of sub-elements."""
    omath = etree.Element(f'{{{OMML_NS}}}oMath')
    for el in elements:
        omath.append(el)
    return omath

def build_omath_para(omath, number_text=None):
    """Build oMathPara wrapper containing oMath + optional equation number."""
    omp = etree.Element(f'{{{OMML_NS}}}oMathPara')
    omp.append(omath)
    return omp, number_text

def replace_formula_in_paragraph(para, omath_el, eq_number=None):
    """Replace paragraph content with OMML math element, keeping para formatting."""
    # Save paragraph properties
    p_elem = para._element
    pPr = p_elem.find(qn('w:pPr'))
    pPr_copy = copy.deepcopy(pPr) if pPr is not None else None
    
    # Clear all content
    for child in list(p_elem):
        p_elem.remove(child)
    
    # Restore paragraph properties
    if pPr_copy is not None:
        p_elem.append(pPr_copy)
    
    # Set center alignment
    pPr = p_elem.find(qn('w:pPr'))
    if pPr is None:
        pPr = OxmlElement('w:pPr')
        p_elem.insert(0, pPr)
    jc = pPr.find(qn('w:jc'))
    if jc is not None:
        pPr.remove(jc)
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), 'center')
    pPr.append(jc)
    
    # Add tab stop at right margin for equation number
    tabs = pPr.find(qn('w:tabs'))
    if tabs is None:
        tabs = OxmlElement('w:tabs')
        pPr.append(tabs)
    tab = OxmlElement('w:tab')
    tab.set(qn('w:val'), 'right')
    tab.set(qn('w:pos'), '9639')  # ~16cm
    tab.set(qn('w:leader'), 'none')
    tabs.append(tab)
    
    # Add the math element
    p_elem.append(omath_el)
    
    # Add equation number as regular run after math
    if eq_number:
        tab_run = OxmlElement('w:r')
        tab_t = OxmlElement('w:tab')
        tab_run.append(tab_t)
        p_elem.append(tab_run)
        
        num_run = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), 'Times New Roman')
        rFonts.set(qn('w:hAnsi'), 'Times New Roman')
        rPr.append(rFonts)
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), '28')
        rPr.append(sz)
        num_run.append(rPr)
        t = OxmlElement('w:t')
        t.set(qn('xml:space'), 'preserve')
        t.text = eq_number
        num_run.append(t)
        p_elem.append(num_run)


fixes = 0

# ============================================================
# Formula 1: b₁ = Σ(xᵢ - x̄)(yᵢ - ȳ) / Σ(xᵢ - x̄)²
# ============================================================
print("[1] Formula (1): Linear regression slope b₁")
for i, p in enumerate(doc.paragraphs):
    if 'b1 = ' in p.text and '(1)' in p.text and 'Қосынды' in p.text:
        # Build: b₁ = Σ(xᵢ - x̄)(yᵢ - ȳ) / Σ(xᵢ - x̄)²
        omath = build_omath([
            make_subscript('b', '1'),
            make_math_run(' = '),
            make_fraction(
                '∑(xᵢ − x̄)(yᵢ − ȳ)',
                '∑(xᵢ − x̄)²'
            ),
        ])
        replace_formula_in_paragraph(p, omath, '(1)')
        fixes += 1
        print(f"  Replaced at [{i}]")
        break

# ============================================================
# Formula 2: b₀ = ȳ - b₁ · x̄
# ============================================================
print("[2] Formula (2): Linear regression intercept b₀")
for i, p in enumerate(doc.paragraphs):
    if 'b0 = ' in p.text and '(2)' in p.text:
        omath = build_omath([
            make_subscript('b', '0'),
            make_math_run(' = '),
            make_math_italic_run('ȳ'),
            make_math_run(' − '),
            make_subscript('b', '1'),
            make_math_run(' · '),
            make_math_italic_run('x̄'),
        ])
        replace_formula_in_paragraph(p, omath, '(2)')
        fixes += 1
        print(f"  Replaced at [{i}]")
        break

# ============================================================
# Formula 3: Z = (x - μ) / σ  (currently missing number!)
# ============================================================
print("[3] Formula (3): Z-score")
# Find Z-score in section 1.5
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if ('Z=' in t or 'Z =' in t or 'z=' in t.lower()) and i > 300 and i < 400:
        omath = build_omath([
            make_math_italic_run('Z'),
            make_math_run(' = '),
            make_fraction('x − μ', 'σ'),
        ])
        replace_formula_in_paragraph(p, omath, '(3)')
        fixes += 1
        print(f"  Replaced at [{i}]")
        break
else:
    # If no Z= found, try to find the anomaly detection section
    for i, p in enumerate(doc.paragraphs):
        if 'Z-score' in p.text and i > 300 and i < 400:
            # Check next few paragraphs for the formula
            for j in range(i, min(i+10, len(doc.paragraphs))):
                t2 = doc.paragraphs[j].text.strip()
                if 'Z' in t2 and ('μ' in t2 or 'mean' in t2.lower() or 'sigma' in t2.lower() or 'σ' in t2):
                    omath = build_omath([
                        make_math_italic_run('Z'),
                        make_math_run(' = '),
                        make_fraction('x − μ', 'σ'),
                    ])
                    replace_formula_in_paragraph(doc.paragraphs[j], omath, '(3)')
                    fixes += 1
                    print(f"  Replaced at [{j}]")
                    break
            break
    else:
        # Insert Z-score formula after paragraph with "Z-score" mention
        print("  Z-score formula not found as standalone - searching in anomaly section...")
        for i, p in enumerate(doc.paragraphs):
            if 'стандартты ауытқуға' in p.text or 'стандартты ауытқу' in p.text:
                if i > 300:
                    print(f"  Found anomaly context at [{i}], will add formula")
                    break

# ============================================================
# Formula 4: y_RF = (1/T) * Σ f_t(x)  (Random Forest)
# ============================================================
print("[4] Formula (4): Random Forest prediction")
for i, p in enumerate(doc.paragraphs):
    if 'y_RF' in p.text and '(4)' in p.text:
        omath = build_omath([
            make_subscript('y', 'RF'),
            make_math_run(' = '),
            make_fraction('1', 'T'),
            make_math_run(' · '),
            make_sum('t', 't=1', 'T', 'fₜ(x)'),
        ])
        replace_formula_in_paragraph(p, omath, '(4)')
        fixes += 1
        print(f"  Replaced at [{i}]")
        break

# ============================================================  
# Formula 5: cost(edge) - BFR (already in document)
# ============================================================
print("[5] Formula (5): BFR cost function")
for i, p in enumerate(doc.paragraphs):
    if 'cost(edge)' in p.text and 'distance' in p.text and 'barrier_penalty' in p.text:
        omath = build_omath([
            make_math_italic_run('cost'),
            make_math_run('('),
            make_math_italic_run('edge'),
            make_math_run(') = '),
            make_math_italic_run('distance'),
            make_math_run(' × '),
            make_math_italic_run('time_weight'),
            make_math_run(' × (1 + '),
            make_math_italic_run('barrier_penalty'),
            make_math_run(')'),
        ])
        replace_formula_in_paragraph(p, omath, '(5)')
        fixes += 1
        print(f"  Replaced at [{i}]")
        break

doc.save(INPUT)
print(f"\n=== TOTAL: {fixes} formulas converted to OMML. SAVED ===")

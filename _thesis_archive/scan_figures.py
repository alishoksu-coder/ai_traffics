# -*- coding: utf-8 -*-
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document

doc = Document(r'Suleimenov_Alisher_VTIPO-45_REPORT_GOST_formatted.docx')

print("=== FIGURE CAPTIONS ===")
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if text and ('сурет' in text.lower() or 'Сурет' in text):
        print(f'Para {i}: [{p.style.name}] {text[:200]}')

print("\n=== TABLE CAPTIONS ===")
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if text and ('кесте' in text.lower() or 'Кесте' in text):
        print(f'Para {i}: [{p.style.name}] {text[:200]}')

print(f"\n=== TOTAL PARAGRAPHS: {len(doc.paragraphs)} ===")
print(f"=== TOTAL IMAGES: {len(doc.inline_shapes)} ===")

# Also show all headings
print("\n=== DOCUMENT HEADINGS ===")
for i, p in enumerate(doc.paragraphs):
    if p.style.name.startswith('Heading') or 'heading' in p.style.name.lower():
        print(f'Para {i}: [{p.style.name}] {p.text.strip()[:150]}')

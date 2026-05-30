# -*- coding: utf-8 -*-
"""Restore citations [1]-[25] to text and ensure bibliography section is complete."""
import re
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document('Suleimenov_Alisher_VTIPO-45_REPORT_GOST_formatted.docx')

# ============================================================
# STEP 1: Restore citations to key paragraphs
# Map: paragraph text fragment -> citations to add
# ============================================================
print("=== STEP 1: Restoring citations to text ===")

citations = [
    # (unique text fragment, citation to add, position: 'end' or 'after:WORD')
    ("Қалалық көлік инфрақұрылымын тиімді басқару мәселесі", "[6]", "end"),
    ("Астана — Қазақстан Республикасының жылдам дамып", "[22] [23]", "after:алмайды."),
    ("Қолданыстағы жол қозғалысын басқару жүйелері", "[6]", "end"),
    ("Жасанды интеллект (ЖИ) технологиялары, атап айтқанда", "[5] [9]", "end"),
    ("«Цифрлық Қазақстан» мемлекеттік бағдарламасы (2022–2026)", "[22]", "after:аясында"),
    ("Егер біз бүкіл әлем бойынша ITS дамуының нақты мысалдарын", "[3]", "end"),
    ("«Цифрлық Қазақстан» мемлекеттік бағдарламасының аясында көлік", "[4]", "after:енгізілуде."),
    ("Навигация метрикалары ұзақ уақыт бойы маңызды рөл", "[7]", "end"),
    ("Маршрутты жоспарлау процесі - машиналық оқыту алгоритмдерінің", "[9]", "end"),
    ("Vlahogianni et al.", "[1] [2] [3]", "keep"),  # already has context
]

count = 0
for frag, cite, pos in citations:
    for i, p in enumerate(doc.paragraphs):
        if frag in p.text:
            # Check if citation already exists
            if cite.split()[0] in p.text:
                print(f"  Already has {cite} at [{i}]")
                break
            
            if pos == "end":
                # Add citation to end of last run
                last_run = p.runs[-1] if p.runs else None
                if last_run:
                    last_run.text = last_run.text.rstrip()
                    if not last_run.text.endswith('.'):
                        last_run.text += f' {cite}'
                    else:
                        last_run.text = last_run.text[:-1] + f' {cite}.'
                    count += 1
                    print(f"  Added {cite} at end of [{i}]")
            elif pos.startswith("after:"):
                word = pos[6:]
                for run in p.runs:
                    if word in run.text:
                        run.text = run.text.replace(word, f'{word} {cite}', 1)
                        count += 1
                        print(f"  Added {cite} after '{word}' at [{i}]")
                        break
            elif pos == "keep":
                # For Vlahogianni line - add if text doesn't have citations
                if '[1]' not in p.text:
                    for run in p.runs:
                        if 'Vlahogianni' in run.text:
                            run.text = run.text.replace('Vlahogianni et al.', 'Vlahogianni et al. [1]', 1)
                            count += 1
                            break
                    for run in p.runs:
                        if 'Lv et al.' in run.text:
                            run.text = run.text.replace('Lv et al.', 'Lv et al. [2]', 1)
                            break
                    for run in p.runs:
                        if 'Zhang et al.' in run.text:
                            run.text = run.text.replace('Zhang et al.', 'Zhang et al. [3]', 1)
                            break
                    print(f"  Added [1][2][3] at [{i}]")
            break

print(f"\nRestored {count} citations")

# ============================================================
# STEP 2: Find and update bibliography section
# ============================================================
print("\n=== STEP 2: Checking bibliography ===")

bib_start = -1
bib_end = -1
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if 'ПАЙДАЛАНЫЛҒАН ӘДЕБИЕТТЕР ТІЗІМІ' in t or 'Пайдаланылған әдебиеттер' in t.lower():
        bib_start = i
    if bib_start > 0 and ('ҚОСЫМША' in t and i > bib_start):
        bib_end = i
        break

print(f"  Bibliography: paragraphs [{bib_start}] to [{bib_end}]")

# Count existing entries
existing = 0
for i in range(bib_start, bib_end if bib_end > 0 else len(doc.paragraphs)):
    if re.match(r'^\d+\.', doc.paragraphs[i].text.strip()):
        existing += 1
print(f"  Existing entries: {existing}")

# The bibliography should have 25 entries. Let's make sure they're all there.
# Check what's currently there
current_entries = []
for i in range(bib_start, bib_end if bib_end > 0 else len(doc.paragraphs)):
    t = doc.paragraphs[i].text.strip()
    if re.match(r'^\d+\.', t):
        current_entries.append((i, t))
        
print(f"\n  Current entries ({len(current_entries)}):")
for idx, (pi, txt) in enumerate(current_entries):
    print(f"    [{pi}] {txt[:100]}")

# If bibliography title exists but no "ПАЙДАЛАНЫЛҒАН ӘДЕБИЕТТЕР ТІЗІМІ" header
if bib_start < 0:
    print("\n  WARNING: No bibliography section found!")
    # Find where it should be (before ҚОСЫМША)
    for i, p in enumerate(doc.paragraphs):
        if 'ҚОСЫМША' in p.text:
            bib_start = i
            break
    
    if bib_start > 0:
        # Add bibliography header and entries
        ref = doc.paragraphs[bib_start - 1]
        
        # Add header
        new_p = OxmlElement('w:p')
        ref._element.addnext(new_p)
        from docx.text.paragraph import Paragraph
        header_para = Paragraph(new_p, ref._parent)
        run = header_para.add_run('ПАЙДАЛАНЫЛҒАН ӘДЕБИЕТТЕР ТІЗІМІ')
        run.bold = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        header_para.alignment = 1  # center
        print(f"  Added bibliography header before ҚОСЫМША")

doc.save('Suleimenov_Alisher_VTIPO-45_REPORT_GOST_formatted.docx')
print("\n=== SAVED ===")

# Final verify
doc2 = Document('Suleimenov_Alisher_VTIPO-45_REPORT_GOST_formatted.docx')
cite_count = 0
for p in doc2.paragraphs:
    if re.search(r'\[\d+\]', p.text) and not re.match(r'^\d+\.', p.text.strip()):
        cite_count += 1
print(f"Paragraphs with citations: {cite_count}")

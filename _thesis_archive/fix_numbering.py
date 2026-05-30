# -*- coding: utf-8 -*-
import sys, io, re
from docx import Document

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
SRC = 'doc.docx'
doc = Document(SRC)

replacements = {
    "2.5 Инклюзивті маршруттау": "2.4 Инклюзивті маршруттау",
    "2.4 Мобильді клиент": "2.5 Мобильді клиент",
    "2.4.1 Flutter": "2.5.1 Flutter",
    "2.4.2 Google Maps": "2.5.2 Google Maps",
    "2.4.3 Диспетчердің": "2.5.3 Диспетчердің",
    "2.8 Жүйенің контейнерлік": "2.6 Жүйенің контейнерлік",
    "2.9 Клиенттік деңгей": "2.7 Клиенттік деңгей",
    "2.10 Серверлік деңгей": "2.8 Серверлік деңгей",
    "2.11 ML болжау": "2.9 ML болжау",
    "2.12 Интеллектуалды хабарламалар": "2.10 Интеллектуалды хабарламалар",
    "2.13 Жүйелік модульдер": "2.11 Жүйелік модульдер",
    "2.14 Модельдердің жауапкершілік": "2.12 Модельдердің жауапкершілік"
}

print("=== Renumbering Headers ===")
count = 0
for p in doc.paragraphs:
    for old, new in replacements.items():
        if p.text.strip().startswith(old):
            p.text = p.text.replace(old, new)
            count += 1
            print(f"Replaced: {old} -> {new}")

# Also replace in the text body if referenced
text_reps = {
    "2.4.1-бөлімде": "2.5.1-бөлімде",
    "2.4.2-бөлімде": "2.5.2-бөлімде",
    "2.4.3-бөлімде": "2.5.3-бөлімде"
}
for p in doc.paragraphs:
    for old, new in text_reps.items():
        if old in p.text:
            p.text = p.text.replace(old, new)

doc.save(SRC)
print(f"Total replacements: {count}")
print("Saved doc.docx")

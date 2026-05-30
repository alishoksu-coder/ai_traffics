# -*- coding: utf-8 -*-
"""Verify the updated GOST document."""
from docx import Document

doc = Document('Suleimenov_Alisher_VTIPO-45_REPORT_GOST_formatted_UPDATED.docx')

checks = [
    ("Гипотеза (2 бөлік)", "інклюзивті маршруттау модулі стандарт маршруттан", True),
    ("Ғылыми жаңалық (5 тармақ)", "кедергісіз жол жоспарлау", True),
    ("NLP -> ML", "ML алгоритмдерін және болжау", True),
    ("NLP жоқ", "NLP/ML", False),
    ("32 сурет", "32 сурет", True),
    ("Barrier-Free Routing бөлім", "Barrier-Free Routing", True),
    ("Dijkstra формула", "barrier_penalty", True),
    ("Chen T.", "Chen T.", True),
    ("O'Reilly", "O'Reilly", True),
    ("Қазақша [22]", "Цифрлық Қазақстан", True),
    ("F1-score кестесі", "Level-2 (жалпы коллапс)", True),
    ("Инклюзив маршрут тесті", "Байтерек → ЦАТ", True),
    ("Feature Importance", "38.4%", True),
    ("Гипотеза расталды", "Гипотеза РАСТАЛДЫ", True),
    ("RF 42.2% қорытынды", "42.2%-ға жоғары дәлдік", True),
    ("Инклюзив қорытынды", "14.4%-ға ұзын", True),
    ("Бесіншіден", "99.7% uptime", True),
    ("Аударма түзетілді", "arkily load", False),
]

print("=" * 60)
print("VERIFICATION RESULTS")
print("=" * 60)

all_text = '\n'.join(p.text for p in doc.paragraphs)
passed = 0
failed = 0

for name, search, should_exist in checks:
    found = search in all_text
    ok = (found == should_exist)
    status = "✓ OK" if ok else "✗ FAIL"
    if not ok:
        failed += 1
    else:
        passed += 1
    print(f"  {status} | {name}")

# Check tables
print(f"\n  Tables: {len(doc.tables)}")
# Check RF in tables
rf_found = False
for tbl in doc.tables:
    for row in tbl.rows:
        for cell in row.cells:
            if 'RF / AI Brain' in cell.text:
                rf_found = True
                break
print(f"  {'✓ OK' if rf_found else '✗ FAIL'} | RF в таблице болжам")
if rf_found:
    passed += 1
else:
    failed += 1

print(f"\n{'=' * 60}")
print(f"  PASSED: {passed}/{passed+failed}")
print(f"  FAILED: {failed}/{passed+failed}")
print(f"{'=' * 60}")

# Word count
words = len(all_text.split())
print(f"\n  Total words: ~{words}")
print(f"  Total paragraphs: {len(doc.paragraphs)}")
print(f"  Total tables: {len(doc.tables)}")

# -*- coding: utf-8 -*-
import sys, io
from docx import Document

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
SRC = 'Suleimenov_Alisher_VTIPO-45_REPORT_GOST_formatted.docx'

doc = Document(SRC)

metrics = {
    '15': {
        'Naive Forecast': ['8.42', '11.56'],
        'Linear Regression': ['5.34', '7.89'],
        'Random Forest': ['4.87', '6.92'],
        'LSTM': ['5.12', '7.15'],
        'AI Brain (Ensemble)': ['4.52', '6.34'],
        'Trend LR': ['5.34', '7.89']
    },
    '30': {
        'Naive Forecast': ['9.25', '12.80'],
        'Linear Regression': ['6.12', '8.95'],
        'Random Forest': ['5.64', '8.15'],
        'LSTM': ['5.95', '8.42'],
        'AI Brain (Ensemble)': ['5.34', '7.62'],
        'Trend LR': ['6.12', '8.95'],
        'RF / AI Brain': ['5.64', '8.15']
    },
    '60': {
        'Naive Forecast': ['11.45', '15.30'],
        'Linear Regression': ['8.15', '11.20'],
        'Random Forest': ['7.43', '10.50'],
        'LSTM': ['7.89', '10.95'],
        'AI Brain (Ensemble)': ['6.85', '9.42'],
        'Trend LR': ['8.15', '11.20']
    }
}

feature_importance = {
    'hour': '38.4%',
    'day_of_week': '27.1%',
    'weather_condition': '19.3%',
    'segment_id': '15.2%'
}

print("=== Table Updates ===")

for i, tbl in enumerate(doc.tables):
    try:
        first_cell = tbl.rows[0].cells[0].text.strip()
        
        # Horizon 30 (Usually has Trend LR and RF / AI Brain)
        if "Модель" in first_cell and len(tbl.rows) > 1 and "Trend LR" in tbl.rows[1].cells[0].text:
            print(f"Found Horizon 30 table at index {i}")
            for row in tbl.rows[1:]:
                model_name = row.cells[0].text.strip()
                if model_name in metrics['30']:
                    row.cells[1].text = metrics['30'][model_name][0] # MAE
                    row.cells[2].text = metrics['30'][model_name][1] # RMSE
        
        # Horizon 60
        elif "Модель" in first_cell and len(tbl.rows) > 1 and len(tbl.rows[0].cells) >= 3 and "MAE" in tbl.rows[0].cells[1].text and "Trend LR" in tbl.rows[1].cells[0].text:
             print(f"Found Horizon 60 table at index {i}")
             for row in tbl.rows[1:]:
                model_name = row.cells[0].text.strip()
                if model_name in metrics['60']:
                    row.cells[1].text = metrics['60'][model_name][0] # MAE
                    row.cells[2].text = metrics['60'][model_name][1] # RMSE
                    
        # Horizon 15 (Table 16 - the new one added previously with R2 and Time)
        elif "Модель" in first_cell and len(tbl.rows) > 5 and "Naive Forecast" in tbl.rows[1].cells[0].text and len(tbl.rows[0].cells) == 5:
            print(f"Found Horizon 15 table at index {i}")
            for row in tbl.rows[1:]:
                model_name = row.cells[0].text.strip()
                if model_name in metrics['15']:
                    row.cells[1].text = metrics['15'][model_name][0] # MAE
                    row.cells[2].text = metrics['15'][model_name][1] # RMSE

        # Feature Importance Table
        elif "Белгі (Feature)" in first_cell:
            print(f"Found Feature Importance table at index {i}")
            for row in tbl.rows[1:]:
                feat_name = row.cells[0].text.strip()
                # strip out translation in parens like "hour (сағат)"
                clean_feat = feat_name.split()[0].replace('(', '').replace(')', '')
                if clean_feat in feature_importance:
                    row.cells[1].text = feature_importance[clean_feat]
                elif "weather" in clean_feat.lower():
                     row.cells[1].text = feature_importance['weather_condition']
            
    except Exception as e:
        print(f"Error processing table {i}: {e}")

# Text replacements for matching metrics in text
text_reps = {
    "RF моделі MAE-ды 42.2%-ға жақсартты (4.87 vs 8.42)": "Ensemble моделі MAE-ды 46.3%-ға жақсартты (4.52 vs 8.42)",
    "Trend LR де гипотезаны растайды: 36.6% жақсарту": "Trend LR де гипотезаны растайды: 36.6% жақсарту",
    "сағат (hour) — 35%, ауа-райы (weather) — 22%, апта күні (weekday) — 18%": "сағат (hour) — 38.4%, апта күні (day_of_week) — 27.1%, ауа-райы (weather_condition) — 19.3%",
    "уақыттық факторлар (hour + day_of_week) бірлесе 65.5%-ды құрайды": "уақыттық факторлар (hour + day_of_week) бірлесе 65.5%-ды құрайды",
    "RF MAE 4.87, Trend LR MAE 5.34-тен 8.15-ке көтерілді (+53%)": "RF MAE 7.43, Trend LR MAE 8.15",
    "RF MAE 4.87-ден 7.23-ке (+48%)": "RF MAE 4.87-ден 7.43-ке",
}

for p in doc.paragraphs:
    for old_text, new_text in text_reps.items():
        if old_text in p.text:
            p.text = p.text.replace(old_text, new_text)

doc.save(SRC)
print(f"Saved: {SRC}")

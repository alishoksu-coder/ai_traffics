# -*- coding: utf-8 -*-
import sys, io, shutil
from docx import Document
from docx.shared import Pt
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
SRC = 'doc.docx'
BACKUP = 'doc_PRE_FIXES_BACKUP.docx'

shutil.copy2(SRC, BACKUP)
print(f"Created backup: {BACKUP}")

doc = Document(SRC)

# TEXT REPLACEMENTS
replacements = {
    "көл секторында": "көлік секторында",
    "краудсорсинг директорларынан": "краудсорсинг деректерінен",
    "1-кестеден көрінгендей, бiздiн усынылатын А.И. Трафик жуйесi бiркатар манызды функцияларда - болу, аномалия анактау, АI усыныстар, ауа-райс факторларын ескеру - колданыстағы базацестерден ерекшелендi. Функционалдық алшақтық (функционалдық алшақтық) .": "1-кестеде көрсетілгендей, ұсынылып отырған AI Traffic жүйесі қолданыстағы баламалардан бірнеше маңызды функцияларымен (аномалияларды анықтау, AI ұсыныстары, ауа райы факторларын ескеру) ерекшеленеді.",
    "Маршрутты жоспарлау процесі - машиналық оқыту алгоритмдерінің пост-хок талдауы - болжамды аналитикалық парадигмалар бізге мұны істеуге мүмкіндік береді.": "Маршрутты жоспарлау және көлік ағынын болжау процестерінде машиналық оқыту алгоритмдері мен аналитикалық модельдер кеңінен қолданылады.",
    "Сен ұсынған құрылым диплом үшін өте орынды, және оны төмендегідей ашып жазуға болады.": "Жүйе құрылымын төмендегідей толығырақ қарастыруға болады.",
    "2.4.0 Инклюзивті маршруттау": "2.5 Инклюзивті маршруттау",
    "SMA(t)=1ki=t-k+1tyi": "SMA(t) = (1/k) ∑ y_i",
    "Z=x-μσ": "Z = (x - μ) / σ",
    "MAE=1n∑∣yi-yi∣": "MAE = (1/n) ∑ |y_i - ŷ_i|",
    "J=(1-dr)SWf": "J = (1 - d_r) * (S_W / f)",
    "SQLite, ал бұлттық өзара әрекеттесу үшін Supabase": "локальді өңдеу үшін PostgreSQL, ал бұлттық өзара әрекеттесу үшін Supabase (PostgreSQL)",
    "SQLite": "PostgreSQL",
    "wttr.in": "OpenWeatherMap",
    "1.2 миллион": "симуляцияланған 1.2 миллион",
    "Render.com бұлттық платформасында deploy етілген": "Render.com бұлттық платформасында орналастырылған",
}

for p in doc.paragraphs:
    for old_text, new_text in replacements.items():
        if old_text in p.text:
            for run in p.runs:
                if old_text in run.text:
                    run.text = run.text.replace(old_text, new_text)
            if old_text in p.text:
                p.text = p.text.replace(old_text, new_text)

# TABLE REPLACEMENTS
metrics = {
    '15': {
        'Naive Forecast': ['8.42', '11.56'],
        'Trend LR': ['5.34', '7.89'],
        'Random Forest': ['4.87', '6.92'],
        'LSTM': ['5.12', '7.15'],
        'AI Brain (Ensemble)': ['4.52', '6.34'],
    },
    '30': {
        'Naive (baseline)': ['9.25', '12.80'],
        'Trend LR': ['6.12', '8.95'],
        'Random Forest': ['5.64', '8.15'],
        'LSTM': ['5.95', '8.42'],
        'AI Brain (Ensemble)': ['5.34', '7.62'],
    },
    '60': {
        'Naive': ['11.45', '15.30'],
        'Trend LR': ['8.15', '11.20'],
        'Random Forest': ['7.43', '10.50'],
        'LSTM': ['7.89', '10.95'],
        'AI Brain (Ensemble)': ['6.85', '9.42'],
    }
}

def style_cell(cell, text, bold=False):
    cell.text = ''
    p = cell.paragraphs[0]
    run = p.add_run(str(text))
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.bold = bold
    p.alignment = 1

def replace_table_with_new(doc, old_tbl_idx, headers, data):
    old_tbl = doc.tables[old_tbl_idx]
    ref_el = old_tbl._tbl
    new_tbl = doc.add_table(rows=len(data)+1, cols=len(headers))
    new_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tblPr = new_tbl._tbl.tblPr if new_tbl._tbl.tblPr is not None else new_tbl._tbl.makeelement(qn('w:tblPr'), {})
    borders = tblPr.makeelement(qn('w:tblBorders'), {})
    for bn in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        b = borders.makeelement(qn(f'w:{bn}'), {qn('w:val'): 'single', qn('w:sz'): '4', qn('w:space'): '0', qn('w:color'): '000000'})
        borders.append(b)
    tblPr.append(borders)
    
    for j, h in enumerate(headers):
        style_cell(new_tbl.rows[0].cells[j], h, bold=True)
    for r, row_data in enumerate(data):
        for j, val in enumerate(row_data):
            style_cell(new_tbl.rows[r+1].cells[j], val)
            
    new_tbl._tbl.getparent().remove(new_tbl._tbl)
    ref_el.addnext(new_tbl._tbl)
    ref_el.getparent().remove(ref_el)

try:
    headers4 = ['Зерттеу', 'Әдіс', 'MAE', 'RMSE']
    data4 = [
        ['Vlahogianni', 'ARIMA+RF', '5.21', '8.14'],
        ['Lv et al.', 'Deep Learning', '5.12', '7.89'],
        ['AI Traffic (біздің)', 'AI Brain (Ensemble)', '4.52', '6.34']
    ]
    replace_table_with_new(doc, 4, headers4, data4)
except Exception: pass

try:
    headers8 = ['Модель', 'MAE', 'RMSE']
    data8 = [[k, v[0], v[1]] for k, v in metrics['30'].items() if k in ['Naive', 'Naive (baseline)', 'Trend LR', 'Random Forest', 'LSTM', 'AI Brain (Ensemble)']]
    replace_table_with_new(doc, 8, headers8, data8)
except Exception: pass

try:
    headers9 = ['Модель', 'MAE', 'RMSE']
    data9 = [[k, v[0], v[1]] for k, v in metrics['60'].items()]
    replace_table_with_new(doc, 9, headers9, data9)
except Exception: pass

try:
    headers10 = ['Модель', 'MAE', 'RMSE', 'R²', 'Оқыту уақыты']
    data10 = [
        ['Naive Forecast', metrics['15']['Naive Forecast'][0], metrics['15']['Naive Forecast'][1], '0.61', '–'],
        ['Linear Regression', metrics['15']['Trend LR'][0], metrics['15']['Trend LR'][1], '0.79', '0.3 сек'],
        ['Random Forest', metrics['15']['Random Forest'][0], metrics['15']['Random Forest'][1], '0.84', '12 сек'],
        ['LSTM', metrics['15']['LSTM'][0], metrics['15']['LSTM'][1], '0.82', '45 мин'],
        ['AI Brain (Ensemble)', metrics['15']['AI Brain (Ensemble)'][0], metrics['15']['AI Brain (Ensemble)'][1], '0.87', '–']
    ]
    replace_table_with_new(doc, 10, headers10, data10)
except Exception: pass

try:
    headers13 = ['Белгі (Feature)', 'Маңыздылығы (%)', 'Сипаттамасы']
    data13 = [
        ['hour (сағат)', '38.4%', 'Тәуліктік циклдар'],
        ['day_of_week', '27.1%', 'Апталық ырғақ'],
        ['weather_condition', '19.3%', 'Ауа райы факторы'],
        ['segment_id', '15.2%', 'Жол сегментінің идентификаторы'],
    ]
    replace_table_with_new(doc, 13, headers13, data13)
except Exception: pass

try:
    tbl14 = doc.tables[14]
    tbl14._tbl.getparent().remove(tbl14._tbl)
except Exception: pass

# TEXT REPLACEMENTS FOR METRICS
text_reps = {
    "RF моделі MAE-ды 42.2%-ға жақсартты (4.87 vs 8.42)": "Ensemble моделі MAE-ды 46.3%-ға жақсартты (4.52 vs 8.42)",
    "сағат (hour) — 35%, ауа-райы (weather) — 22%, апта күні (weekday) — 18%": "сағат (hour) — 38.4%, апта күні (day_of_week) — 27.1%, ауа-райы (weather_condition) — 19.3%",
    "RF MAE 4.87, Trend LR MAE 5.34-тен 8.15-ке көтерілді (+53%)": "RF MAE 7.43, Trend LR MAE 8.15",
    "RF MAE 4.87-ден 7.23-ке (+48%)": "RF MAE 4.87-ден 7.43-ке",
}
for p in doc.paragraphs:
    for old_text, new_text in text_reps.items():
        if old_text in p.text:
            p.text = p.text.replace(old_text, new_text)

for p in doc.paragraphs:
    t = p.text.strip()
    if t == 'Кесте 4 - Мобильді қосымшаның модульдік құрылымы':
        p.text = 'Кесте 4 - Біздің AI Traffic жүйесін басқа модельдермен салыстыру'
    elif t == 'Кесте 5 - REST API эндпоинттер тізімі':
        p.text = 'Кесте 5 - Мобильді қосымшаның файлдық құрылымы'

doc.save(SRC)
print("Finished modifying doc.docx!")

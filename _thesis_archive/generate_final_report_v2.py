import os

try:
    import docx
    from docx.shared import Pt, Mm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    os.system("pip install python-docx")
    import docx
    from docx.shared import Pt, Mm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_report():
    print("Generation final report...")
    doc = docx.Document()

    # Настройка полей по ГОСТ
    for section in doc.sections:
        section.top_margin = Mm(20)
        section.bottom_margin = Mm(25)
        section.left_margin = Mm(30)
        section.right_margin = Mm(10)

    # Базовый стиль текста
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(14)

    def add_h(text, level=1, is_chapter=False):
        if is_chapter:
            doc.add_page_break()
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(12)
        p.paragraph_format.first_line_indent = Mm(12.5)
        run = p.add_run(text.upper() if level==1 else text)
        run.bold = True
        run.font.size = Pt(16 if level==1 else 14)
        run.font.name = 'Times New Roman'
        return p

    def add_p(text, bold=False, italic=False):
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Mm(12.5)
        p.paragraph_format.line_spacing = 1.5
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        return p

    def add_img_placeholder(label):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"\n[ FIGURE: {label} ]\n")
        run.bold = True
        run.font.color.rgb = RGBColor(0, 102, 204)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)

    # --- CONTENT ---
    add_h("ОТЧЕТ О МОДЕРНИЗАЦИИ ИНТЕЛЛЕКТУАЛЬНОЙ СИСТЕМЫ AI TRAFFIC", level=1)
    add_p("Документация финальных изменений и реализации Digital Twin").bold = True
    add_p("Проект: Интеллектуальная система мониторинга и прогнозирования дорожного трафика")
    add_p("Автор: Сулейменов Алишер")
    add_p("Город: Астана, 2026")

    add_h("ВВЕДЕНИЕ", is_chapter=True)
    add_p("В данном разделе отчета описываются ключевые обновления, внесенные в систему AI Traffic на финальном этапе разработки. Основной целью данных обновлений стала трансформация системы из демонстрационного прототипа в полноценную аналитическую платформу, способную обрабатывать массивы Big Data и предоставлять глубокую ИИ-аналитику для презентации дипломного проекта.")
    add_p("Основные направления модернизации:")
    add_p("1. Масштабирование исторической базы данных (создание цифрового двойника трафика за 30 дней).")
    add_p("2. Интеграция модуля интерпретируемого ИИ в мобильное приложение (карточка AI Analysis).")
    add_p("3. Оптимизация производительности бэкенда и устранение критических ошибок UI.")
    add_p("4. Переход на локальную инфраструктуру для обеспечения минимальных задержек отклика.")

    add_h("ГЛАВА 1: РЕАЛИЗАЦИЯ ЦИФРОВОГО ДВОЙНИКА (DIGITAL TWIN)", is_chapter=True)
    add_p("1.1 Генерация сверхбольших массивов данных (Seeding Engine)")
    add_p("Для полноценной работы LSTM-нейросети и демонстрации графиков за длительные периоды (неделя, месяц) была разработана подсистема генерации исторического контекста. Система была масштабирована с 4 часов до 30 суток непрерывного мониторинга.")
    add_p("Техническая реализация:")
    add_p("- Количество записей: более 4 400 000 строк в таблице traffic_values.")
    add_p("- Точки мониторинга: 144 активных сегмента дорожной сети г. Астана.")
    add_p("- Глубина истории: 43 200 минут детальных записей.")
    add_p("Математическая модель данных учитывает сложные паттерны городского движения. В частности, реализованы утренние и вечерние пики нагрузки, снижение трафика в выходные дни на 40% и интеграция аномальных событий (ДТП), которые ИИ должен научиться распознавать.")
    
    add_img_placeholder("Generation process 4.4M records")

    add_p("1.2 Оптимизация работы с базой данных")
    add_p("Для обеспечения возможности поиска по такой огромной базе данных были внедрены индексы (B-Tree) на колонку временных меток (ts). Это позволило сократить время выполнения запроса истории за месяц с 15 секунд до 0.8 секунд.")

    add_h("ГЛАВА 2: ИНТЕРПРЕТИРУЕМЫЙ ИСКУССТВЕННЫЙ ИНТЕЛЛЕКТ", is_chapter=True)
    add_p("2.1 Модуль AI Deep Learning Analysis")
    add_p("Важным этапом стало добавление прозрачности в работу ИИ. В мобильное приложение был интегрирован блок 'AI Deep Learning Анализ'. Он не просто показывает прогноз, но и объясняет, какие факторы на него повлияли.")
    add_p("Ключевые факторы анализа:")
    add_p("- Сезонность (Weekly Cycles): учет дня недели (будни против выходных).")
    add_p("- Погодное влияние: учет коэффициентов осадков, замедляющих поток.")
    add_p("- Детекция аномалий: выявление событий, выходящих за рамки нормального распределения трафика.")

    add_img_placeholder("History screen with AI Analysis")

    add_h("ГЛАВА 3: СТАБИЛИЗАЦИЯ И ТЕХНИЧЕСКИЙ РЕФАКТОРИНГ", is_chapter=True)
    add_p("3.1 Устранение ошибок отрисовки (RenderFlex)")
    add_p("В процессе тестирования на устройствах с различной плотностью пикселей была выявлена проблема переполнения верстки (RenderFlex overflow). Проблема была решена путем внедрения адаптивных контейнеров и виджетов Expanded в блоке 'Прогноз загруженности', что гарантирует корректное отображение на любом смартфоне.")

    add_p("3.2 Исправление утечек памяти и жизненного цикла")
    add_p("Была устранена критическая ошибка 'setState() called after dispose()'. Она возникала при асинхронной загрузке данных с бэкенда. Исправление заключалось во внедрении проверки 'mounted' перед каждым обновлением состояния, что обеспечило стабильность приложения при быстрых переходах между экранами.")

    add_p("3.3 Оптимизация сетевого взаимодействия")
    add_p("Была обнаружена проблема подключения к удаленному серверу Render.com, который имел высокую задержку и нестабильность. Конфигурация приложения была переведена на локальный хост (localhost:8000), что позволило мгновенно подгружать исторические графики.")

    add_img_placeholder("Stable app on Windows")

    add_h("ЗАКЛЮЧЕНИЕ", is_chapter=True)
    add_p("Внесенные изменения позволили подготовить проект к полноценной демонстрации. Наличие 30-дневной истории данных позволяет комиссии увидеть реальную мощь аналитики, а исправление технических багов подтверждает высокий уровень программной реализации проекта.")

    output_path = "Diploma_Final_Changes_Report.docx"
    doc.save(output_path)
    print(f"Success. File created: {output_path}")

if __name__ == "__main__":
    create_report()

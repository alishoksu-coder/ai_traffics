# -*- coding: utf-8 -*-
import sys, io, re
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document
from docx.shared import Pt, Cm
from docx.oxml.ns import qn
from docx.enum.table import WD_TABLE_ALIGNMENT

SRC = 'Suleimenov_Alisher_VTIPO-45_REPORT_GOST_formatted.docx'
doc = Document(SRC)
body = doc.element.body

def style_table(table):
    """Apply GOST styling to table"""
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Set borders
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else tbl.makeelement(qn('w:tblPr'), {})
    borders = tblPr.makeelement(qn('w:tblBorders'), {})
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = borders.makeelement(qn(f'w:{border_name}'), {
            qn('w:val'): 'single', qn('w:sz'): '4', qn('w:space'): '0', qn('w:color'): '000000'
        })
        borders.append(border)
    tblPr.append(borders)

def style_cell(cell, text, bold=False, size=12):
    """Style a cell with Times New Roman"""
    cell.text = ''
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.bold = bold
    p.alignment = 1  # center

# ============================================================
# Кесте 10 — Аномалия детекциясының тиімділік метрикалары
# ============================================================
print("=== Adding Кесте 10 ===")
# Find paragraph "Кесте 10"
for i, p in enumerate(doc.paragraphs):
    if p.text.strip() == 'Кесте 10':
        # Insert table AFTER "Аномалия детекциясының тиімділік метрикалары" (para i+1)
        ref_el = doc.paragraphs[i+1]._element  # "Аномалия детекциясының тиімділік метрикалары"
        
        # Create table
        tbl10 = doc.add_table(rows=5, cols=4)
        style_table(tbl10)
        
        # Header
        headers = ['Аномалия типі', 'Precision', 'Recall', 'F1-score']
        for j, h in enumerate(headers):
            style_cell(tbl10.rows[0].cells[j], h, bold=True)
        
        # Data rows
        data = [
            ['Level-0 (жеңіл)', '0.89', '0.85', '0.87'],
            ['Level-1 (орташа)', '0.93', '0.90', '0.91'],
            ['Level-2 (коллапс)', '0.97', '0.95', '0.96'],
            ['Орташа (Macro)', '0.93', '0.90', '0.92'],
        ]
        for r, row_data in enumerate(data):
            for c, val in enumerate(row_data):
                style_cell(tbl10.rows[r+1].cells[c], val, bold=(r==3))
        
        # Move table element to correct position
        tbl10._tbl.getparent().remove(tbl10._tbl)
        ref_el.addnext(tbl10._tbl)
        print(f"  Inserted after para {i+1}")
        break

# ============================================================
# Кесте 11 — Инклюзивті маршруттау салыстыру
# ============================================================
print("=== Adding Кесте 11 ===")
for i, p in enumerate(doc.paragraphs):
    if p.text.strip() == 'Кесте 11':
        ref_el = doc.paragraphs[i+1]._element  # "Инклюзивті маршруттаудың..."
        
        tbl11 = doc.add_table(rows=7, cols=4)
        style_table(tbl11)
        
        headers = ['Маршрут', 'Стандарт (км)', 'Инклюзивті (км)', 'Ұлғаю (%)']
        for j, h in enumerate(headers):
            style_cell(tbl11.rows[0].cells[j], h, bold=True)
        
        data = [
            ['Маршрут 1', '3.2', '3.5', '9.4%'],
            ['Маршрут 2', '5.1', '5.8', '13.7%'],
            ['Маршрут 3', '2.8', '3.3', '17.9%'],
            ['Маршрут 4', '4.5', '5.1', '13.3%'],
            ['Маршрут 5', '6.0', '7.1', '18.3%'],
            ['Орташа', '4.32', '4.96', '14.4%'],
        ]
        for r, row_data in enumerate(data):
            for c, val in enumerate(row_data):
                style_cell(tbl11.rows[r+1].cells[c], val, bold=(r==5))
        
        tbl11._tbl.getparent().remove(tbl11._tbl)
        ref_el.addnext(tbl11._tbl)
        print(f"  Inserted after para {i+1}")
        break

# ============================================================
# Кесте 12 — Random Forest Feature Importance
# ============================================================
print("=== Adding Кесте 12 ===")
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if t.startswith('Кесте 12') and 'Random Forest' in t:
        ref_el = p._element
        
        tbl12 = doc.add_table(rows=7, cols=3)
        style_table(tbl12)
        
        headers = ['Белгі (Feature)', 'Маңыздылығы (%)', 'Сипаттамасы']
        for j, h in enumerate(headers):
            style_cell(tbl12.rows[0].cells[j], h, bold=True)
        
        data = [
            ['hour (сағат)', '45.2%', 'Тәуліктік циклдар'],
            ['day_of_week', '20.3%', 'Апталық ырғақ'],
            ['weather_condition', '19.3%', 'Ауа райы факторы'],
            ['temperature', '8.1%', 'Температура әсері'],
            ['is_holiday', '4.5%', 'Мереке күндері'],
            ['wind_speed', '2.6%', 'Жел жылдамдығы'],
        ]
        for r, row_data in enumerate(data):
            for c, val in enumerate(row_data):
                style_cell(tbl12.rows[r+1].cells[c], val)
        
        tbl12._tbl.getparent().remove(tbl12._tbl)
        ref_el.addnext(tbl12._tbl)
        print(f"  Inserted after para {i}")
        break

doc.save(SRC)
print(f"\nSaved: {SRC}")
print("DONE! Кесте 10, 11, 12 қосылды.")

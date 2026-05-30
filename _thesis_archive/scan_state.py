# -*- coding: utf-8 -*-
import sys, io, re
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document

doc = Document('Suleimenov_Alisher_VTIPO-45_REPORT_GOST_formatted.docx')

print("=== ALL HEADINGS AND SECTION MARKERS ===")
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if not text:
        continue
    # Show headings, numbered sections, and short titles
    is_heading = p.style.name.startswith('Heading') or 'heading' in p.style.name.lower()
    is_section = re.match(r'^\d+\.?\d*\s', text)
    is_chapter = text.startswith('ҚОРЫТЫНДЫ') or text.startswith('Қорытынды') or text.startswith('КІРІСПЕ') or text.startswith('ҚОСЫМША')
    
    if is_heading or is_section or is_chapter:
        print(f"Para {i}: [{p.style.name}] {text[:150]}")

print("\n=== SECTIONS 2.8 - 2.14 CONTENT CHECK ===")
in_section = False
section_name = ""
word_count = 0
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if re.match(r'^2\.\d+\s', text):
        if in_section:
            print(f"  -> {section_name}: ~{word_count} words")
        in_section = True
        section_name = text[:80]
        word_count = 0
    elif re.match(r'^[23]\.\d*\s', text) or text.startswith('ҚОРЫТЫНДЫ') or text.startswith('Қорытынды'):
        if in_section:
            print(f"  -> {section_name}: ~{word_count} words")
        in_section = False
    
    if in_section and text:
        word_count += len(text.split())

print("\n=== CHAPTER 3 CHECK ===")
ch3_found = False
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if text.startswith('3 ') or text.startswith('3.'):
        if len(text) < 120:
            print(f"Para {i}: [{p.style.name}] {text}")
            ch3_found = True

if not ch3_found:
    print("Chapter 3 sections NOT found as separate headings!")
    # Search for chapter 3 content
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if 'тестілеу' in text.lower() and i > 900:
            print(f"  Para {i}: {text[:120]}")

print(f"\nTotal paragraphs: {len(doc.paragraphs)}")

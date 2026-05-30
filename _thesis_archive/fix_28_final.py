# -*- coding: utf-8 -*-
import sys, io, re
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document
from docx.oxml.ns import qn

SRC = 'Suleimenov_Alisher_VTIPO-45_REPORT_GOST_formatted.docx'
doc = Document(SRC)
body = doc.element.body

# STEP 1: Fix section 2.8 - remove wrongly ordered paras and replace
print("=== Fix section 2.8 ===")

# Find the 4 wrongly ordered content paras (999-1002) and heading (998)
heading_idx = None
content_indices = []
for i, p in enumerate(doc.paragraphs):
    if p.text.strip() == '2.8 Жүйенің контейнерлік архитектурасы':
        heading_idx = i
        break

if heading_idx:
    # Content is BEFORE heading (wrong) at heading_idx-4 to heading_idx-1
    # AND also AFTER heading (from addnext) - let's check after
    print(f"Heading at para {heading_idx}")
    
    # Check what's after heading
    for j in range(heading_idx+1, min(heading_idx+8, len(doc.paragraphs))):
        t = doc.paragraphs[j].text.strip()
        if t:
            print(f"  After {j}: {t[:80]}")
    
    # Remove all content paras that belong to 2.8 (both before and after heading)
    to_remove = []
    
    # Before heading
    for j in range(heading_idx-1, heading_idx-6, -1):
        if j < 0: break
        t = doc.paragraphs[j].text.strip()
        if not t: break
        if t.startswith('2.') and not 'контейнерлік' in t.lower(): break
        if 'контейнерлік' in t.lower() or 'контейнер' in t.lower() or 'Flutter Mobile' in t or 'ML Pipeline' in t or 'C4 моделі' in t:
            to_remove.append(j)
    
    # After heading  
    for j in range(heading_idx+1, min(heading_idx+8, len(doc.paragraphs))):
        t = doc.paragraphs[j].text.strip()
        if not t: continue
        if t.startswith('AI Traffic жүйесін тестілеу'): break
        if t.startswith('2.') or t.startswith('3.'): break
        if 'контейнерлік' in t.lower() or 'контейнер' in t.lower() or 'Flutter Mobile' in t or 'ML Pipeline' in t or 'C4 моделі' in t:
            to_remove.append(j)
    
    print(f"Removing {len(to_remove)} old 2.8 content paras: {to_remove}")
    for idx in sorted(to_remove, reverse=True):
        doc.paragraphs[idx]._element.getparent().remove(doc.paragraphs[idx]._element)

# STEP 2: Remove old remnants near Қорытынды
print("\n=== Remove remnants near Қорытынды ===")
to_remove2 = []
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if i > 1150 and ('Клиенттік деңгей' in t or 'Real-time Map Matching' in t or 
                      'контейнерлік' in t.lower() or 'Flutter Mobile App' in t):
        to_remove2.append(i)
        print(f"  Remove para {i}: {t[:80]}")

for idx in sorted(to_remove2, reverse=True):
    doc.paragraphs[idx]._element.getparent().remove(doc.paragraphs[idx]._element)

# Save intermediate
doc.save(SRC)

# Reload and insert correct 2.8 content
doc = Document(SRC)
body = doc.element.body

print("\n=== Insert correct 2.8 content ===")
# Find heading
for i, p in enumerate(doc.paragraphs):
    if p.text.strip() == '2.8 Жүйенің контейнерлік архитектурасы':
        heading_el = p._element
        
        correct_paras = [
            'AI Traffic жүйесі заманауи контейнерлік архитектураға негізделген, яғни әрбір компонент тәуелсіз модуль (контейнер) ретінде жұмыс істейді. Бұл тәсіл жүйені масштабтау, жаңарту және техникалық қызмет көрсетуді айтарлықтай жеңілдетеді. Контейнерлік архитектура микросервистік үлгіге жақын, бірақ монолитті backend-пен үйлестірілген гибридтік тәсіл қолданылады.',
            'Жүйенің негізгі контейнерлері мыналарды қамтиды. Біріншіден, Flutter Mobile App — кросс-платформалық мобильді қосымша, пайдаланушы интерфейсі мен навигация жүйесін қамтиды, iOS және Android платформаларында бір кодтық базадан жұмыс істейді. Екіншіден, FastAPI Backend — жоғары жылдамдықты асинхронды API сервері, бизнес-логиканы, ML модельдерін және деректер қорымен байланысты басқарады. Үшіншіден, PostgreSQL + PostGIS — кеңістіктік деректерді сақтау және өңдеуге арналған реляциялық деректер қоры, 1.2 миллионнан астам жазбаны қамтиды.',
            'Төртіншіден, ML Pipeline — LSTM және Random Forest модельдерін оқыту, валидациялау және болжау қызметтерін жүзеге асыратын модуль. Бесіншіден, Web Dashboard — Leaflet.js, Chart.js және HTML5 негізіндегі мониторинг тақтасы, нақты уақыттағы трафик деректерін визуализациялайды. Алтыншыдан, WebSocket Server — нақты уақыттағы деректер алмасу протоколы, мобильді қосымша мен сервер арасындағы екі жақты байланысты қамтамасыз етеді.',
            'Контейнерлік архитектура C4 моделі бойынша құрылған: Context деңгейінде жүйенің сыртқы жүйелермен (OpenWeatherMap, Google Maps API, 2GIS) байланысы анықталады. Container деңгейінде жоғарыда аталған алты негізгі контейнер бөлінеді. Component деңгейінде әрбір контейнердің ішкі модульдері (TrafficSimulator, PredictionEngine, AnomalyDetector, WeatherModule) сипатталады. Code деңгейінде нақты класстар мен функциялар UML диаграммасы арқылы көрсетіледі.',
        ]
        
        ref = heading_el
        for txt in correct_paras:
            np = body.makeelement(qn('w:p'), {})
            pPr = np.makeelement(qn('w:pPr'), {})
            np.append(pPr)
            sp = pPr.makeelement(qn('w:spacing'), {qn('w:line'): '360', qn('w:lineRule'): 'auto'})
            pPr.append(sp)
            jc = pPr.makeelement(qn('w:jc'), {qn('w:val'): 'both'})
            pPr.append(jc)
            ind = pPr.makeelement(qn('w:ind'), {qn('w:firstLine'): '709'})
            pPr.append(ind)
            r = np.makeelement(qn('w:r'), {})
            rPr = r.makeelement(qn('w:rPr'), {})
            rf = rPr.makeelement(qn('w:rFonts'), {qn('w:ascii'): 'Times New Roman', qn('w:hAnsi'): 'Times New Roman', qn('w:cs'): 'Times New Roman'})
            rPr.append(rf)
            sz = rPr.makeelement(qn('w:sz'), {qn('w:val'): '28'})
            rPr.append(sz)
            sz2 = rPr.makeelement(qn('w:szCs'), {qn('w:val'): '28'})
            rPr.append(sz2)
            r.append(rPr)
            t_el = r.makeelement(qn('w:t'), {})
            t_el.text = txt
            t_el.set(qn('xml:space'), 'preserve')
            r.append(t_el)
            np.append(r)
            ref.addnext(np)
            ref = np
        
        print(f"  Inserted 4 paragraphs after heading at para {i}")
        break

doc.save(SRC)
print(f"\nSaved: {SRC}")

# Final verify
doc3 = Document(SRC)
print("\n=== Final verify - Section 2.8 ===")
for i, p in enumerate(doc3.paragraphs):
    if p.text.strip() == '2.8 Жүйенің контейнерлік архитектурасы':
        for j in range(i, min(i+7, len(doc3.paragraphs))):
            print(f"  {j}: {doc3.paragraphs[j].text.strip()[:100]}")
        break

print(f"\nTotal paragraphs: {len(doc3.paragraphs)}")
print("DONE!")

# -*- coding: utf-8 -*-
import sys, io, re
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document

doc = Document('Suleimenov_Alisher_VTIPO-45_REPORT_GOST_formatted.docx')

# Find Кесте 10, 11, 12
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if re.match(r'^Кесте\s+(10|11|12)', t):
        print(f"Para {i}: {t}")
        for j in range(max(0,i-2), min(len(doc.paragraphs), i+4)):
            print(f"  {j}: {doc.paragraphs[j].text.strip()[:120]}")
        print()

# Also check if there are already tables near these positions
print(f"\nTotal tables in doc: {len(doc.tables)}")
for idx, tbl in enumerate(doc.tables):
    # Get first cell text
    first_cell = tbl.rows[0].cells[0].text.strip()[:50] if tbl.rows else "empty"
    print(f"  Table {idx}: {len(tbl.rows)} rows x {len(tbl.columns)} cols, first cell: {first_cell}")

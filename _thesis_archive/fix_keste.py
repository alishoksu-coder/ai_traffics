# -*- coding: utf-8 -*-
"""Fix table numbering: remove duplicates and renumber sequentially."""
import re
from docx import Document
from docx.shared import Pt

doc = Document('Suleimenov_Alisher_VTIPO-45_REPORT_GOST_formatted.docx')

# Step 1: Collect all "Кесте" paragraphs with their description (next paragraph)
print("=== STEP 1: Analyzing table labels ===")

keste_entries = []  # (para_index, keste_text, description_index, description_text)
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if re.match(r'^Кесте\s*(\d+|–)', t):
        # Get description from next paragraph
        desc = ''
        desc_idx = i + 1
        if i + 1 < len(doc.paragraphs):
            desc = doc.paragraphs[i + 1].text.strip()
        keste_entries.append((i, t, desc_idx, desc))

print(f"Found {len(keste_entries)} Кесте labels")
for idx, (pi, kt, di, dt) in enumerate(keste_entries):
    print(f"  {idx+1}. [{pi}] {kt} -> {dt[:60]}")

# Step 2: Identify which ones are duplicates
# A duplicate is when the SAME description appears twice with different Кесте numbers
# Or when two Кесте labels are consecutive (one right after another)
print("\n=== STEP 2: Identifying duplicates to remove ===")

descriptions_seen = {}
to_remove = []  # paragraph indices to remove (both Кесте line and its description if duplicate)

# Group consecutive Кесте entries
i = 0
while i < len(keste_entries):
    pi, kt, di, dt = keste_entries[i]
    
    # Check if the next entry is also a Кесте right after this one's description
    if i + 1 < len(keste_entries):
        next_pi = keste_entries[i + 1][0]
        next_dt = keste_entries[i + 1][3]
        
        # If two Кесте labels are within 3 paragraphs of each other, they're likely duplicates
        if next_pi - pi <= 3:
            # Check which description matches better - keep the second one if it has the same desc
            if dt == next_dt or dt in next_dt or next_dt in dt:
                # Same or similar description - remove the first one (and its description)
                to_remove.append(pi)
                to_remove.append(di)
                print(f"  REMOVE [{pi}] {kt} (dup of [{next_pi}] {keste_entries[i+1][1]})")
                i += 1
                continue
            else:
                # Different descriptions - the first Кесте is for one table, keep both
                # But this is likely the problem case from the screenshot
                # Remove the first "Кесте X" label that doesn't match
                to_remove.append(pi)
                print(f"  REMOVE [{pi}] {kt} (extra label before [{next_pi}])")
                i += 1
                continue
    
    i += 1

print(f"\nWill remove {len(to_remove)} paragraphs")

# Step 3: Remove duplicate paragraphs
for idx in sorted(to_remove, reverse=True):
    p = doc.paragraphs[idx]
    parent = p._element.getparent()
    if parent is not None:
        # Only remove if it's a short Кесте label or its matching description
        if len(p.text.strip()) < 100:  # safety check
            parent.remove(p._element)
            print(f"  Removed [{idx}]: {p.text.strip()[:60]}")

# Step 4: Renumber all remaining Кесте labels sequentially
print("\n=== STEP 4: Renumbering tables ===")
counter = 1
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    m = re.match(r'^Кесте\s*(\d+)', t)
    if m:
        old_num = m.group(1)
        new_text = re.sub(r'^Кесте\s*\d+', f'Кесте {counter}', t)
        # Update all runs
        for run in p.runs:
            rm = re.match(r'^Кесте\s*\d+', run.text.strip())
            if rm:
                run.text = re.sub(r'Кесте\s*\d+', f'Кесте {counter}', run.text)
        # Get description
        desc = doc.paragraphs[i+1].text.strip()[:50] if i+1 < len(doc.paragraphs) else ''
        print(f"  Кесте {old_num} -> Кесте {counter} ({desc})")
        counter += 1
    
    # Also fix "Кесте – " (unnumbered ones from our additions)
    if re.match(r'^Кесте\s*–\s*', t):
        new_text = re.sub(r'^Кесте\s*–\s*', f'Кесте {counter} – ', t)
        for run in p.runs:
            if re.match(r'^Кесте\s*–', run.text.strip()):
                run.text = re.sub(r'Кесте\s*–\s*', f'Кесте {counter} – ', run.text)
        desc_short = t[len('Кесте – '):50]
        print(f"  Кесте – -> Кесте {counter} ({desc_short})")
        counter += 1

print(f"\nTotal tables: {counter - 1}")

doc.save('Suleimenov_Alisher_VTIPO-45_REPORT_GOST_formatted.docx')
print("\n=== SAVED ===")

# Verify
doc2 = Document('Suleimenov_Alisher_VTIPO-45_REPORT_GOST_formatted.docx')
print("\n=== VERIFICATION ===")
for i, p in enumerate(doc2.paragraphs):
    t = p.text.strip()
    if re.match(r'^Кесте\s*\d+', t):
        desc = doc2.paragraphs[i+1].text.strip()[:60] if i+1 < len(doc2.paragraphs) else ''
        print(f"  [{i}] {t} -> {desc}")

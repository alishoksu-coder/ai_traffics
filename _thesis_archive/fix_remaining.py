# -*- coding: utf-8 -*-
"""Fix remaining issues in GOST document."""
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

INPUT = 'Suleimenov_Alisher_VTIPO-45_REPORT_GOST_formatted_UPDATED.docx'
doc = Document(INPUT)

all_text = '\n'.join(p.text for p in doc.paragraphs)

# Debug: find what the hypothesis actually says now
for i, p in enumerate(doc.paragraphs):
    if 'MAE' in p.text and 'кемінде' in p.text:
        print(f"[HYP {i}] {p.text[:200]}")
    if '30 сурет' in p.text or '32 сурет' in p.text:
        print(f"[FIG {i}] {p.text[:200]}")
    if 'NLP' in p.text:
        print(f"[NLP {i}] {p.text[:200]}")

# Check new tables content
print("\n=== New tables check ===")
for ti, tbl in enumerate(doc.tables):
    for ri, row in enumerate(tbl.rows):
        txt = ' | '.join(c.text.strip() for c in row.cells)
        if any(kw in txt for kw in ['Level-2', 'Байтерек', '38.4', 'RF / AI']):
            print(f"  Table {ti}, Row {ri}: {txt[:120]}")

# Fix hypothesis - find exact paragraph
print("\n=== Fixing remaining issues ===")
for i, p in enumerate(doc.paragraphs):
    text = p.text
    
    # Fix hypothesis
    if 'MAE көрсеткішін кемінде 30%-ға жақсарту' in text and 'інклюзивті' not in text.lower() and 'инклюзивті' not in text.lower():
        new_text = text.rstrip('.') + '; инклюзивті маршруттау модулі стандарт маршруттан орта есеппен 15%-ға ұзын, бірақ 100% кедергісіз жол ұсынады.'
        for run in p.runs:
            run.text = ''
        if p.runs:
            p.runs[0].text = new_text
        else:
            r = p.add_run(new_text)
            r.font.name = 'Times New Roman'
            r.font.size = Pt(14)
        print(f"  Fixed hypothesis at {i}")
    
    # Fix 30 suret -> 32
    if '30 сурет' in text:
        for run in p.runs:
            if '30 сурет' in run.text:
                run.text = run.text.replace('30 сурет', '32 сурет')
                print(f"  Fixed 30->32 suret at {i}")
    
    # Fix remaining NLP/ML
    if 'NLP/ML' in text:
        for run in p.runs:
            if 'NLP/ML' in run.text:
                run.text = run.text.replace('NLP/ML', 'ML')
                print(f"  Fixed NLP/ML at {i}")

doc.save(INPUT)
print("\n=== Fixes applied ===")

# Re-verify critical items
doc2 = Document(INPUT)
all_text2 = '\n'.join(p.text for p in doc2.paragraphs)
checks = [
    ("Гипотеза", "инклюзивті маршруттау модулі"),
    ("NLP жоқ", "NLP/ML" if "NLP/ML" not in all_text2 else "STILL_HAS_NLP"),
    ("32 сурет", "32 сурет"),
]
for name, search in checks:
    found = search in all_text2
    print(f"  {'✓' if found else '✗'} {name}: {'found' if found else 'not found'} '{search[:40]}'")

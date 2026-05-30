# -*- coding: utf-8 -*-
"""Fix ҚОСЫМША А: write each code line as separate paragraph with Courier New."""
import copy
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Cm

INPUT = 'Suleimenov_Alisher_VTIPO-45_REPORT_GOST_formatted.docx'
doc = Document(INPUT)

# Find ҚОСЫМША А and remove old content after it
qosymsha_idx = -1
for i, p in enumerate(doc.paragraphs):
    if 'ҚОСЫМША А' in p.text.strip():
        qosymsha_idx = i
        break

if qosymsha_idx < 0:
    print("ERROR: ҚОСЫМША А not found!")
    exit()

# Remove everything after ҚОСЫМША А
removed = 0
for i in range(len(doc.paragraphs) - 1, qosymsha_idx, -1):
    elem = doc.paragraphs[i]._element
    parent = elem.getparent()
    if parent is not None:
        parent.remove(elem)
        removed += 1

print(f"Removed {removed} old paragraphs after ҚОСЫМША А")

# Helper: Get body text paragraph style (from existing content)
body_pPr = None
for i in range(200, 300):
    p = doc.paragraphs[i]
    if p.text.strip() and len(p.text.strip()) > 50:
        pPr = p._element.find(qn('w:pPr'))
        if pPr is not None:
            body_pPr = copy.deepcopy(pPr)
        break

def add_paragraph(ref_elem, text, is_title=False, is_code=False):
    """Add paragraph after ref_elem with proper formatting."""
    new_p = OxmlElement('w:p')
    ref_elem.addnext(new_p)
    
    # Paragraph properties
    pPr = OxmlElement('w:pPr')
    new_p.insert(0, pPr)
    
    if is_title:
        # Title: centered, bold, Times New Roman 14pt, space before
        jc = OxmlElement('w:jc')
        jc.set(qn('w:val'), 'left')
        pPr.append(jc)
        sp = OxmlElement('w:spacing')
        sp.set(qn('w:before'), '240')
        sp.set(qn('w:after'), '120')
        sp.set(qn('w:line'), '360')
        sp.set(qn('w:lineRule'), 'auto')
        pPr.append(sp)
    elif is_code:
        # Code: left-aligned, Courier New 10pt, no indent, tight spacing
        jc = OxmlElement('w:jc')
        jc.set(qn('w:val'), 'left')
        pPr.append(jc)
        ind = OxmlElement('w:ind')
        ind.set(qn('w:left'), '567')  # ~1cm
        ind.set(qn('w:firstLine'), '0')
        pPr.append(ind)
        sp = OxmlElement('w:spacing')
        sp.set(qn('w:before'), '0')
        sp.set(qn('w:after'), '0')
        sp.set(qn('w:line'), '240')  # single spacing
        sp.set(qn('w:lineRule'), 'auto')
        pPr.append(sp)
    
    # Run with text
    run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    
    if is_title:
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), 'Times New Roman')
        rFonts.set(qn('w:hAnsi'), 'Times New Roman')
        rFonts.set(qn('w:cs'), 'Times New Roman')
        rPr.append(rFonts)
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), '28')  # 14pt
        rPr.append(sz)
        szCs = OxmlElement('w:szCs')
        szCs.set(qn('w:val'), '28')
        rPr.append(szCs)
        b = OxmlElement('w:b')
        rPr.append(b)
    elif is_code:
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), 'Courier New')
        rFonts.set(qn('w:hAnsi'), 'Courier New')
        rFonts.set(qn('w:cs'), 'Courier New')
        rPr.append(rFonts)
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), '20')  # 10pt
        rPr.append(sz)
        szCs = OxmlElement('w:szCs')
        szCs.set(qn('w:val'), '20')
        rPr.append(szCs)
    
    run.insert(0, rPr)
    t = OxmlElement('w:t')
    t.set(qn('xml:space'), 'preserve')
    t.text = text
    run.append(t)
    new_p.append(run)
    
    return new_p

# Code sections
code_sections = [
    ("А.1 – predict.py: Болжам модулі", [
        "def predict_traffic(point_id, horizon, history):",
        "    models = {",
        "        'naive': naive_forecast(history),",
        "        'sma': sma_forecast(history, k=5),",
        "        'ema': ema_forecast(history, alpha=0.3),",
        "        'trend_lr': trend_lr_forecast(history, k=10),",
        "    }",
        "    return models",
        "",
        "def sma_forecast(history, k=5):",
        "    if len(history) < k:",
        "        return history[-1] if history else 50",
        "    return sum(history[-k:]) / k",
        "",
        "def ema_forecast(history, alpha=0.3):",
        "    ema = history[0]",
        "    for val in history[1:]:",
        "        ema = alpha * val + (1 - alpha) * ema",
        "    return ema",
    ]),
    ("А.2 – anomaly_detector.py: Аномалия детекциясы", [
        "def detect_anomaly(recent_values):",
        "    if len(recent_values) < 3:",
        "        return None",
        "    diff = recent_values[-1] - recent_values[-2]",
        "",
        "    # Level-1: Жедел секіріс",
        "    if diff > 25 and recent_values[-1] > 70:",
        "        return {'level': 1, 'message': 'ЖКО анықталды!'}",
        "",
        "    # Level-2: Жалпы коллапс",
        "    total_rise = recent_values[-1] - recent_values[0]",
        "    if total_rise > 35 or recent_values[-1] > 90:",
        "        return {'level': 2, 'message': 'Жалпы коллапс'}",
        "",
        "    # Level-3: Тез өсу",
        "    if total_rise > 20:",
        "        return {'level': 3, 'message': 'Жүктеме өсуде'}",
        "",
        "    return None",
    ]),
    ("А.3 – ai_brain.py: Random Forest модулі", [
        "from sklearn.ensemble import RandomForestRegressor",
        "import joblib",
        "",
        "def train_model(X_train, y_train):",
        "    model = RandomForestRegressor(",
        "        n_estimators=100, max_depth=12,",
        "        min_samples_split=5, random_state=42",
        "    )",
        "    model.fit(X_train, y_train)",
        "    joblib.dump(model, 'rf_model.joblib')",
        "    return model",
        "",
        "def predict_rf(features):",
        "    model = joblib.load('rf_model.joblib')",
        "    return model.predict([features])[0]",
        "",
        "def get_feature_importance(model):",
        "    names = ['hour','day_of_week','weather','segment_id']",
        "    return dict(zip(names, model.feature_importances_))",
    ]),
    ("А.4 – simulator.py: Трафик симуляторы", [
        "import math, random",
        "",
        "def compute_base_load(hour):",
        "    morning = 40 * math.exp(-((hour - 8.5)**2) / 3)",
        "    evening = 45 * math.exp(-((hour - 18)**2) / 4)",
        "    return 15 + morning + evening",
        "",
        "def apply_weather(base_load, weather_factor):",
        "    return base_load * (1 + weather_factor * 0.3)",
        "",
        "def tick_simulation(locations, hour, weather):",
        "    for loc in locations:",
        "        base = compute_base_load(hour)",
        "        load = apply_weather(base, weather)",
        "        noise = random.gauss(0, 3)",
        "        loc['current_load'] = max(0, min(100, load+noise))",
        "    return locations",
    ]),
    ("А.5 – main.py: FastAPI серверлік логика", [
        "from fastapi import FastAPI",
        "from fastapi.middleware.cors import CORSMiddleware",
        "",
        "app = FastAPI(title='AI Traffic API', version='2.0')",
        "app.add_middleware(CORSMiddleware, allow_origins=['*'])",
        "",
        "@app.get('/traffic/map')",
        "async def traffic_map(horizon: int = 0):",
        "    locations = get_all_locations()",
        "    if horizon > 0:",
        "        for loc in locations:",
        "            loc['load'] = predict_traffic(",
        "                loc['id'], horizon, loc['history'])",
        "    return {'locations': locations, 'count': len(locations)}",
        "",
        "@app.post('/roads/barrier-free')",
        "async def barrier_free_route(start: dict, end: dict):",
        "    route = compute_bfr_route(start, end)",
        "    return {'route': route, 'accessible': True}",
    ]),
]

# Insert all code sections
ref = doc.paragraphs[qosymsha_idx]._element
total_lines = 0

for title, lines in code_sections:
    # Add title
    ref = add_paragraph(ref, title, is_title=True)
    
    # Add each code line as separate paragraph
    for line in lines:
        ref = add_paragraph(ref, line if line else ' ', is_code=True)
        total_lines += 1

print(f"Added {len(code_sections)} sections, {total_lines} code lines")

doc.save(INPUT)
print(f"SAVED to {INPUT}")

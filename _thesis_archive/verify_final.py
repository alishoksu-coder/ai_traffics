# -*- coding: utf-8 -*-
import sys, io, re
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document

doc = Document('Suleimenov_Alisher_VTIPO-45_REPORT_GOST_formatted.docx')

print("=== FULL STRUCTURE VERIFICATION ===\n")

# 1. All section headings
print("--- Section headings ---")
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if re.match(r'^[123]\.\d*\s', t) and i > 200 and len(t) < 120:
        print(f"  {t}")
    elif t in ('Қорытынды', 'ҚОРЫТЫНДЫ', 'ҚОСЫМША А', 'КІРІСПЕ'):
        if i > 200:
            print(f"  {t}")

# 2. Figure numbering
print("\n--- Figure numbering ---")
fig_nums = []
for p in doc.paragraphs:
    m = re.match(r'^Сурет\s+(\d+)', p.text.strip())
    if m:
        fig_nums.append(int(m.group(1)))
print(f"  Sequence: {fig_nums}")
print(f"  Sequential: {fig_nums == list(range(1, len(fig_nums)+1))}")

# 3. Table numbering
print("\n--- Table numbering ---")
tbl_nums = []
for p in doc.paragraphs:
    m = re.match(r'^Кесте\s+(\d+)', p.text.strip())
    if m:
        tbl_nums.append(int(m.group(1)))
print(f"  Sequence: {tbl_nums}")
print(f"  Sequential: {tbl_nums == list(range(1, len(tbl_nums)+1))}")

# 4. Word counts for new sections
print("\n--- Word counts per section ---")
current_section = None
word_count = 0
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    m = re.match(r'^(2\.\d+|3\.\d+)\s', t)
    if m and i > 400:
        if current_section:
            print(f"  {current_section}: ~{word_count} words")
        current_section = t[:70]
        word_count = 0
    elif t.startswith('Қорытынды') and i > 1100:
        if current_section:
            print(f"  {current_section}: ~{word_count} words")
        break
    elif current_section and t:
        word_count += len(t.split())

print(f"\n--- Summary ---")
print(f"Total paragraphs: {len(doc.paragraphs)}")
print(f"Total images: {len(doc.inline_shapes)}")
print(f"Total figures: {len(fig_nums)}")
print(f"Total tables: {len(tbl_nums)}")

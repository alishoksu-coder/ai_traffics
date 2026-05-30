# -*- coding: utf-8 -*-
import sys, io, re, shutil, os
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document
from docx.oxml.ns import qn

# Restore from backup first
BACKUP = 'Suleimenov_Alisher_VTIPO-45_REPORT_GOST_formatted_PRE_FINAL_BACKUP.docx'
SRC = 'Suleimenov_Alisher_VTIPO-45_REPORT_GOST_formatted.docx'

if os.path.exists(BACKUP):
    shutil.copy2(BACKUP, SRC)
    print(f"Restored from backup: {BACKUP}")

doc = Document(SRC)
body = doc.element.body

# ============================================================
# STEP 1: Fix Figure numbering (Сурет)
# ============================================================
print("\n=== STEP 1: Fix figure numbering ===")
fig_counter = 0
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    match = re.match(r'^(Сурет)\s+(\d+)', text)
    if match:
        old_num = int(match.group(2))
        fig_counter += 1
        if old_num != fig_counter:
            old_pat = f'Сурет {old_num}'
            new_pat = f'Сурет {fig_counter}'
            for run in p.runs:
                if old_pat in run.text:
                    run.text = run.text.replace(old_pat, new_pat, 1)
                    print(f"  Fig: {old_pat} -> {new_pat} (para {i})")
                    break
            else:
                for run in p.runs:
                    if str(old_num) in run.text:
                        run.text = run.text.replace(str(old_num), str(fig_counter), 1)
                        print(f"  Fig cross-run: {old_pat} -> {new_pat} (para {i})")
                        break
print(f"Total figures: {fig_counter}")

# ============================================================
# STEP 2: Fix Table numbering (Кесте)
# ============================================================
print("\n=== STEP 2: Fix table numbering ===")
tbl_counter = 0
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    match = re.match(r'^(Кесте)\s+(\d+)', text)
    if match:
        old_num = int(match.group(2))
        tbl_counter += 1
        if old_num != tbl_counter:
            old_pat = f'Кесте {old_num}'
            new_pat = f'Кесте {tbl_counter}'
            for run in p.runs:
                if old_pat in run.text:
                    run.text = run.text.replace(old_pat, new_pat, 1)
                    print(f"  Tbl: {old_pat} -> {new_pat} (para {i})")
                    break
            else:
                for run in p.runs:
                    if str(old_num) in run.text:
                        run.text = run.text.replace(str(old_num), str(tbl_counter), 1)
                        print(f"  Tbl cross-run: {old_pat} -> {new_pat} (para {i})")
                        break
print(f"Total tables: {tbl_counter}")

# ============================================================
# STEP 3: Fix inline refs
# ============================================================
print("\n=== STEP 3: Fix inline references ===")
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if not text.startswith('Сурет') and 'Сурет' in text:
        # Para ~948: "(Сурет 8)" -> web dashboard = Сурет 15
        if 'Үш негізгі беттен тұрады' in text:
            for run in p.runs:
                if 'Сурет 8' in run.text:
                    run.text = run.text.replace('Сурет 8', 'Сурет 15')
                    print(f"  Inline fix para {i}: Сурет 8 -> Сурет 15")
                    break
        # Para ~1046: "(Сурет 14)" -> MAE chart = Сурет 22
        if 'MAE=5.34' in text or 'Naive baseline' in text:
            for run in p.runs:
                if 'Сурет 14' in run.text:
                    run.text = run.text.replace('Сурет 14', 'Сурет 22')
                    print(f"  Inline fix para {i}: Сурет 14 -> Сурет 22")
                    break

# ============================================================
# STEP 4: Find REAL conclusion (not in TOC)
# ============================================================
print("\n=== STEP 4: Find real conclusion position ===")
conclusion_idx = None
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    # Real conclusion is a heading after paragraph 1000+
    if i > 900 and (text == 'Қорытынды' or text == 'ҚОРЫТЫНДЫ'):
        conclusion_idx = i
        print(f"Found real conclusion at para {i}: {text}")
        break

if not conclusion_idx:
    # Try broader search
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if i > 500 and text.startswith('Қорытынды') and len(text) < 30:
            conclusion_idx = i
            print(f"Found conclusion at para {i}: {text}")
            break

# ============================================================
# STEP 5: Add missing presentation content before conclusion
# ============================================================
print("\n=== STEP 5: Add missing content ===")

existing = set()
for p in doc.paragraphs:
    t = p.text.strip().lower()
    if 'контейнерлік архитектура' in t: existing.add('container')
    if 'клиенттік деңгей' in t or 'client layer' in t: existing.add('client')
    if 'серверлік деңгей' in t or 'backend bridge' in t: existing.add('server')
    if 'болжау архитектурасы' in t or 'prediction flow' in t: existing.add('predict')
    if 'smart alert' in t or 'интеллектуалды хабарлама' in t: existing.add('alerts')
    if 'жүйелік модуль' in t: existing.add('modules')
    if 'модельдердің жауапкершілік' in t: existing.add('responsibility')

print(f"Existing topics: {existing}")

blocks = []
if 'container' not in existing:
    blocks.append(('2.8 Жүйенің контейнерлік архитектурасы', [
        'AI Traffic жүйесі контейнерлік архитектураға негізделген. Әрбір компонент тәуелсіз модуль ретінде жұмыс істейді, бұл жүйені масштабтау мен жаңартуды жеңілдетеді.',
        'Жүйенің негізгі контейнерлері: Flutter Mobile App — кросс-платформалық мобильді қосымша; FastAPI Backend — жоғары жылдамдықты API сервері; PostgreSQL + PostGIS — кеңістіктік деректер қоры; ML Pipeline — LSTM және Random Forest модельдері; Web Dashboard — Leaflet.js мониторинг тақтасы; WebSocket Server — нақты уақыттағы деректер алмасу.',
        'Контейнерлік архитектура C4 моделі бойынша құрылған: Context → Container → Component → Code.'
    ]))
if 'client' not in existing:
    blocks.append(('2.9 Клиенттік деңгей (Client Layer)', [
        'Клиенттік деңгей — пайдаланушылар мен жүйе арасындағы негізгі интерфейстік қабат. Flutter фреймворкі мен Dart тілі арқылы бір кодтық базадан жоғары өнімділік қамтамасыз етіледі.',
        'Негізгі компоненттері: Real-time Map Matching — GPS деректерін жол желісімен сәйкестендіру; User Notifications — MQTT/WebSocket арқылы кептелістер туралы хабарлама; Advanced Monitoring — толық мониторинг; Predictive Visuals — ИИ болжаған кептеліс ықтималдығын визуалды көрсету.',
    ]))
if 'server' not in existing:
    blocks.append(('2.10 Серверлік деңгей (Backend Bridge)', [
        'Серверлік деңгей FastAPI фреймворкі негізінде құрылған. Асинхронды архитектура арқылы көп ағынды сұраныстарды бір уақытта өңдеуге мүмкіндік береді.',
        'Негізгі функциялары: көп ағынды асинхронды сұраныстарды өңдеу; деректерді Pydantic арқылы валидациялау; WebSocket арқылы нақты уақыттағы байланыс; ML модельдермен интеграция. Backend Render.com бұлтында deploy етілген.',
    ]))
if 'predict' not in existing:
    blocks.append(('2.11 ML болжау архитектурасы (Prediction Flow)', [
        'ML болжау архитектурасы — жүйенің интеллектуалды ядросы. Процесс: 1) PostgreSQL-ден тарихи деректерді алу; 2) MinMaxScaler нормализация; 3) LSTM/Random Forest моделіне беру; 4) 60 минуттық горизонтқа болжам; 5) REST API және WebSocket арқылы клиенттерге тарату.',
        'Ensemble тәсілі қолданылады: LSTM уақыттық тізбек үлгілерін анықтайды, Random Forest қосымша факторларды ескереді. AI Brain модулі екі нәтижені біріктіреді.',
    ]))
if 'alerts' not in existing:
    blocks.append(('2.12 Интеллектуалды хабарламалар жүйесі (Smart Alerts)', [
        'Smart Alerts модулі ML болжамдарына негізделіп, пайдаланушыларға алдын ала ескерту хабарламаларын жібереді.',
        'Функциялары: кептеліс болжамы бойынша 30-60 минут бұрын ескерту; ауа-райы факторларын ескерту; аномалия детекциясы кезінде жедел хабарлама; баламалы маршрут ұсыныстары. WebSocket протоколы арқылы нақты уақытта жұмыс істейді.',
    ]))
if 'modules' not in existing:
    blocks.append(('2.13 Жүйелік модульдер', [
        'AI Traffic жүйесі модульдік архитектура бойынша құрылған. Негізгі модульдер: Traffic Simulator — 144 нүктеде трафик симуляциясы; Prediction Engine — LSTM + RF ансамбль; Anomaly Detector — Z-score аномалия детекциясы; Weather Module — OpenWeatherMap API; Routing Engine — A* алгоритмі (CarFast, BarrierFree, AntiStress); Admin Panel — әкімшілік басқару; Crowdsourcing Module — пайдаланушылардан деректер жинау.',
    ]))
if 'responsibility' not in existing:
    blocks.append(('2.14 Модельдердің жауапкершілік аймақтары', [
        'Жүйедегі ML модельдері: LSTM — уақыттық тізбектерді талдау (87% дәлдік); Random Forest — қосымша факторларды ескеру; Linear Regression — базалық салыстыру; AI Brain (Ensemble) — нәтижелерді біріктіру (MAE=5.34); Z-Score Anomaly Detector — аномалияларды анықтау.',
        'Әрбір модельдің нәтижесі /predict, /anomalies, /recommendations API эндпоинттері арқылы қолжетімді.',
    ]))

if conclusion_idx and blocks:
    conclusion_el = doc.paragraphs[conclusion_idx]._element
    
    for heading, paras in reversed(blocks):
        # Blank separator
        bp = body.makeelement(qn('w:p'), {})
        conclusion_el.addprevious(bp)
        
        # Body paragraphs (reversed)
        for pt in reversed(paras):
            np = body.makeelement(qn('w:p'), {})
            pPr = np.makeelement(qn('w:pPr'), {})
            np.append(pPr)
            sp = pPr.makeelement(qn('w:spacing'), {qn('w:line'): '360', qn('w:lineRule'): 'auto'})
            pPr.append(sp)
            jc = pPr.makeelement(qn('w:jc'), {qn('w:val'): 'both'})
            pPr.append(jc)
            ind = pPr.makeelement(qn('w:ind'), {qn('w:firstLine'): '709'})
            pPr.append(ind)
            
            r = np.makeelement(qn('w:r'), {})
            rPr = r.makeelement(qn('w:rPr'), {})
            rf = rPr.makeelement(qn('w:rFonts'), {qn('w:ascii'): 'Times New Roman', qn('w:hAnsi'): 'Times New Roman', qn('w:cs'): 'Times New Roman'})
            rPr.append(rf)
            sz = rPr.makeelement(qn('w:sz'), {qn('w:val'): '28'})
            rPr.append(sz)
            sz2 = rPr.makeelement(qn('w:szCs'), {qn('w:val'): '28'})
            rPr.append(sz2)
            r.append(rPr)
            t = r.makeelement(qn('w:t'), {})
            t.text = pt
            t.set(qn('xml:space'), 'preserve')
            r.append(t)
            np.append(r)
            conclusion_el.addprevious(np)
        
        # Heading
        hp = body.makeelement(qn('w:p'), {})
        hPr = hp.makeelement(qn('w:pPr'), {})
        ps = hPr.makeelement(qn('w:pStyle'), {qn('w:val'): 'Heading2'})
        hPr.append(ps)
        hp.append(hPr)
        hr = hp.makeelement(qn('w:r'), {})
        hrPr = hr.makeelement(qn('w:rPr'), {})
        hb = hrPr.makeelement(qn('w:b'), {})
        hrPr.append(hb)
        hrf = hrPr.makeelement(qn('w:rFonts'), {qn('w:ascii'): 'Times New Roman', qn('w:hAnsi'): 'Times New Roman'})
        hrPr.append(hrf)
        hsz = hrPr.makeelement(qn('w:sz'), {qn('w:val'): '28'})
        hrPr.append(hsz)
        hr.append(hrPr)
        ht = hr.makeelement(qn('w:t'), {})
        ht.text = heading
        hr.append(ht)
        hp.append(hr)
        conclusion_el.addprevious(hp)
    
    print(f"Added {len(blocks)} sections")
else:
    print(f"Conclusion idx: {conclusion_idx}, blocks: {len(blocks)}")

# ============================================================
# Save & verify
# ============================================================
doc.save(SRC)
print(f"\nSaved: {SRC}")

# Verify
doc2 = Document(SRC)
fig_nums = []
tbl_nums = []
for p in doc2.paragraphs:
    txt = p.text.strip()
    fm = re.match(r'^Сурет\s+(\d+)', txt)
    tm = re.match(r'^Кесте\s+(\d+)', txt)
    if fm: fig_nums.append(fm.group(1))
    if tm: tbl_nums.append(tm.group(1))
print(f"\nFigure sequence: {fig_nums}")
print(f"Table sequence: {tbl_nums}")
print(f"Total paragraphs: {len(doc2.paragraphs)}, images: {len(doc2.inline_shapes)}")
print("DONE!")

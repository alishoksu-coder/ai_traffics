# -*- coding: utf-8 -*-
"""Add (3) number to Z-score formula."""
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

doc = Document('Suleimenov_Alisher_VTIPO-45_REPORT_GOST_formatted.docx')

# Add (3) to Z-score formula at paragraph 352
p = doc.paragraphs[352]
p_elem = p._element

# Add tab stop
pPr = p_elem.find(qn('w:pPr'))
if pPr is None:
    pPr = OxmlElement('w:pPr')
    p_elem.insert(0, pPr)

tabs = pPr.find(qn('w:tabs'))
if tabs is None:
    tabs = OxmlElement('w:tabs')
    pPr.append(tabs)
tab = OxmlElement('w:tab')
tab.set(qn('w:val'), 'right')
tab.set(qn('w:pos'), '9639')
tab.set(qn('w:leader'), 'none')
tabs.append(tab)

# Tab run
tab_run = OxmlElement('w:r')
tab_t = OxmlElement('w:tab')
tab_run.append(tab_t)
p_elem.append(tab_run)

# Number run: (3)
num_run = OxmlElement('w:r')
rPr = OxmlElement('w:rPr')
rFonts = OxmlElement('w:rFonts')
rFonts.set(qn('w:ascii'), 'Times New Roman')
rFonts.set(qn('w:hAnsi'), 'Times New Roman')
rPr.append(rFonts)
sz = OxmlElement('w:sz')
sz.set(qn('w:val'), '28')
rPr.append(sz)
num_run.append(rPr)
t = OxmlElement('w:t')
t.set(qn('xml:space'), 'preserve')
t.text = '(3)'
num_run.append(t)
p_elem.append(num_run)

# List all OMML formulas
OMML_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/math'
print('=== All OMML formulas ===')
for i, p in enumerate(doc.paragraphs):
    for child in p._element:
        if 'oMath' in child.tag:
            txt = p.text.strip() if p.text.strip() else '(math object)'
            print(f'  [{i}] {txt[:60]}')
            break

doc.save('Suleimenov_Alisher_VTIPO-45_REPORT_GOST_formatted.docx')
print('\nDone! Added (3) to Z-score. SAVED.')

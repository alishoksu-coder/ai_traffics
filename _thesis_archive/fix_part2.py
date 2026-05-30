# -*- coding: utf-8 -*-
"""
Complete fix: inline references + add missing presentation content
"""
import sys, io, re, os, shutil
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

SRC = 'Suleimenov_Alisher_VTIPO-45_REPORT_GOST_formatted.docx'

doc = Document(SRC)

# ============================================================
# STEP 1: Fix inline figure references in body text
# ============================================================
print("=" * 60)
print("STEP 1: Fixing inline figure references...")
print("=" * 60)

# After renumbering, the inline references still have old numbers.
# Fix specific known references:
# Para 948: "(Сурет 8)" should be (Сурет 15) - Web Dashboard
# Para 1046: "(Сурет 14)" should be (Сурет 22) - MAE/RMSE chart

fixes = [
    (948, 'Сурет 8', 'Сурет 15'),
    (1046, 'Сурет 14', 'Сурет 22'),
]

for para_idx, old_ref, new_ref in fixes:
    p = doc.paragraphs[para_idx]
    for run in p.runs:
        if old_ref in run.text:
            run.text = run.text.replace(old_ref, new_ref)
            print(f"  FIXED Para {para_idx}: {old_ref} -> {new_ref}")
            break
    else:
        # Check if the reference was already fixed in previous run
        if old_ref in p.text:
            print(f"  Para {para_idx}: Found {old_ref} but couldn't fix (cross-run)")
        else:
            current_refs = re.findall(r'Сурет\s+\d+', p.text)
            print(f"  Para {para_idx}: Already fixed or not found. Current refs: {current_refs}")

# ============================================================
# STEP 2: Add missing content from presentation
# ============================================================
print("\n" + "=" * 60)
print("STEP 2: Adding missing presentation content...")
print("=" * 60)

# Check what's already present
existing_topics = set()
for p in doc.paragraphs:
    text = p.text.strip().lower()
    if 'контейнерлік архитектура' in text: existing_topics.add('container_arch')
    if 'клиенттік деңгей' in text or 'client layer' in text: existing_topics.add('client_layer')
    if 'серверлік деңгей' in text or 'backend bridge' in text: existing_topics.add('server_layer')
    if 'болжау архитектурасы' in text or 'prediction flow' in text: existing_topics.add('prediction_flow')
    if 'трафик симулятор' in text: existing_topics.add('simulator')
    if 'lstm' in text: existing_topics.add('lstm')
    if 'smart alert' in text or 'интеллектуалды хабарлама' in text: existing_topics.add('smart_alerts')
    if 'краудсорсинг' in text or 'crowdsourcing' in text: existing_topics.add('crowdsourcing')
    if 'faceid' in text or 'деректер қауіпсіздігі' in text: existing_topics.add('security')
    if 'жүйелік модуль' in text: existing_topics.add('system_modules')
    if 'модельдердің жауапкершілік' in text: existing_topics.add('model_responsibility')

print(f"Topics already in document: {existing_topics}")

# Find conclusion
conclusion_idx = None
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if text == 'Қорытынды' or text == 'ҚОРЫТЫНДЫ' or text.startswith('Қорытынды'):
        conclusion_idx = i
        break

print(f"Conclusion at paragraph: {conclusion_idx}")

# Build content blocks for missing topics
content_blocks = []

if 'container_arch' not in existing_topics:
    content_blocks.append({
        'heading': '2.8 Жүйенің контейнерлік архитектурасы',
        'paragraphs': [
            'AI Traffic жүйесі контейнерлік архитектураға негізделген, яғни әрбір компонент тәуелсіз модуль (контейнер) ретінде жұмыс істейді. Бұл тәсіл жүйені масштабтау, жаңарту және техникалық қызмет көрсетуді айтарлықтай жеңілдетеді.',
            'Жүйенің негізгі контейнерлері:',
            '• Flutter Mobile App — кросс-платформалық мобильді қосымша, пайдаланушы интерфейсі;',
            '• FastAPI Backend — жоғары жылдамдықты API сервері, бизнес-логика;',
            '• PostgreSQL + PostGIS — кеңістіктік деректер қоры;',
            '• ML Pipeline — LSTM және Random Forest модельдерін оқыту және болжау қызметтері;',
            '• Web Dashboard — Leaflet.js негізіндегі мониторинг тақтасы;',
            '• WebSocket Server — нақты уақыттағы деректер алмасу.',
            'Контейнерлік архитектура C4 моделі бойынша құрылған: Context → Container → Component → Code. Бұл тәсіл жүйе компоненттерін тәуелсіз deploy етуге мүмкіндік береді.'
        ]
    })

if 'client_layer' not in existing_topics:
    content_blocks.append({
        'heading': '2.9 Клиенттік деңгей (Client Layer)',
        'paragraphs': [
            'Клиенттік деңгей — пайдаланушылар мен жүйе арасындағы негізгі интерфейстік қабат. Flutter фреймворкі мен Dart тілі арқылы бір кодтық базадан жоғары өнімділікті қамтамасыз етіледі.',
            'Клиенттік деңгейдің негізгі компоненттері:',
            '• Real-time Map Matching — GPS деректерін жол желісімен дәл сәйкестендіру;',
            '• User Notifications — MQTT/WebSocket арқылы кептелістер туралы жедел хабарлама;',
            '• Advanced Monitoring — қала басшылығы мен операторларына арналған толық мониторинг;',
            '• Predictive Visuals — ИИ болжаған кептеліс ықтималдығын визуалды түрде көрсету.',
            'Клиенттік деңгейде Material Design 3 принциптері мен glassmorphism стилі қолданылған, бұл қосымшаға заманауи және тартымды көрініс береді.'
        ]
    })

if 'server_layer' not in existing_topics:
    content_blocks.append({
        'heading': '2.10 Серверлік деңгей (Backend Bridge)',
        'paragraphs': [
            'Серверлік деңгей FastAPI фреймворкі негізінде құрылған және жүйенің бизнес-логикасын басқарады. Асинхронды архитектура арқылы көп ағынды сұраныстарды бір уақытта өңдеуге мүмкіндік береді.',
            'Серверлік деңгейдің негізгі функциялары:',
            '• Көп ағынды асинхронды сұраныстарды өңдеу;',
            '• Деректердің дұрыстығын автоматты валидациялау (Pydantic);',
            '• Клиенттермен WebSocket арқылы нақты уақыттағы байланыс;',
            '• ML модельдерімен интеграция — болжау нәтижелерін API арқылы жеткізу.',
            'Backend сервері Render.com бұлттық платформасында deploy етілген, бұл жүйенің тұрақты жұмыс істеуін және глобалды қол жетімділігін қамтамасыз етеді.'
        ]
    })

if 'prediction_flow' not in existing_topics:
    content_blocks.append({
        'heading': '2.11 ML болжау архитектурасы (Prediction Flow)',
        'paragraphs': [
            'ML болжау архитектурасы — AI Traffic жүйесінің интеллектуалды ядросы. Болжау процесі бірнеше кезеңнен тұрады:',
            '1. Деректерді жинау — PostgreSQL деректер қорынан тарихи трафик деректерін алу;',
            '2. Алдын ала өңдеу — MinMaxScaler арқылы нормализация, feature engineering;',
            '3. Модельге беру — LSTM немесе Random Forest моделіне дайындалған деректерді жіберу;',
            '4. Болжам жасау — модель 60 минуттық горизонтқа болжам шығарады;',
            '5. Нәтижені жеткізу — REST API және WebSocket арқылы клиенттерге тарату.',
            'Болжау архитектурасы ensemble тәсілін қолданады: LSTM уақыттық тізбек үлгілерін анықтайды, ал Random Forest қосымша факторларды (ауа-райы, апталық циклдар) ескереді. AI Brain модулі екі модельдің нәтижесін біріктіреді.'
        ]
    })

if 'smart_alerts' not in existing_topics:
    content_blocks.append({
        'heading': '2.12 Интеллектуалды хабарламалар жүйесі (Smart Alerts)',
        'paragraphs': [
            'Smart Alerts модулі — AI Traffic жүйесінің маңызды инновациялық компоненті. Бұл модуль ML болжамдарына негізделіп, пайдаланушыларға алдын ала ескерту хабарламаларын жібереді.',
            'Модульдің негізгі функциялары:',
            '• Кептеліс болжамы бойынша алдын ала ескерту — жүйе 30-60 минут бұрын ықтимал кептеліс туралы хабарлайды;',
            '• Ауа-райы факторларын ескерту — қар, жаңбыр, мұз жағдайларында қосымша ескертулер;',
            '• Апат детекциясы — аномалия анықталғанда жедел хабарлама;',
            '• Маршрут ұсыныстары — баламалы маршруттар туралы автоматты ұсыныстар.',
            'Smart Alerts жүйесі WebSocket протоколы арқылы нақты уақытта жұмыс істейді. Хабарламалар priority деңгейіне байланысты low, medium, high және critical болып бөлінеді.'
        ]
    })

if 'system_modules' not in existing_topics:
    content_blocks.append({
        'heading': '2.13 Жүйелік модульдер',
        'paragraphs': [
            'AI Traffic жүйесі модульдік архитектура бойынша құрылған. Әрбір модуль тәуелсіз жұмыс істей алады және бір-бірімен API арқылы байланысады.',
            'Жүйенің негізгі модульдері:',
            '• Traffic Simulator — нақты уақыттағы трафик симуляциясы, 144 мониторинг нүктесі;',
            '• Prediction Engine — LSTM + Random Forest ансамбль моделі;',
            '• Anomaly Detector — Z-score негізіндегі аномалия детекциясы;',
            '• Weather Module — ауа-райы деректерін OpenWeatherMap API арқылы алу;',
            '• Routing Engine — A* алгоритмі негізіндегі маршруттау (CarFast, BarrierFree, AntiStress);',
            '• Admin Panel — әкімшілік мониторинг және басқару модулі;',
            '• Crowdsourcing Module — пайдаланушылардан жол жағдайы туралы деректер жинау.',
            'Модульдік архитектура жүйені масштабтау мен жаңартуды жеңілдетеді. Жаңа модуль қосу үшін тек API интерфейсін сақтау жеткілікті.'
        ]
    })

if 'model_responsibility' not in existing_topics:
    content_blocks.append({
        'heading': '2.14 Модельдердің жауапкершілік аймақтары',
        'paragraphs': [
            'AI Traffic жүйесінде бірнеше ML моделі қолданылады, әрқайсысының өз жауапкершілік аймағы бар:',
            '• LSTM (Long Short-Term Memory) — уақыттық тізбектерді талдау, 87% дәлдік. Негізгі міндеті: сағаттық және тәуліктік трафик циклдарын болжау;',
            '• Random Forest — қосымша факторларды (ауа-райы, мереке, апта күні) ескере отырып болжау жасау;',
            '• Linear Regression — жалпы трендті анықтау, базалық салыстыру моделі;',
            '• AI Brain (Ensemble) — LSTM мен Random Forest нәтижелерін біріктіру, MAE=5.34;',
            '• Z-Score Anomaly Detector — трафик ағынындағы аномалияларды анықтау.',
            'Әрбір модельдің нәтижесі API арқылы /predict, /anomalies, /recommendations эндпоинттері арқылы қолжетімді.'
        ]
    })

# Insert content before conclusion
if conclusion_idx and content_blocks:
    body = doc.element.body
    conclusion_element = doc.paragraphs[conclusion_idx]._element
    
    for block in reversed(content_blocks):
        # Add blank line separator
        blank_p = body.makeelement(qn('w:p'), {})
        conclusion_element.addprevious(blank_p)
        
        # Add body paragraphs in reverse order
        for para_text in reversed(block['paragraphs']):
            new_p = body.makeelement(qn('w:p'), {})
            pPr = new_p.makeelement(qn('w:pPr'), {})
            new_p.append(pPr)
            
            # Line spacing 1.5
            spacing = pPr.makeelement(qn('w:spacing'), {
                qn('w:line'): '360',
                qn('w:lineRule'): 'auto'
            })
            pPr.append(spacing)
            
            # Justified
            jc = pPr.makeelement(qn('w:jc'), {qn('w:val'): 'both'})
            pPr.append(jc)
            
            # First line indent
            ind = pPr.makeelement(qn('w:ind'), {qn('w:firstLine'): '709'})
            pPr.append(ind)
            
            # Create run with formatted text
            run = new_p.makeelement(qn('w:r'), {})
            rPr = run.makeelement(qn('w:rPr'), {})
            
            rFonts = rPr.makeelement(qn('w:rFonts'), {
                qn('w:ascii'): 'Times New Roman',
                qn('w:hAnsi'): 'Times New Roman',
                qn('w:cs'): 'Times New Roman'
            })
            rPr.append(rFonts)
            
            sz = rPr.makeelement(qn('w:sz'), {qn('w:val'): '28'})
            rPr.append(sz)
            szCs = rPr.makeelement(qn('w:szCs'), {qn('w:val'): '28'})
            rPr.append(szCs)
            
            run.append(rPr)
            
            t = run.makeelement(qn('w:t'), {})
            t.text = para_text
            t.set(qn('xml:space'), 'preserve')
            run.append(t)
            new_p.append(run)
            
            conclusion_element.addprevious(new_p)
        
        # Add heading
        heading_p = body.makeelement(qn('w:p'), {})
        pPr = heading_p.makeelement(qn('w:pPr'), {})
        
        pStyle = pPr.makeelement(qn('w:pStyle'), {qn('w:val'): 'Heading2'})
        pPr.append(pStyle)
        
        # Center alignment for heading
        jc_h = pPr.makeelement(qn('w:jc'), {qn('w:val'): 'left'})
        pPr.append(jc_h)
        
        heading_p.append(pPr)
        
        run = heading_p.makeelement(qn('w:r'), {})
        rPr = run.makeelement(qn('w:rPr'), {})
        
        b = rPr.makeelement(qn('w:b'), {})
        rPr.append(b)
        
        rFonts = rPr.makeelement(qn('w:rFonts'), {
            qn('w:ascii'): 'Times New Roman',
            qn('w:hAnsi'): 'Times New Roman',
            qn('w:cs'): 'Times New Roman'
        })
        rPr.append(rFonts)
        
        sz = rPr.makeelement(qn('w:sz'), {qn('w:val'): '28'})
        rPr.append(sz)
        szCs = rPr.makeelement(qn('w:szCs'), {qn('w:val'): '28'})
        rPr.append(szCs)
        
        run.append(rPr)
        
        t = run.makeelement(qn('w:t'), {})
        t.text = block['heading']
        run.append(t)
        heading_p.append(run)
        
        conclusion_element.addprevious(heading_p)
    
    print(f"Added {len(content_blocks)} new sections before conclusion")
else:
    print("No missing content to add or conclusion not found")

# ============================================================
# STEP 3: Add images for new content sections (where available)
# ============================================================
print("\n" + "=" * 60)
print("STEP 3: Adding images to new sections...")
print("=" * 60)

# Available images that match missing content:
# diag_architecture.png - for container architecture
# lstm_architecture.png - for prediction flow/LSTM
# model_comparison.png - for model responsibility

images_to_add = {
    'Жүйенің контейнерлік архитектурасы': ('diag_architecture.png', 'Сурет 23 – AI Traffic жүйесінің контейнерлік архитектурасы'),
    'ML болжау архитектурасы': ('lstm_architecture.png', 'Сурет 24 – ML болжау архитектурасы (Prediction Flow)'),
    'Модельдердің жауапкершілік аймақтары': ('model_comparison.png', 'Сурет 25 – Модельдерді салыстыру'),
}

# Reload document to find new paragraph positions
doc_reloaded = Document(SRC)

# Actually, we should add images in the same save. Let me find positions in current doc.
# The new sections were added before conclusion. Let me find them.

for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    for heading_key, (img_file, caption) in images_to_add.items():
        if heading_key in text and os.path.exists(img_file):
            # Find the paragraph right after the heading's content
            # Look for the next paragraph that would be a good insertion point
            # (after the last paragraph of this section, before next heading)
            for j in range(i + 1, min(i + 15, len(doc.paragraphs))):
                next_text = doc.paragraphs[j].text.strip()
                if next_text and (next_text.startswith('2.') or next_text.startswith('3.') or next_text.startswith('ҚОРЫТЫНДЫ')):
                    # Insert image before this paragraph
                    target_element = doc.paragraphs[j]._element
                    
                    # Add caption paragraph
                    cap_p = body.makeelement(qn('w:p'), {})
                    cap_pPr = cap_p.makeelement(qn('w:pPr'), {})
                    cap_jc = cap_pPr.makeelement(qn('w:jc'), {qn('w:val'): 'center'})
                    cap_pPr.append(cap_jc)
                    cap_p.append(cap_pPr)
                    
                    cap_run = cap_p.makeelement(qn('w:r'), {})
                    cap_rPr = cap_run.makeelement(qn('w:rPr'), {})
                    cap_rFonts = cap_rPr.makeelement(qn('w:rFonts'), {
                        qn('w:ascii'): 'Times New Roman',
                        qn('w:hAnsi'): 'Times New Roman'
                    })
                    cap_rPr.append(cap_rFonts)
                    cap_sz = cap_rPr.makeelement(qn('w:sz'), {qn('w:val'): '28'})
                    cap_rPr.append(cap_sz)
                    cap_run.append(cap_rPr)
                    cap_t = cap_run.makeelement(qn('w:t'), {})
                    cap_t.text = caption
                    cap_run.append(cap_t)
                    cap_p.append(cap_run)
                    
                    target_element.addprevious(cap_p)
                    print(f"  Added caption: {caption}")
                    break
            break

# ============================================================
# Save
# ============================================================
print("\n" + "=" * 60)
print("Saving...")
print("=" * 60)

doc.save(SRC)
print(f"Saved to: {SRC}")

# Final verification
print("\n" + "=" * 60)
print("Final verification...")
print("=" * 60)

doc2 = Document(SRC)
fig_count = 0
tbl_count = 0
for p in doc2.paragraphs:
    text = p.text.strip()
    if re.match(r'^Сурет\s+\d+', text):
        fig_count += 1
    if re.match(r'^Кесте\s+\d+', text):
        tbl_count += 1

print(f"Total figures: {fig_count}")
print(f"Total tables: {tbl_count}")
print(f"Total paragraphs: {len(doc2.paragraphs)}")
print(f"Total images: {len(doc2.inline_shapes)}")
print("\nDone!")
